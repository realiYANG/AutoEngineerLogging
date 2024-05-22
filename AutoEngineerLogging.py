# coding=utf-8
import os
import random
import shutil
import socket
import sys
import threading
import time
from datetime import datetime, timedelta
import warnings
from dateutil.parser import parse

import numpy as np
import openpyxl
import pandas as pd
import xlrd
import xlwings as xw
import xlwt
from xlutils.copy import copy
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import load_workbook
from PIL import Image
from PyQt5.Qt import QPoint, QPropertyAnimation, QEasingCurve, QAbstractAnimation
from PyQt5 import QtCore, QtGui, QtWidgets, QtNetwork
from PyQt5.QtCore import QDate, QTime, QBasicTimer, QDateTime, Qt, QTimer, QCoreApplication
from PyQt5.QtPrintSupport import QPageSetupDialog, QPrintDialog, QPrinter
from PyQt5.QtWidgets import (QApplication, QColorDialog, QDialog, QFileDialog,
                             QFontDialog, QLabel, QLineEdit, QMainWindow,
                             QMessageBox, QPushButton, QRadioButton,
                             QTableWidgetItem, QTextEdit, QWidget, QStyleFactory)
from CLASSES.ui_PROCESSING_CHAIN import Ui_Form
from CLASSES.ui_ENGINEER_LOGGING import Ui_MainWindow
from CLASSES.FTP_UP_DOWN_CLASS import MyFTP
from CLASSES.EMITTINGSTR_CLASS import EmittingStr
from CLASSES.SUPERVISOR_BY_EMAIL_CLASS import Supervisor

warnings.filterwarnings('ignore', category=DeprecationWarning)

class Main_window(QMainWindow, Ui_MainWindow):
    def __init__(self):
        super(Main_window, self).__init__()
        self.setupUi(self)
        self.statusBar().showMessage('Ready')
        # self.setWindowOpacity(1.0)
        self.setObjectName("mainWindow")
        # qss = "QMainWindow#mainWindow{background-color:black;}"

        randon_num = random.randint(1, 100)
        if 1 <= randon_num < 5:
            qss = "QMainWindow#mainWindow{border-image:url(./resources/image/Background.jpg);}"
            self.setStyleSheet(qss)
        else:
            pass
        self.lock = threading.Lock()  # 数据锁

        # Release OR Debug 版本切换控制
        # TODO
        # 将控制台输出重定向到textBrowser中
        # sys.stdout = EmittingStr(textWritten=self.outputWritten)
        # sys.stderr = EmittingStr(textWritten=self.outputWritten)

        # 许可时间
        # TODO
        end_license_time = '2024-09-01 12:00:00'
        print('当前时间是: ' + self.now() + '\n' + '本版本使用期限为: ' + end_license_time + '，建议届时更新:)')

        # 网络版开关
        '''
        reply = QMessageBox.question(self, '联网提示',
                                     "选择 Yes则网络正式版 No则非网络试用版", QMessageBox.Yes |
                                     QMessageBox.No, QMessageBox.No)
        '''
        ########### 强制选择 Yes则网络版 No则非网络版
        reply = QMessageBox.No
        ###########
        if reply == QMessageBox.Yes:
            '''
            # 设置内网代理
            proxy = QtNetwork.QNetworkProxy()
            proxy.setType(QtNetwork.QNetworkProxy.HttpProxy)
            proxy.setHostName("https://10.22.19.21")
            proxy.setPort(8080)
            QtNetwork.QNetworkProxy.setApplicationProxy(proxy)
            '''
            self.run_on_net = True
        else:
            self.run_on_net = False
            pass

        if self.run_on_net == True:
            try:
                Supervisor.usage_supervisor()
                error = False
            except:
                error = True
                QMessageBox.information(self, "提示", "网络连接失败，请确认网络是否连接正常")
                input('网络连接失败，请确认网络是否连接正常')

            # with open('.\\resources\\延期码.txt', "r") as f:
            #     license_str = f.read()
            # if reply == QMessageBox.Yes:  # 网络正式版
            #     end_license_time = '3000-01-01 12:00:00'
            # elif reply == QMessageBox.No:  # 非网络试用版
            #     if 'yang' in license_str:
            #         end_license_time = '2021-05-01 12:00:00'
            #     else:
            #         end_license_time = '2020-10-31 12:00:00'

            if self.now() > end_license_time and error == True:
                print('模块需要升级，请联系软件开发人员')
                QMessageBox.information(self, "提示", "模块需要升级，请联系研发人员，电话：18580367621")
            elif self.now() > end_license_time and error == False:
                print('模块需要升级，请联系软件开发人员')
                QMessageBox.information(self, "提示", "模块需要升级，请联系研发人员，电话：18580367621")
            elif self.now() <= end_license_time and error == True:
                print('模块无需升级，网络连接失败，请确认网络是否连接正常')
                QMessageBox.information(self, "提示", "模块无需升级，网络连接失败，请确认网络是否连接正常")
            elif self.now() <= end_license_time and error == False:
                print('模块已经更新，可以正常使用')
                self.main_initialization()
        else:
            if self.now() > end_license_time:
                print('模块需要升级，请联系软件开发人员')
                QMessageBox.information(self, "提示", "模块需要升级，请联系研发人员，电话：18580367621")
            elif self.now() <= end_license_time:
                self.main_initialization()
            else:
                pass
            pass

    def main_initialization(self):
        # 防止上传不能保存空文件夹的bug
        dir1_path = '.\\WorkSpace\\报告生成工区\\原始资料'
        dir2_path = '.\\WorkSpace\\报告生成工区\\成果表'
        dir3_path = '.\\WorkSpace\\报告生成工区\\储层表'
        dir4_path = '.\\WorkSpace\\报告生成工区\\储层图'
        dir5_path = '.\\WorkSpace\\报告生成工区\\胶结差图'
        dir6_path = '.\\WorkSpace\\分层和成果表工区'
        dir7_path = '.\\WorkSpace\\合并统计工区'
        dir8_path = '.\\WorkSpace'
        dir_paths = [dir1_path, dir2_path, dir3_path, dir4_path, dir5_path, dir6_path, dir7_path, dir8_path]
        for item in dir_paths:
            if not os.path.exists(item):
                os.makedirs(item)
                print(item, ' 已创建。')

        # 水泥胶结评价模块初始化
        ###################################################
        self.pushButton.clicked.connect(self.read_raw_info_docx)
        self.pushButton.setToolTip('请确保原始记录登记表格式正确')
        self.pushButton_53.clicked.connect(self.automate_table_helper)  # 智能补充解释评价顶底深度
        self.pushButton_53.setToolTip('成果表放置后该按钮才有用')
        self.pushButton_3.clicked.connect(self.generate_txt_file)
        self.pushButton_3.setToolTip('生成可导入LEAD4.0的TXT井信息文件')
        self.pushButton_67.clicked.connect(self.result_table_process_in_report_module)  # 成果表规范化
        self.pushButton_4.clicked.connect(self.generate_report_thread)
        self.pushButton_4.clicked.connect(self.progressbar_action_thread)
        self.pushButton_5.clicked.connect(self.clean_report_workspace)  # 清理报告生成工区目录（除了result，因为整理后的会输出到这个目录）
        self.pushButton_31.clicked.connect(self.clean_report_workspace_all)  # 清理报告生成工区目录
        self.pushButton_40.clicked.connect(self.clean_workspace_all)  # 清理所有工区目录
        self.pushButton_48.clicked.connect(self.open_report_workspace_directory)  # 打开报告生成工区
        self.pushButton_41.clicked.connect(self.open_result_table_directory)  # 打开成果表文件夹
        self.pushButton_42.clicked.connect(self.open_formation_table_directory)  # 打开储层表文件夹
        self.pushButton_47.clicked.connect(self.open_formation_pictures_directory)  # 打开储层图文件夹
        self.pushButton_46.clicked.connect(self.open_bad_cement_pictures_directory)  # 打开胶结差图文件夹
        self.pushButton_49.clicked.connect(self.flush_on_textEdits)  # 在textEdit上刷新显示

        # 钻头数据表和套管数据表初始化
        self.bit_info_table()
        self.casing_info_table()

        self.comboBox_8.addItems(['合格', '无连续25m', '不确定', '/'])
        self.comboBox_8.setCurrentText('/')

        self.comboBox_9.addItems(['合格', '无连续25m', '不确定', '/'])
        self.comboBox_9.setCurrentText('/')

        self.comboBox_11.addItems(['合格', '无连续25m', '不确定', '/'])
        self.comboBox_11.setCurrentText('/')

        self.comboBox_12.addItems(['合格', '无连续25m', '不确定', '/'])
        self.comboBox_12.setCurrentText('/')

        choices_list1 = ["李柯沁", "刘佳露", "杨玉竹", "杨晨曦", "周政英", "闫跃星", "赵晓军", "王遂华", "/"]
        self.comboBox_2.addItems(choices_list1)
        self.comboBox_2.setCurrentText('/')

        choices_list2 = ["李柯沁", "刘佳露", "杨玉竹", "杨晨曦", "周政英", "闫跃星", "赵晓军", "王遂华", "朱莉", "王昌德", "王参文", "刘静", "李海军", "/"]
        self.comboBox.addItems(choices_list2)
        self.comboBox.setCurrentText('/')

        choices_list3 = ["朱莉", "王昌德", "王参文", "刘静", "李海军", "/"]
        self.comboBox_4.addItems(choices_list3)
        self.comboBox_4.setCurrentText('/')

        self.all_Types_Simple = ['直井', '定向井', '大斜度井', '水平井', '分支井', '其它类型井']
        self.comboBox_6.addItems(self.all_Types_Simple)
        self.comboBox_6.setCurrentText('/')

        self.all_Categories_Simple = ['探井', '开发井', '注入井', '地热井', '预探井', '科探井', '报废井', '储气库井', '其它井']
        self.comboBox_7.addItems(self.all_Categories_Simple)
        self.comboBox_7.setCurrentText('/')

        self.logging_Equipments = ['ECLIPS-5700', 'MAXIS-500', 'LOGIQ', 'HH2530', 'SKD-3000', 'MCET-1000', 'CPLog', '其他', '/']
        self.comboBox_5.addItems(self.logging_Equipments)
        self.comboBox_5.setCurrentText('/')

        # 进度条
        self.step = 0
        self.count = 0  # LCD显示的数字
        self.progressBar.setRange(0, 100)
        self.progressBar.setValue(0)
        self.timer = QBasicTimer()

        # 动态显示时间在label上
        timer_of_time = QTimer(self)
        timer_of_time.timeout.connect(self.showtime)
        timer_of_time.start()
        ###################################################

        # 套损评价模块初始化
        ###################################################
        # self.id = 1
        # self.lines = []
        self.editable = True
        self.des_sort = True
        self.table2()
        self.table3()
        self.table4()
        self.table_casing()

        # self.dateTimeEdit()
        # self.dateTimeEdit_2()

        # 打开解析LAS文件
        self.pushButton_38.clicked.connect(self.open_las_file)
        self.pushButton_34.clicked.connect(self.read_las_file)

        # 损伤评价按钮组
        self.pushButton_11.clicked.connect(self.add_line_for_tableWidget_2)
        self.pushButton_12.clicked.connect(self.delete_line_for_tableWidget_2)
        self.pushButton_13.clicked.connect(self.generate_results_from_tableWidget_2)

        # 结垢评价按钮组
        self.pushButton_17.clicked.connect(self.add_line_for_tableWidget_3)
        self.pushButton_18.clicked.connect(self.delete_line_for_tableWidget_3)
        self.pushButton_19.clicked.connect(self.generate_results_from_tableWidget_3)

        # 变形评价按钮组
        self.pushButton_20.clicked.connect(self.add_line_for_tableWidget_4)
        self.pushButton_21.clicked.connect(self.delete_line_for_tableWidget_4)
        self.pushButton_22.clicked.connect(self.generate_results_from_tableWidget_4)

        # 增加空行
        self.pushButton_64.clicked.connect(self.add_blank_line_for_tableWidget_2)
        self.pushButton_63.clicked.connect(self.add_blank_line_for_tableWidget_3)
        self.pushButton_62.clicked.connect(self.add_blank_line_for_tableWidget_4)

        # 填写确认生成本地Excel文件
        self.pushButton_35.clicked.connect(self.add_line_for_tableWidget_5)
        self.pushButton_36.clicked.connect(self.delete_line_for_tableWidget_5)
        self.pushButton_37.clicked.connect(self.casing_info_save)

        # 清理目录
        self.pushButton_24.clicked.connect(self.clean_the_dir)

        # 测量项目checkbox的状态
        self.checkBox_6.setChecked(True)  # 默认为MIT24
        self.checkBox_6.stateChanged.connect(self.changecb_type1)
        self.checkBox_7.stateChanged.connect(self.changecb_type2)
        self.checkBox_8.stateChanged.connect(self.changecb_type3)
        self.checkBox_12.stateChanged.connect(self.changecb_type4)

        # 类型checkbox的状态
        self.checkBox_4.stateChanged.connect(self.changecb1)
        self.checkBox.stateChanged.connect(self.changecb2)
        self.checkBox_2.stateChanged.connect(self.changecb2)
        self.checkBox_3.stateChanged.connect(self.changecb2)

        # 生成快速解释结论
        self.pushButton_23.clicked.connect(self.generate_fast_report)
        ###################################################

        # 图片添加签名模块初始化
        ###################################################
        self.radioButton_5.toggled.connect(lambda: self.btnstate(self.radioButton_5))
        self.radioButton_6.toggled.connect(lambda: self.btnstate(self.radioButton_6))
        self.pushButton_25.clicked.connect(self.open_picture_file)
        self.pushButton_30.clicked.connect(self.reset_add_signature)
        ###################################################

        # 图片转PDF模块初始化
        ###################################################
        self.pushButton_28.clicked.connect(self.open_picture_file_to_pdf)
        self.pushButton_32.clicked.connect(self.convert_picture_to_pdf)
        ###################################################

        # 成果表の小工具集
        ###################################################
        self.radioButton.toggled.connect(lambda: self.btnstate_table(self.radioButton))
        self.radioButton_2.toggled.connect(lambda: self.btnstate_table(self.radioButton_2))

        self.pushButton_9.clicked.connect(self.open_file1)
        self.pushButton_10.clicked.connect(self.open_file2)
        self.pushButton_54.clicked.connect(self.open_file3)
        self.pushButton_55.clicked.connect(self.open_file4)
        self.pushButton_59.clicked.connect(self.open_file5)
        self.pushButton_60.clicked.connect(self.open_file6)

        self.pushButton_29.clicked.connect(self.reset_table_process)
        self.pushButton_52.clicked.connect(self.open_table_process_directory)  # 打开合并统计工区文件夹
        self.pushButton_56.clicked.connect(self.open_table_fusion_directory)  # 打开综合评价工区文件夹

        # 生成综合评价表
        self.pushButton_57.clicked.connect(self.table_process3)  # 规范化
        self.pushButton_58.clicked.connect(self.table_fusion_reaction)

        self.action.triggered.connect(self.menubar_simple_instruction)
        self.action_2.triggered.connect(self.menubar_author_info)

        # 查询好中差比例小工具
        self.pushButton_66.clicked.connect(self.table_process4)
        self.pushButton_61.clicked.connect(self.search_for_statistic_result)
        ###################################################

        # list文件拼接模块初始化
        ###################################################
        self.pushButton_43.clicked.connect(self.open_list_file1)
        self.pushButton_45.clicked.connect(self.open_list_file2)
        self.pushButton_44.clicked.connect(self.assemble_list_together)
        ###################################################

        # 分层表和储层表整理模块初始化
        ###################################################
        self.pushButton_7.clicked.connect(self.select_layer_table)
        self.pushButton_2.clicked.connect(self.layer_table_process)
        self.pushButton_6.clicked.connect(self.select_formation_table)
        self.pushButton_8.clicked.connect(self.formation_table_process)
        self.pushButton_50.clicked.connect(self.open_layer_result_directory)  # 打开分层和成果表工区文件夹
        self.pushButton_51.clicked.connect(self.open_layer_result_directory)  # 打开分层和成果表工区文件夹
        ###################################################

        # 解释评价顶底深度编辑后发送信号
        ###################################################
        self.lineEdit_103.editingFinished.connect(self.set_well_detail_name)
        self.lineEdit_105.editingFinished.connect(self.set_well_detail_name)
        ###################################################

        # 报告生成进程管控
        ###################################################
        self.checkBox_9.setChecked(True)  # 字段替换
        self.checkBox_10.setChecked(True)  # 表格添加
        self.checkBox_11.setChecked(True)  # 表格调整
        ###################################################

        # 接箍对比表自动整理
        ###################################################
        self.pushButton_33.clicked.connect(self.select_ccl_table)
        self.pushButton_39.clicked.connect(self.organize_ccl_table)
        ###################################################

        # 菜单actions
        ###################################################
        self.action_2.triggered.connect(self.menubar_simple_instruction)
        self.action.triggered.connect(self.menubar_author_info)
        ###################################################

    # 解析混合日期的函数
    def parse_mixed_datetime(self, date_string):
        try:
            parsed_date = parse(date_string, fuzzy=True)
            return parsed_date
        except ValueError:
            # Handle invalid date strings here (if needed)
            return None

    # 在textBrowser中显示程序运行状态
    def outputWritten(self, text):
        cursor = self.textBrowser.textCursor()
        cursor.movePosition(QtGui.QTextCursor.End)
        cursor.insertText(text)
        # self.textBrowser.setTextCursor(cursor)
        # self.textBrowser.ensureCursorVisible()

    def menubar_simple_instruction(self):
        QMessageBox.information(self, "简介",
                                "工程测井工作助手主要针对工程所生产过程中的LEAD固井质量处理、套损检测中的工作实现自动化的实现，能够有效地提升生产效率和规范报告图件")

    def menubar_author_info(self):
        QMessageBox.information(self, "联系方式",
                                "软件开发: 杨艺  软件测试: 刘恒 王参文 何强\n电话：18580367621，邮箱：978030836@qq.com")

    # 添加一个计时器事件
    def timerEvent(self, e):
        if self.step >= 100:
            self.step = 0
            self.progressBar.setValue(self.step)
            # self.timer.stop()
            self.count += 1
            self.lcdNumber.display(self.count)
            return
        self.step = self.step + 1
        self.progressBar.setValue(self.step)

    def progressbar_action_thread(self):
        self.timer.start(10, self)
        progressbar_action = threading.Thread(target=self.progressbar_action)
        progressbar_action.start()

    def progressbar_action(self):
        if self.timer.isActive():
            self.pushButton_4.setText('请等待……')
            self.pushButton_4.setEnabled(False)
        else:
            pass

    def select_ccl_table(self):
        fnames = QFileDialog.getOpenFileNames(self, '打开文件', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.lineEdit_58.setText(fname)

    def organize_ccl_table(self):
        PATH = self.lineEdit_58.text()
        workbook = load_workbook(PATH)
        print(workbook.sheetnames)
        sheet = workbook[workbook.sheetnames[0]]
        print(sheet.dimensions)

        # cell1 = sheet.cell(row=1, column=1)
        # cell2 = sheet.cell(row=11, column=3)
        # print(cell1.value, cell2.value)

        # 获取该表相应的行数和列数
        nrows = sheet.max_row
        ncols = sheet.max_column
        print(nrows, ncols)

        # 将表格中的None替换为空字符串
        for row in range(1, nrows + 1):
            for col in range(1, ncols + 1):
                if sheet.cell(row, col).value == None:
                    sheet.cell(row, col).value = ''

        noise = random.random() * 0.02  # 随机数
        ##########################################################第一次处理
        # 定义偏移列数
        add_number1 = 9

        # 将表格复制到后面列
        for row in range(1, nrows + 1):
            for col in range(1, ncols + 1):
                sheet.cell(row, col + add_number1).value = sheet.cell(row, col).value

        # 修正异常长度导致的深度误差
        for row in range(4, nrows):  # 为了不溢出，只循环到倒数第二行；从4开始，是为了避免row-2等于第一行文字
            if sheet.cell(row, 5).value != '':
                if float(sheet.cell(row, 5).value) < 0:
                    sheet.cell(row - 1, 1 + add_number1).value = round(
                        float(sheet.cell(row - 2, 1 + add_number1).value) + float(
                            sheet.cell(row - 2, 4 + add_number1).value) + noise, 3)
                    sheet.cell(row, 1 + add_number1).value = round(
                        float(sheet.cell(row - 1, 1 + add_number1).value) + float(
                            sheet.cell(row - 1, 4 + add_number1).value) + noise, 3)
                    sheet.cell(row + 1, 1 + add_number1).value = round(
                        float(sheet.cell(row, 1 + add_number1).value) + float(
                            sheet.cell(row, 4 + add_number1).value) + noise, 3)
                elif float(sheet.cell(row, 5).value) >= 0:
                    sheet.cell(row - 1, 1 + add_number1).value = round(
                        float(sheet.cell(row - 2, 1 + add_number1).value) + float(
                            sheet.cell(row - 2, 4 + add_number1).value) + noise, 3)
                    sheet.cell(row, 1 + add_number1).value = round(
                        float(sheet.cell(row - 1, 1 + add_number1).value) + float(
                            sheet.cell(row - 1, 4 + add_number1).value) + noise, 3)
                    sheet.cell(row + 1, 1 + add_number1).value = round(
                        float(sheet.cell(row, 1 + add_number1).value) + float(
                            sheet.cell(row, 4 + add_number1).value) + noise, 3)

        # 第一次处理后的长度误差
        for row in range(2, nrows):  # 为了不溢出，只循环到倒数第二行
            sheet.cell(row, 2 + add_number1).value = round(
                float(sheet.cell(row + 1, 1 + add_number1).value) - float(sheet.cell(row, 1 + add_number1).value), 3)
            error1 = float(sheet.cell(row + 1, 1 + add_number1).value) - float(
                sheet.cell(row, 1 + add_number1).value) - float(sheet.cell(row, 4 + add_number1).value)
            error1 = round(error1, 3)
            if abs(error1) > 0.2:
                sheet.cell(row, 5 + add_number1).value = error1
            else:
                sheet.cell(row, 5 + add_number1).value = ''

        # 第一次处理后的深度误差
        for row in range(2, nrows + 1):
            error2 = float(sheet.cell(row, 1 + add_number1).value) - float(sheet.cell(row, 3 + add_number1).value)
            error2 = round(error2, 3)
            if abs(error2) > 0.2:
                sheet.cell(row, 6 + add_number1).value = error2
            else:
                sheet.cell(row, 6 + add_number1).value = ''

        # workbook.save('Output.xlsx')

        ##########################################################第二次处理
        # 定义偏移列数
        add_number2 = 9

        # 将表格复制到后面列
        for row in range(1, nrows + 1):
            for col in range(1 + add_number1, ncols + 1 + add_number1):
                sheet.cell(row, col + add_number2).value = sheet.cell(row, col).value

        # 修正异常长度导致的深度误差
        for row in range(4, nrows):  # 为了不溢出，只循环到倒数第二行
            if sheet.cell(row, 5 + add_number1).value != '':
                if float(sheet.cell(row, 5 + add_number1).value) < 0:
                    sheet.cell(row - 1, 1 + add_number1 + add_number2).value = round(
                        float(sheet.cell(row - 2, 1 + add_number1 + add_number2).value) + float(
                            sheet.cell(row - 2, 4 + add_number1 + add_number2).value) - noise, 3)
                    sheet.cell(row, 1 + add_number1 + add_number2).value = round(
                        float(sheet.cell(row - 1, 1 + add_number1 + add_number2).value) + float(
                            sheet.cell(row - 1, 4 + add_number1 + add_number2).value) + noise, 3)
                    sheet.cell(row + 1, 1 + add_number1 + add_number2).value = round(
                        float(sheet.cell(row, 1 + add_number1 + add_number2).value) + float(
                            sheet.cell(row, 4 + add_number1 + add_number2).value) - noise, 3)
                elif float(sheet.cell(row, 5 + add_number1).value) >= 0:
                    sheet.cell(row - 1, 1 + add_number1 + add_number2).value = round(
                        float(sheet.cell(row - 2, 1 + add_number1 + add_number2).value) + float(
                            sheet.cell(row - 2, 4 + add_number1 + add_number2).value) + noise, 3)
                    sheet.cell(row, 1 + add_number1 + add_number2).value = round(
                        float(sheet.cell(row - 1, 1 + add_number1 + add_number2).value) + float(
                            sheet.cell(row - 1, 4 + add_number1 + add_number2).value) - noise, 3)
                    sheet.cell(row + 1, 1 + add_number1 + add_number2).value = round(
                        float(sheet.cell(row, 1 + add_number1 + add_number2).value) + float(
                            sheet.cell(row, 4 + add_number1 + add_number2).value) + noise, 3)

        # 第二次处理后的长度误差
        for row in range(2, nrows):  # 为了不溢出，只循环到倒数第二行
            sheet.cell(row, 2 + add_number1 + add_number2).value = round(
                float(sheet.cell(row + 1, 1 + add_number1 + add_number2).value) - float(
                    sheet.cell(row, 1 + add_number1 + add_number2).value), 3)
            error = float(sheet.cell(row + 1, 1 + add_number1 + add_number2).value) - float(
                sheet.cell(row, 1 + add_number1 + add_number2).value) - float(
                sheet.cell(row, 4 + add_number1 + add_number2).value)
            error = round(error, 3)
            if abs(error) > 0.2:
                sheet.cell(row, 5 + add_number1 + add_number2).value = error
            else:
                sheet.cell(row, 5 + add_number1 + add_number2).value = ''

        # 第二次处理后的深度误差
        for row in range(2, nrows + 1):
            error2 = float(sheet.cell(row, 1 + add_number1 + add_number2).value) - float(
                sheet.cell(row, 3 + add_number1 + add_number2).value)
            error2 = round(error2, 3)
            if abs(error2) > 0.2:
                sheet.cell(row, 6 + add_number1 + add_number2).value = error2
            else:
                sheet.cell(row, 6 + add_number1 + add_number2).value = ''

        temp = PATH.split('/')[-1]
        PATH = PATH.replace(temp, '')
        new_path = ''.join([PATH, '深度移动后的套管接箍对比表.xlsx'])
        workbook.save(new_path)

        QMessageBox.information(self, "提示", "深度移动完毕，请查看接箍对比表的同级目录")

    def set_well_detail_name(self):
        well_Name = self.lineEdit.text()
        start_Evaluation = self.lineEdit_103.text()
        end_Evaluation = self.lineEdit_105.text()
        logging_Date = self.lineEdit_40.text().replace('-', '')

        if '.' not in start_Evaluation:
            start_Evaluation_1_digits = ''.join([start_Evaluation, '.0'])
            start_Evaluation_2_digits = ''.join([start_Evaluation, '.00'])
        elif '.0' in start_Evaluation and '.00' not in start_Evaluation:
            start_Evaluation_1_digits = start_Evaluation
            start_Evaluation_2_digits = start_Evaluation.replace('.0', '.00')
        elif '.00' in start_Evaluation:
            start_Evaluation_1_digits = start_Evaluation.replace('.00', '.0')
            start_Evaluation_2_digits = start_Evaluation
        elif '.5' in start_Evaluation and '.50' not in start_Evaluation:
            start_Evaluation_1_digits = start_Evaluation
            start_Evaluation_2_digits = start_Evaluation.replace('.5', '.50')
        else:
            start_Evaluation_1_digits = start_Evaluation
            start_Evaluation_2_digits = start_Evaluation

        if '.' not in end_Evaluation:
            end_Evaluation_1_digits = ''.join([end_Evaluation, '.0'])
            end_Evaluation_2_digits = ''.join([end_Evaluation, '.00'])
        elif '.0' in end_Evaluation and '.00' not in end_Evaluation:
            end_Evaluation_1_digits = end_Evaluation
            end_Evaluation_2_digits = end_Evaluation.replace('.0', '.00')
        elif '.00' in end_Evaluation:
            end_Evaluation_1_digits = end_Evaluation.replace('.00', '.0')
            end_Evaluation_2_digits = end_Evaluation
        elif '.5' in end_Evaluation and '.50' not in end_Evaluation:
            end_Evaluation_1_digits = end_Evaluation
            end_Evaluation_2_digits = end_Evaluation.replace('.5', '.50')
        else:
            end_Evaluation_1_digits = end_Evaluation
            end_Evaluation_2_digits = end_Evaluation

        start_Evaluation = start_Evaluation.replace('.00', '').replace('.0', '').replace('.50', '.5')
        end_Evaluation = end_Evaluation.replace('.00', '').replace('.0', '').replace('.50', '.5')
        # 新规定，起始评价深度从0米开始
        # TODO
        if float(start_Evaluation) < 200:
            start_Evaluation = '0'
        well_Times_Name = ''.join([well_Name, '_固井质量_', start_Evaluation, '-', end_Evaluation, '_', logging_Date])
        self.lineEdit_3.setText(well_Times_Name)

        # self.lineEdit_107.setText(start_Evaluation_1_digits)
        self.lineEdit_103.setText(start_Evaluation_2_digits)
        self.lineEdit_105.setText(end_Evaluation_2_digits)

    def select_layer_table(self):
        fnames = QFileDialog.getOpenFileNames(self, '打开分层表', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.lineEdit_6.setText(fname)

    def layer_table_process(self):
        fileDirLayer = self.lineEdit_6.text()
        wb1 = xlrd.open_workbook(fileDirLayer)
        wb1_openpyxl = openpyxl.load_workbook(fileDirLayer)
        ##############################################
        # 处理Layer表
        sheet1 = wb1.sheets()[0]
        nrow1 = sheet1.nrows
        ncol1 = sheet1.ncols
        delete_Row_V = []
        delete_Row_H = []

        # 用openpyxl进行处理
        sheet1_openpyxl = wb1_openpyxl[wb1_openpyxl.sheetnames[0]]
        for row in range(nrow1):
            for col in range(ncol1):
                if sheet1.cell_value(row, col) == '垂直定位':
                    delete_Row_V.append(row)
                else:
                    pass
        if delete_Row_V != []:
            for i in range(len(delete_Row_V)):
                sheet1_openpyxl.delete_rows(delete_Row_V[i] + 1)
                delete_Row_V = (np.array(delete_Row_V) - 1).tolist()  # 所有元素减1

        for row in range(nrow1):
            for col in range(ncol1):
                if sheet1.cell_value(row, col) == '水平定位':
                    delete_Row_H.append(row)
                else:
                    pass

        if delete_Row_H != []:
            for i in range(len(delete_Row_H)):
                sheet1_openpyxl.delete_rows(delete_Row_H[i] + 1)
                delete_Row_H = (np.array(delete_Row_H) - 1).tolist()  # 所有元素减1

        # deleterows(sheet1_openpyxl, delete_Row + 1)#openpyxl中数行数从1开始
        sheet1_openpyxl['C2'] = None

        for row in range(1, nrow1 + 1):
            for col in range(1, ncol1):
                if sheet1_openpyxl[row][col].value == '龙一2':
                    sheet1_openpyxl[row][col].value = '龙一^2'
                elif sheet1_openpyxl[row][col].value == '龙一１1':
                    sheet1_openpyxl[row][col].value = '龙一^１1'
                elif sheet1_openpyxl[row][col].value == '龙一１2':
                    sheet1_openpyxl[row][col].value = '龙一^１2'
                elif sheet1_openpyxl[row][col].value == '龙一１3':
                    sheet1_openpyxl[row][col].value = '龙一^１3'
                elif sheet1_openpyxl[row][col].value == '龙一１4':
                    sheet1_openpyxl[row][col].value = '龙一^１4'
                elif sheet1_openpyxl[row][col].value == '龙一11':
                    sheet1_openpyxl[row][col].value = '龙一^11'
                elif sheet1_openpyxl[row][col].value == '龙一12':
                    sheet1_openpyxl[row][col].value = '龙一^12'
                elif sheet1_openpyxl[row][col].value == '龙一13':
                    sheet1_openpyxl[row][col].value = '龙一^13'
                elif sheet1_openpyxl[row][col].value == '龙一14':
                    sheet1_openpyxl[row][col].value = '龙一^14'
                elif sheet1_openpyxl[row][col].value == '嘉二1':
                    sheet1_openpyxl[row][col].value = '嘉二^1'
                elif sheet1_openpyxl[row][col].value == '嘉二2':
                    sheet1_openpyxl[row][col].value = '嘉二^2'
                elif sheet1_openpyxl[row][col].value == '嘉二3':
                    sheet1_openpyxl[row][col].value = '嘉二^3'
                elif sheet1_openpyxl[row][col].value == '嘉四1':
                    sheet1_openpyxl[row][col].value = '嘉四^1'
                elif sheet1_openpyxl[row][col].value == '嘉四2':
                    sheet1_openpyxl[row][col].value = '嘉四^2'
                elif sheet1_openpyxl[row][col].value == '嘉四3':
                    sheet1_openpyxl[row][col].value = '嘉四^3'
                elif sheet1_openpyxl[row][col].value == '嘉四4':
                    sheet1_openpyxl[row][col].value = '嘉四^4'
                elif sheet1_openpyxl[row][col].value == '嘉五1':
                    sheet1_openpyxl[row][col].value = '嘉五^1'
                elif sheet1_openpyxl[row][col].value == '嘉五2':
                    sheet1_openpyxl[row][col].value = '嘉五^2'
                elif sheet1_openpyxl[row][col].value == '雷一1':
                    sheet1_openpyxl[row][col].value = '雷一^1'
                elif sheet1_openpyxl[row][col].value == '雷一2':
                    sheet1_openpyxl[row][col].value = '雷一^2'
                elif sheet1_openpyxl[row][col].value == '雷三1':
                    sheet1_openpyxl[row][col].value = '雷三^1'
                elif sheet1_openpyxl[row][col].value == '雷三2':
                    sheet1_openpyxl[row][col].value = '雷三^2'
                elif sheet1_openpyxl[row][col].value == '雷三3':
                    sheet1_openpyxl[row][col].value = '雷三^3'
                elif sheet1_openpyxl[row][col].value == '嘉三3':
                    sheet1_openpyxl[row][col].value = '嘉三^3'
                elif sheet1_openpyxl[row][col].value == '嘉三2':
                    sheet1_openpyxl[row][col].value = '嘉三^2'
                elif sheet1_openpyxl[row][col].value == '嘉三1':
                    sheet1_openpyxl[row][col].value = '嘉三^1'
                else:
                    pass
        wb1_openpyxl.save('.\\WorkSpace\\分层和成果表工区\\Layer_整理后.xlsx')
        QMessageBox.information(self, "提示", "整理完毕，请到目录中查看")

    def select_formation_table(self):
        fnames = QFileDialog.getOpenFileNames(self, '打开储层表', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.lineEdit_2.setText(fname)

    def formation_table_process(self):
        fileDirResult = self.lineEdit_2.text()
        wb2 = xlrd.open_workbook(fileDirResult)
        wb2_openpyxl = openpyxl.load_workbook(fileDirResult)
        ##############################################
        # 处理Result表
        sheet2 = wb2.sheets()[0]
        nrow2 = sheet2.nrows
        ncol2 = sheet2.ncols

        # 用openpyxl进行处理
        sheet2_openpyxl = wb2_openpyxl[wb2_openpyxl.sheetnames[0]]
        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '自然伽马':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '补偿声波':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '补偿密度':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '电阻率':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '孔隙度':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '含水饱和度':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '有机碳含量':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '总含气量':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '自然伽玛':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '补偿中子':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '深侧向':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '浅侧向':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '深感应':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '浅感应':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '吸附气含量':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '游离气含量':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)

        for row in range(1, sheet2_openpyxl.max_row):
            for col in range(1, sheet2_openpyxl.max_column):
                if sheet2_openpyxl[row][col - 1].value == '储层厚度':
                    delete_Col = col
                    sheet2_openpyxl.delete_cols(delete_Col)
        wb2_openpyxl.save('.\\WorkSpace\\分层和成果表工区\\Result_报告格式(需转xls).xlsx')

        # insert column
        sheet2_openpyxl.insert_cols(4)
        sheet2_openpyxl.insert_cols(4)
        for row in range(3, sheet2_openpyxl.max_row + 1):
            sheet2_openpyxl[row][3].value = sheet2_openpyxl[row][2].value.split('--')[1]
            sheet2_openpyxl[row][2].value = sheet2_openpyxl[row][2].value.split('--')[0]
            sheet2_openpyxl[row][4].value = sheet2_openpyxl[row][5].value
        wb2_openpyxl.save('.\\WorkSpace\\分层和成果表工区\\Result_整理后.xlsx')
        QMessageBox.information(self, "提示", "整理完毕，请到目录中查看")

    def open_layer_result_directory(self):
        path = '.\\WorkSpace\\分层和成果表工区'
        os.startfile(path)

    ################################################################################## 添加签名模块
    def open_picture_file_to_pdf(self):
        fnames = QFileDialog.getOpenFileNames(self, '打开图片文件', './')  # 注意这里返回值是元组
        if fnames[0]:
            fname = '|'.join(fnames[0])
            self.lineEdit_57.setText(fname)

    def convert_picture_to_pdf(self):
        paths = self.lineEdit_57.text()
        path = paths.split('|')
        count = 0
        for a_path in path:
            img = Image.open(a_path)
            # 转pdf
            path_without_suffix = a_path.split('.')[0]
            img.save(path_without_suffix + '.pdf', "PDF", resolution=300.0, save_all=True)
            count += 1
        info = ''.join(["共", str(count), "个图片格式转换成功，请到源文件所在目录中查看"])
        QMessageBox.information(self, "提示", info)

    ################################################################################## 添加签名模块
    def open_picture_file(self):
        fnames = QFileDialog.getOpenFileNames(self, '打开图片文件', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.lineEdit_56.setText(fname)

    def btnstate(self, btn):
        # 输出按钮1与按钮2的状态，选中还是没选中
        if btn.text() == '100分辨率，宽度952像素':
            if btn.isChecked() == True:
                print(btn.text() + " 被选中")
                self.pushButton_26.clicked.connect(self.add_signature_on_pic_100)
            else:
                pass

        if btn.text() == "150分辨率，宽度1427像素":
            if btn.isChecked() == True:
                print(btn.text() + " 被选中")
                self.pushButton_26.clicked.connect(self.add_signature_on_pic_150)
            else:
                pass

    def reset_add_signature(self):
        try:
            self.radioButton_5.toggled.disconnect(lambda: self.btnstate(self.radioButton_5))
        except:
            print('Error1')
        else:
            print('Disconnected1')

        try:
            self.radioButton_6.toggled.disconnect(lambda: self.btnstate(self.radioButton_6))
        except:
            print('Error2')
        else:
            print('Disconnected2')

        try:
            self.pushButton_26.clicked.disconnect(self.add_signature_on_pic_100)
            self.pushButton_26.clicked.disconnect(self.add_signature_on_pic_150)
        except:
            print('Error3')
        else:
            print('Disconnected3')
        self.radioButton_5.toggled.connect(lambda: self.btnstate(self.radioButton_5))
        self.radioButton_6.toggled.connect(lambda: self.btnstate(self.radioButton_6))

    def addImg_100(self, img):
        # 审核
        mark1 = Image.open('.\\resources\\签名\\签名-' + self.comboBox_4.currentText() + '.jpg')
        mark1 = mark1.resize((90, 40), Image.ANTIALIAS)
        img.paste(mark1, (800, 1014))
        # 校对
        mark2 = Image.open('.\\resources\\签名\\签名-' + self.comboBox.currentText() + '.jpg')
        mark2 = mark2.resize((90, 40), Image.ANTIALIAS)
        img.paste(mark2, (522, 1014))
        # 处理
        mark3 = Image.open('.\\resources\\签名\\签名-' + self.comboBox_2.currentText() + '.jpg')
        mark3 = mark3.resize((90, 40), Image.ANTIALIAS)
        img.paste(mark3, (190, 1014))
        img.save(self.lineEdit_56.text())
        # 截图头
        img2 = img.crop((0, 0, 952, 1080))
        temp = self.lineEdit_56.text().split('.')[-1]
        temp = ''.join(['.', temp])
        temp2 = self.lineEdit_56.text().replace(temp, '')
        img2.save(temp2 + 'head.jpg')
        # 转pdf
        img.save(temp2 + '.pdf', "PDF", resolution=300.0, save_all=True)

    def addImg_150(self, img):
        # 审核
        mark1 = Image.open('.\\resources\\签名\\签名-' + self.comboBox_4.currentText() + '.jpg')
        mark1 = mark1.resize((130, 60), Image.ANTIALIAS)
        img.paste(mark1, (1205, 1520))
        # 校对
        mark2 = Image.open('.\\resources\\签名\\签名-' + self.comboBox.currentText() + '.jpg')
        mark2 = mark2.resize((130, 60), Image.ANTIALIAS)
        img.paste(mark2, (790, 1520))
        # 处理
        mark3 = Image.open('.\\resources\\签名\\签名-' + self.comboBox_2.currentText() + '.jpg')
        mark3 = mark3.resize((130, 60), Image.ANTIALIAS)
        img.paste(mark3, (290, 1520))
        img.save(self.lineEdit_56.text())
        # 截图头
        img2 = img.crop((0, 0, 1427, 1620))
        temp = self.lineEdit_56.text().split('.')[-1]
        temp = ''.join(['.', temp])
        temp2 = self.lineEdit_56.text().replace(temp, '')
        img2.save(temp2 + 'head.jpg')
        # 转pdf
        img.save(temp2 + '.pdf', "PDF", resolution=300.0, save_all=True)

    def add_signature_on_pic_100(self):
        if self.run_on_net == True:
            Supervisor.generate_signature_usage_supervisor_100()
        else:
            pass
        fileDir = self.lineEdit_56.text()
        try:
            oriImg = Image.open(fileDir)
            self.addImg_100(oriImg)
            QMessageBox.information(self, "提示", "签名成功")
            # oriImg.show()
        except IOError:
            QMessageBox.information(self, "提示", "不能打开，请确认路径是否正确")

    def add_signature_on_pic_150(self):
        if self.run_on_net == True:
            Supervisor.generate_signature_usage_supervisor_150()
        else:
            pass
        fileDir = self.lineEdit_56.text()
        try:
            oriImg = Image.open(fileDir)
            self.addImg_150(oriImg)
            QMessageBox.information(self, "提示", "签名成功")
            # oriImg.show()
        except IOError:
            QMessageBox.information(self, "提示", "不能打开，请确认路径是否正确")

    def showtime(self):
        # datetime = QDateTime.currentDateTime()
        date = QDate.currentDate()
        time = QTime.currentTime()
        text1 = date.toString(Qt.DefaultLocaleLongDate)
        text2 = time.toString(Qt.DefaultLocaleLongDate)
        self.label_110.setText(text1 + text2)

    def read_raw_info_docx(self):
        fnames = QFileDialog.getOpenFileNames(self, '打开文件', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.lineEdit_5.setText(fname)
        fileDir1 = self.lineEdit_5.text()

        ######################################################################## 初始化
        self.lineEdit_29.setText('西南分公司')
        self.lineEdit_49.setText('套管井测井')
        self.lineEdit_15.setText('中油测井西南分公司')
        self.lineEdit_51.setText('固井质量检测')
        self.lineEdit_42.setText('CBL/VDL')
        self.lineEdit_103.setText('')
        self.lineEdit_105.setText('')
        self.lineEdit_106.setText('')

        document = Document(fileDir1)
        ######################################################################## 井名
        well_Name_Raw = document.tables[0].cell(1, 2).text
        well_Name = well_Name_Raw.replace(' ', '').replace('井', '')
        self.lineEdit.setText(well_Name)
        if self.run_on_net == True:
            Supervisor.load_raw_table_usage_supervisor(well_Name)  # 监控加载
        else:
            pass

        ######################################################################## 井型
        well_Type = document.tables[0].cell(2, 2).text
        well_Type = well_Type.replace(' ', '')
        if '直' in well_Type:
            well_Type = '直井（单井）'
        elif '斜' in well_Type:
            well_Type = '斜直井'
        elif '水平' in well_Type:
            well_Type = '水平井（单井）'
        if well_Type in self.all_Types_Simple:
            self.comboBox_6.setCurrentText(well_Type)

        ######################################################################## 井别
        well_Category = document.tables[0].cell(3, 2).text
        well_Category = well_Category.replace(' ', '')
        if well_Category in self.all_Categories_Simple:
            self.comboBox_7.setCurrentText(well_Category)

        ######################################################################## 井深
        try:
            well_Depth = document.tables[0].cell(4, 2).text.split('m')[0]
        except:
            pass
        well_Depth = well_Depth.replace(' ', '').replace('m', '')
        if well_Depth == '':
            well_Depth = '-99999'
        self.lineEdit_4.setText(well_Depth)

        ######################################################################## X坐标
        x_Coordinate = document.tables[0].cell(5, 2).text.replace(' ', '').replace('（X）', '').replace('横', '').replace(
            '纵', '').replace('m', '').replace(':', '').replace('：', '').replace('(X)', '').replace('（Y）', '').replace(
            '(Y)', '')
        if x_Coordinate == '':
            x_Coordinate = '-99999'
        self.lineEdit_9.setText(x_Coordinate)
        self.lineEdit_8.setText('-99999')

        ######################################################################## Y坐标
        y_Coordinate = document.tables[0].cell(5, 4).text.replace(' ', '').replace('（Y）', '').replace('横', '').replace(
            '纵', '').replace('m', '').replace(':', '').replace('：', '').replace('(Y)', '').replace('（X）', '').replace(
            '(X)', '')
        if y_Coordinate == '':
            y_Coordinate = '-99999'
        self.lineEdit_10.setText(y_Coordinate)
        self.lineEdit_7.setText('-99999')

        ######################################################################## 补心高度
        bushing_Height = document.tables[0].cell(9, 2).text.replace(' ', '').replace('m', '')
        if bushing_Height == '':
            bushing_Height = '-99999'
        self.lineEdit_13.setText(bushing_Height)

        ######################################################################## 地面海拔
        ground_Elevation = document.tables[0].cell(10, 2).text.replace(' ', '').replace('m', '')
        if ground_Elevation == '':
            ground_Elevation = '-99999'
        self.lineEdit_12.setText(ground_Elevation)
        ######################################################################## 补心海拔
        kelly_Bushing = document.tables[0].cell(11, 2).text.replace(' ', '').replace('m', '')
        if kelly_Bushing == '':
            kelly_Bushing = '-99999'
        self.lineEdit_11.setText(kelly_Bushing)
        ######################################################################## 磁偏角
        magnetic_Declination = document.tables[0].cell(12, 2).text.replace(' ', '').replace('°', '')
        if magnetic_Declination == '':
            magnetic_Declination = '-99999'
        self.lineEdit_14.setText(magnetic_Declination)
        ######################################################################## 甲方单位
        client_Name = document.tables[0].cell(14, 2).text.replace(' ', '')
        self.lineEdit_50.setText(client_Name)
        ######################################################################## 推测油田
        if '大庆' in client_Name:
            oil_Field = '大庆油田'
        elif '西南' in client_Name or '四川' in client_Name or '勘探' in client_Name or '蜀南' in client_Name or '重庆' in client_Name or '开发' or '致密' in client_Name:
            oil_Field = '西南油气田'
        else:
            oil_Field = ''
        self.lineEdit_17.setText(oil_Field)
        ######################################################################## 钻井单位
        drilling_Unit = document.tables[0].cell(15, 2).text.replace(' ', '')
        # if drilling_Unit == '':
        #     drilling_Unit = '-99999'
        self.lineEdit_74.setText(drilling_Unit)
        ######################################################################## 通过钻井单位推断固井单位
        cement_Unit = ''
        if '川' in drilling_Unit:
            cement_Unit = '川庆钻探'
        elif '中原' in drilling_Unit:
            cement_Unit = '中原钻井'
        elif '长城' in drilling_Unit:
            cement_Unit = '长城钻探'
        elif '大庆' in drilling_Unit:
            cement_Unit = '大庆钻探'
        elif '渤' in drilling_Unit:
            cement_Unit = '渤海钻探'
        else:
            cement_Unit == ''
        # if cement_Unit == '':
        #     cement_Unit = '-99999'
        self.lineEdit_25.setText(cement_Unit)
        ######################################################################## 开钻日期
        spud_Date = document.tables[0].cell(16, 2).text
        if spud_Date != '':
            spud_Date = spud_Date.replace(' ', '')
        if '年' in spud_Date and '月' in spud_Date and '日' in spud_Date:
            spud_Date1 = spud_Date.split('年')[0]
            spud_Date2 = spud_Date.split('年')[1].split('月')[0]
            spud_Date3 = spud_Date.split('月')[1].split('日')[0]
            spud_Date = '-'.join([spud_Date1, spud_Date2, spud_Date3])
            if spud_Date1 == '' and spud_Date2 == '' and spud_Date3 == '':
                spud_Date = ''  # '1900-01-01'
            try:
                spud_Date = self.parse_mixed_datetime(spud_Date).strftime('%Y-%m-%d %H:%M:%S')
            except:
                # QMessageBox.information(self, "提示", "请检查开钻日期是否为空")
                pass
        elif '-' in spud_Date:
            spud_Date1 = spud_Date.split('-')[0]
            spud_Date2 = spud_Date.split('-')[1]
            spud_Date3 = spud_Date.split('-')[2]
            spud_Date = '-'.join([spud_Date1, spud_Date2, spud_Date3])
            if spud_Date1 == '' and spud_Date2 == '' and spud_Date3 == '':
                spud_Date = ''  # '1900-01-01'
            try:
                spud_Date = self.parse_mixed_datetime(spud_Date).strftime('%Y-%m-%d %H:%M:%S')
            except:
                # QMessageBox.information(self, "提示", "请检查开钻日期是否为空")
                pass
        elif '.' in spud_Date:
            spud_Date1 = spud_Date.split('.')[0]
            spud_Date2 = spud_Date.split('.')[1]
            spud_Date3 = spud_Date.split('.')[2]
            spud_Date = '-'.join([spud_Date1, spud_Date2, spud_Date3])
            if spud_Date1 == '' and spud_Date2 == '' and spud_Date3 == '':
                spud_Date = ''  # '1900-01-01'
            try:
                spud_Date = self.parse_mixed_datetime(spud_Date).strftime('%Y-%m-%d %H:%M:%S')
            except:
                # QMessageBox.information(self, "提示", "请检查开钻日期是否为空")
                pass
        if spud_Date == '':
            spud_Date = ''  # '1900-01-01'
        spud_Date = spud_Date.split(' ')[0]
        self.lineEdit_20.setText(spud_Date)
        ######################################################################## 完钻日期
        end_Drilling_Date = document.tables[0].cell(17, 2).text
        if end_Drilling_Date != '':
            end_Drilling_Date = end_Drilling_Date.replace(' ', '')
        if '年' in end_Drilling_Date and '月' in end_Drilling_Date and '日' in end_Drilling_Date:
            end_Drilling_Date1 = end_Drilling_Date.split('年')[0]
            end_Drilling_Date2 = end_Drilling_Date.split('年')[1].split('月')[0]
            end_Drilling_Date3 = end_Drilling_Date.split('月')[1].split('日')[0]
            end_Drilling_Date = '-'.join([end_Drilling_Date1, end_Drilling_Date2, end_Drilling_Date3])
            if end_Drilling_Date1 == '' and end_Drilling_Date2 == '' and end_Drilling_Date3 == '':
                end_Drilling_Date = ''  # '1900-01-01'
            try:
                end_Drilling_Date = self.parse_mixed_datetime(end_Drilling_Date).strftime('%Y-%m-%d %H:%M:%S')
            except:
                # QMessageBox.information(self, "提示", "请检查完钻日期是否为空")
                pass
        elif '-' in end_Drilling_Date:
            end_Drilling_Date1 = end_Drilling_Date.split('-')[0]
            end_Drilling_Date2 = end_Drilling_Date.split('-')[1]
            end_Drilling_Date3 = end_Drilling_Date.split('-')[2]
            end_Drilling_Date = '-'.join([end_Drilling_Date1, end_Drilling_Date2, end_Drilling_Date3])
            if end_Drilling_Date1 == '' and end_Drilling_Date2 == '' and end_Drilling_Date3 == '':
                end_Drilling_Date = ''  # '1900-01-01'
            try:
                end_Drilling_Date = self.parse_mixed_datetime(end_Drilling_Date).strftime('%Y-%m-%d %H:%M:%S')
            except:
                # QMessageBox.information(self, "提示", "请检查完钻日期是否为空")
                pass
        elif '.' in end_Drilling_Date:
            end_Drilling_Date1 = end_Drilling_Date.split('.')[0]
            end_Drilling_Date2 = end_Drilling_Date.split('.')[1]
            end_Drilling_Date3 = end_Drilling_Date.split('.')[2]
            end_Drilling_Date = '-'.join([end_Drilling_Date1, end_Drilling_Date2, end_Drilling_Date3])
            if end_Drilling_Date1 == '' and end_Drilling_Date2 == '' and end_Drilling_Date3 == '':
                end_Drilling_Date = ''  # '1900-01-01'
            try:
                end_Drilling_Date = self.parse_mixed_datetime(end_Drilling_Date).strftime('%Y-%m-%d %H:%M:%S')
            except:
                # QMessageBox.information(self, "提示", "请检查完钻日期是否为空")
                pass
        if end_Drilling_Date == '':
            end_Drilling_Date = ''  # '1900-01-01'
        end_Drilling_Date = end_Drilling_Date.split(' ')[0]
        self.lineEdit_21.setText(end_Drilling_Date)
        ######################################################################## 完井日期
        completion_Date = document.tables[0].cell(18, 2).text
        if completion_Date != '':
            completion_Date = completion_Date.replace(' ', '')
        if '年' in completion_Date and '月' in completion_Date and '日' in completion_Date:
            completion_Date1 = completion_Date.split('年')[0].replace(' ', '')
            completion_Date2 = completion_Date.split('年')[1].split('月')[0].replace(' ', '')
            completion_Date3 = completion_Date.split('月')[1].split('日')[0].replace(' ', '')
            completion_Date = '-'.join([completion_Date1, completion_Date2, completion_Date3])
            if completion_Date1 == '' and completion_Date2 == '' and completion_Date3 == '':
                completion_Date = ''  # '1900-01-01'
            try:
                completion_Date = self.parse_mixed_datetime(completion_Date).strftime('%Y-%m-%d %H:%M:%S')
            except:
                # QMessageBox.information(self, "提示", "请检查完井日期是否为空")
                pass
        elif '-' in completion_Date:
            completion_Date1 = completion_Date.split('-')[0]
            completion_Date2 = completion_Date.split('-')[1]
            completion_Date3 = completion_Date.split('-')[2]
            completion_Date = '-'.join([completion_Date1, completion_Date2, completion_Date3])
            if completion_Date1 == '' and completion_Date2 == '' and completion_Date3 == '':
                completion_Date = ''  # '1900-01-01'
            try:
                completion_Date = self.parse_mixed_datetime(completion_Date).strftime('%Y-%m-%d %H:%M:%S')
            except:
                # QMessageBox.information(self, "提示", "请检查完井日期是否为空")
                pass
        elif '.' in completion_Date:
            completion_Date1 = completion_Date.split('.')[0]
            completion_Date2 = completion_Date.split('.')[1]
            completion_Date3 = completion_Date.split('.')[2]
            completion_Date = '-'.join([completion_Date1, completion_Date2, completion_Date3])
            if completion_Date1 == '' and completion_Date2 == '' and completion_Date3 == '':
                completion_Date = ''  # '1900-01-01'
            try:
                completion_Date = self.parse_mixed_datetime(completion_Date).strftime('%Y-%m-%d %H:%M:%S')
            except:
                # QMessageBox.information(self, "提示", "请检查完井日期是否为空")
                pass
        if completion_Date == '':
            completion_Date = ''  # '1900-01-01'
        completion_Date = completion_Date.split(' ')[0]
        self.lineEdit_39.setText(completion_Date)
        ######################################################################## 钻头数据
        bit1_Diameter = document.tables[0].cell(20, 2).text.strip()
        bit1_Diameter = bit1_Diameter.replace(' ', '')
        bit1_Diameter = bit1_Diameter.split('mm')
        bit1_Diameter = bit1_Diameter[0]
        bit1_Depth = document.tables[0].cell(20, 5).text.strip()
        bit1_Depth = bit1_Depth.replace(' ', '').replace('M', '')
        bit1_Depth = bit1_Depth.split('m')
        bit1_Depth = bit1_Depth[0]
        # 默认值
        if bit1_Diameter == '':
            bit1_Diameter = '-99999'
        if bit1_Depth == '':
            bit1_Depth = '-99999'
        self.tableWidget_6.setItem(0, 0, QTableWidgetItem(str(bit1_Diameter)))
        self.tableWidget_6.setItem(0, 1, QTableWidgetItem(str(bit1_Depth)))
        bit2_Diameter = document.tables[0].cell(21, 2).text.strip()
        bit2_Diameter = bit2_Diameter.replace(' ', '')
        bit2_Diameter = bit2_Diameter.split('mm')
        bit2_Diameter = bit2_Diameter[0]
        bit2_Depth = document.tables[0].cell(21, 5).text.strip()
        bit2_Depth = bit2_Depth.replace(' ', '').replace('M', '')
        bit2_Depth = bit2_Depth.split('m')
        bit2_Depth = bit2_Depth[0]
        self.tableWidget_6.setItem(1, 0, QTableWidgetItem(str(bit2_Diameter)))
        self.tableWidget_6.setItem(1, 1, QTableWidgetItem(str(bit2_Depth)))

        bit3_Diameter = document.tables[0].cell(22, 2).text.strip()
        bit3_Diameter = bit3_Diameter.replace(' ', '')
        bit3_Diameter = bit3_Diameter.split('mm')
        bit3_Diameter = bit3_Diameter[0]
        bit3_Depth = document.tables[0].cell(22, 5).text.strip()
        bit3_Depth = bit3_Depth.replace(' ', '').replace('M', '')
        bit3_Depth = bit3_Depth.split('m')
        bit3_Depth = bit3_Depth[0]
        self.tableWidget_6.setItem(2, 0, QTableWidgetItem(str(bit3_Diameter)))
        self.tableWidget_6.setItem(2, 1, QTableWidgetItem(str(bit3_Depth)))

        bit4_Diameter = document.tables[0].cell(23, 2).text.strip()
        bit4_Diameter = bit4_Diameter.replace(' ', '')
        bit4_Diameter = bit4_Diameter.split('mm')
        bit4_Diameter = bit4_Diameter[0]
        bit4_Depth = document.tables[0].cell(23, 5).text.strip()
        bit4_Depth = bit4_Depth.replace(' ', '').replace('M', '')
        bit4_Depth = bit4_Depth.split('m')
        bit4_Depth = bit4_Depth[0]
        self.tableWidget_6.setItem(3, 0, QTableWidgetItem(str(bit4_Diameter)))
        self.tableWidget_6.setItem(3, 1, QTableWidgetItem(str(bit4_Depth)))

        bit5_Diameter = document.tables[0].cell(24, 2).text.strip()
        bit5_Diameter = bit5_Diameter.replace(' ', '')
        bit5_Diameter = bit5_Diameter.split('mm')
        bit5_Diameter = bit5_Diameter[0]
        bit5_Depth = document.tables[0].cell(24, 5).text.strip()
        bit5_Depth = bit5_Depth.replace(' ', '').replace('M', '')
        bit5_Depth = bit5_Depth.split('m')
        bit5_Depth = bit5_Depth[0]
        self.tableWidget_6.setItem(4, 0, QTableWidgetItem(str(bit5_Diameter)))
        self.tableWidget_6.setItem(4, 1, QTableWidgetItem(str(bit5_Depth)))

        bit6_Diameter = document.tables[0].cell(25, 2).text.strip()
        bit6_Diameter = bit6_Diameter.replace(' ', '')
        bit6_Diameter = bit6_Diameter.split('mm')
        bit6_Diameter = bit6_Diameter[0]
        bit6_Depth = document.tables[0].cell(25, 5).text.strip()
        bit6_Depth = bit6_Depth.replace(' ', '').replace('M', '')
        bit6_Depth = bit6_Depth.split('m')
        bit6_Depth = bit6_Depth[0]

        bit7_Diameter = document.tables[0].cell(26, 2).text.strip()
        bit7_Diameter = bit7_Diameter.replace(' ', '')
        bit7_Diameter = bit7_Diameter.split('mm')
        bit7_Diameter = bit7_Diameter[0]
        bit7_Depth = document.tables[0].cell(26, 5).text.strip()
        bit7_Depth = bit7_Depth.replace(' ', '').replace('M', '')
        bit7_Depth = bit7_Depth.split('m')
        bit7_Depth = bit7_Depth[0]

        bit8_Diameter = document.tables[0].cell(27, 2).text.strip()
        bit8_Diameter = bit8_Diameter.replace(' ', '')
        bit8_Diameter = bit8_Diameter.split('mm')
        bit8_Diameter = bit8_Diameter[0]
        bit8_Depth = document.tables[0].cell(27, 5).text.strip()
        bit8_Depth = bit8_Depth.replace(' ', '').replace('M', '')
        bit8_Depth = bit8_Depth.split('m')
        bit8_Depth = bit8_Depth[0]

        bit9_Diameter = document.tables[0].cell(28, 2).text.strip()
        bit9_Diameter = bit9_Diameter.replace(' ', '')
        bit9_Diameter = bit9_Diameter.split('mm')
        bit9_Diameter = bit9_Diameter[0]
        bit9_Depth = document.tables[0].cell(28, 5).text.strip()
        bit9_Depth = bit9_Depth.replace(' ', '').replace('M', '')
        bit9_Depth = bit9_Depth.split('m')
        bit9_Depth = bit9_Depth[0]

        bit10_Diameter = document.tables[0].cell(29, 2).text.strip()
        bit10_Diameter = bit10_Diameter.replace(' ', '')
        bit10_Diameter = bit10_Diameter.split('mm')
        bit10_Diameter = bit10_Diameter[0]
        bit10_Depth = document.tables[0].cell(29, 5).text.strip()
        bit10_Depth = bit10_Depth.replace(' ', '').replace('M', '')
        bit10_Depth = bit10_Depth.split('m')
        bit10_Depth = bit10_Depth[0]

        # 找出最深的钻头深度deepest_bit
        if bit10_Depth != '':
            deepest_bit = bit10_Depth
        elif bit9_Depth != '':
            deepest_bit = bit9_Depth
        elif bit8_Depth != '':
            deepest_bit = bit8_Depth
        elif bit7_Depth != '':
            deepest_bit = bit7_Depth
        elif bit6_Depth != '':
            deepest_bit = bit6_Depth
        elif bit5_Depth != '':
            deepest_bit = bit5_Depth
        elif bit4_Depth != '':
            deepest_bit = bit4_Depth
        elif bit3_Depth != '':
            deepest_bit = bit3_Depth
        elif bit2_Depth != '':
            deepest_bit = bit2_Depth
        elif bit1_Depth != '':
            deepest_bit = bit1_Depth
        self.lineEdit_102.setText(deepest_bit)
        self.lineEdit_27.setText(deepest_bit)
        self.lineEdit_99.setText(deepest_bit)
        self.lineEdit_53.setText(deepest_bit)
        self.lineEdit_23.setText('-99999')
        # 地理位置geo_Position
        geographic_Position = document.tables[0].cell(30, 2).text.strip()
        if geographic_Position != '':
            if '省' in geographic_Position:
                geographic_Position = geographic_Position.split('省')
                geographic_Position1 = ''.join([geographic_Position[0], '省'])
                if '县' not in geographic_Position[1]:
                    geographic_Position2 = geographic_Position[1].split('市')
                    geographic_Position2 = ''.join([geographic_Position2[0], '市'])
                else:
                    geographic_Position2 = geographic_Position[1].split('县')
                    geographic_Position2 = geographic_Position2[0]
                    if '市' in geographic_Position2:
                        geographic_Position2 = geographic_Position2.split('市')[1]
                    geographic_Position2 = ''.join([geographic_Position2, '县'])
            elif '省' not in geographic_Position:
                geographic_Position = geographic_Position.split('市')
                geographic_Position1 = ''.join([geographic_Position[0], '市'])
                if '区' in geographic_Position[1]:
                    geographic_Position2 = geographic_Position[1].split('区')
                    geographic_Position2 = ''.join([geographic_Position2[0], '区'])
                elif '县' in geographic_Position[1]:
                    geographic_Position2 = geographic_Position[1].split('县')
                    geographic_Position2 = ''.join([geographic_Position2[0], '县'])
            geo_Position = ''.join([geographic_Position1, geographic_Position2])
        self.lineEdit_16.setText(geo_Position)
        # 构造位置stru_Position
        structure_Position = document.tables[0].cell(31, 2).text
        if structure_Position != '':
            structure_Position = structure_Position.replace(' ', '')
            structure_Position = structure_Position.replace('四川盆地', '')
            structure_Position = structure_Position.split('构造')
            stru_Position = structure_Position[0]
            if '高石' in stru_Position:
                stru_Position = '高石梯'
            elif '磨溪' in stru_Position:
                stru_Position = '磨溪'
            elif '威远' in stru_Position:
                stru_Position = '威远'
            elif '龙岗' in stru_Position:
                stru_Position = '龙岗'
            elif '龙会' in stru_Position:
                stru_Position = '龙会场'
            elif '中坝' in stru_Position:
                stru_Position = '中坝'
            elif '双鱼' in stru_Position:
                stru_Position = '双鱼石'
            elif '龙女寺' in stru_Position:
                stru_Position = '龙女寺'
            elif '安岳' in stru_Position:
                stru_Position = '安岳'
            elif '长宁' in stru_Position:
                stru_Position = '长宁'
            elif '黄草峡' in stru_Position:
                stru_Position = '黄草峡'

            if '蓬探' in well_Name:
                stru_Position = '太和含气区'
            elif '蓬深' in well_Name:
                stru_Position = '太和含气区'
            elif '金浅' in well_Name:
                stru_Position = '金华'
            elif '磨溪' in well_Name:
                stru_Position = '磨溪'
            elif '高石' in well_Name:
                stru_Position = '高石梯'
            elif '双鱼' in well_Name:
                stru_Position = '双鱼石'
            elif '双探' in well_Name:
                stru_Position = '双鱼石'
            elif '秋林' in well_Name:
                stru_Position = '秋林'
            elif '永浅' in well_Name:
                stru_Position = '成都-简阳宽缓单斜'
            elif '中浅' in well_Name:
                stru_Position = '川中低缓构造带'

        self.lineEdit_18.setText(stru_Position)
        # 目的层层位goal_layer_name
        goal_layer_name = document.tables[0].cell(33, 2).text.replace(' ', '')
        self.lineEdit_118.setText(goal_layer_name)

        ######################################################################## 任务单号
        task_Number = document.tables[1].cell(1, 3).text.strip()
        self.lineEdit_19.setText(task_Number)

        ######################################################################## 井次类型
        well_Times_Type = document.tables[1].cell(2, 2).text.strip()
        # 已默认设置为套管井测井

        ######################################################################## 钻井液flu_Property, flu_Density, flu_Viscosity
        flu_Property = document.tables[1].cell(10, 2).text.strip()
        flu_Property = flu_Property.replace(' ', '')
        self.lineEdit_60.setText(flu_Property)
        self.lineEdit_64.setText(flu_Property)

        drilling_Fluid_Density = document.tables[1].cell(11, 2).text.strip()
        drilling_Fluid_Density = drilling_Fluid_Density.replace(' ', '')
        drilling_Fluid_Density = drilling_Fluid_Density.split('g')
        flu_Density = drilling_Fluid_Density[0]
        self.lineEdit_61.setText(flu_Density)

        drilling_Fluid_Viscosity = document.tables[1].cell(12, 2).text
        flu_Viscosity = drilling_Fluid_Viscosity.replace(' ', '').replace('s', '').replace('S', '').replace('秒', '')
        self.lineEdit_62.setText(flu_Viscosity)

        # 泥浆电阻率
        drilling_Fluid_RT= document.tables[1].cell(13, 2).text
        fluid_RT = drilling_Fluid_RT.replace(' ', '')
        if fluid_RT == '':
            self.lineEdit_66.setText('-99999')
        else:
            self.lineEdit_66.setText(fluid_RT)

        # 泥浆温度 (原始资料登记表上无，但LEAD要求填写)
        self.lineEdit_67.setText('-99999')

        # 泥浆PH值
        drilling_Fluid_PH = document.tables[1].cell(14, 2).text
        fluid_PH = drilling_Fluid_PH.replace(' ', '')
        self.lineEdit_68.setText(fluid_PH)

        # 泥浆失水
        drilling_Fluid_Lost_Water = document.tables[1].cell(15, 2).text
        fluid_Lost_Water = drilling_Fluid_Lost_Water.replace(' ', '').replace('m', '').replace('M', '').replace('l',
                                                                                                                '').replace(
            'L', '')
        self.lineEdit_63.setText(fluid_Lost_Water)

        ######################################################################## 测井装备
        logging_Equipment = document.tables[1].cell(16, 1).text.strip()
        logging_Equipment = logging_Equipment.replace(' ', '')
        if '5700' in logging_Equipment:
            logging_Equipment = 'ECLIPS-5700'
        elif '2530' in logging_Equipment:
            logging_Equipment = 'HH2530'
        elif '3000' in logging_Equipment:
            logging_Equipment = 'SKD-3000'
        elif 'MCET' in logging_Equipment:
            logging_Equipment = 'MCET-1000'
        elif 'mcet' in logging_Equipment:
            logging_Equipment = 'MCET-1000'
        elif 'CP' in logging_Equipment:
            logging_Equipment = 'CPLog'
        elif 'cp' in logging_Equipment:
            logging_Equipment = 'CPLog'
        if logging_Equipment in self.logging_Equipments:
            self.comboBox_5.setCurrentText(logging_Equipment)
        ######################################################################## 测井方式
        logging_Method = document.tables[1].cell(16, 3).text.strip()
        self.lineEdit_100.setText(logging_Method)
        ######################################################################## 测井小队
        logging_Group = document.tables[1].cell(17, 1).text.strip().replace('队', '')
        self.lineEdit_43.setText(logging_Group)
        ######################################################################## 小队长
        logging_Leader = document.tables[1].cell(18, 1).text.strip()
        self.lineEdit_54.setText(logging_Leader)
        ######################################################################## 操作员
        logging_Operator = document.tables[1].cell(19, 1).text.strip()
        self.lineEdit_44.setText(logging_Operator)
        ######################################################################## 时间cement_End_Time, logging_Start_Time, logging_End_Time
        cement_End_Date = ''
        cement_End_Time = document.tables[1].cell(17, 3).text.strip().replace('.', '-').replace('  ', ' ').replace(': ',
                                                                                                                   ':').replace(
            ' :', ':').replace(' : ', ':')
        if cement_End_Time != '':
            try:
                cement_End_Time = self.parse_mixed_datetime(cement_End_Time).strftime('%Y-%m-%d %H:%M:%S')
            except:
                QMessageBox.information(self, "提示", "格式化cement_End_Time时报错")
        else:
            cement_End_Time = ''  # '1900-01-01'
        self.lineEdit_101.setText(cement_End_Time)

        cement_End_Date = cement_End_Time.split(' ')[0]
        self.lineEdit_22.setText(cement_End_Date)

        logging_Start_Time = document.tables[1].cell(18, 3).text.strip().replace('.', '-').replace('  ', ' ').replace(
            ': ', ':').replace(' :', ':').replace(' : ', ':')
        if logging_Start_Time != '':
            try:
                logging_Start_Time = self.parse_mixed_datetime(logging_Start_Time).strftime('%Y-%m-%d %H:%M:%S')
            except:
                QMessageBox.information(self, "提示", "格式化logging_Start_Time时报错")
        else:
            logging_Start_Time = ''  # '1900-01-01 00:00'
        self.lineEdit_104.setText(logging_Start_Time)

        logging_End_Time = document.tables[1].cell(19, 3).text.strip().replace('.', '-').replace('  ', ' ').replace(
            ': ', ':').replace(' :', ':').replace(' : ', ':')
        if logging_End_Time != '':
            try:
                logging_End_Time = self.parse_mixed_datetime(logging_End_Time).strftime('%Y-%m-%d %H:%M:%S')
            except:
                QMessageBox.information(self, "提示", "格式化logging_End_Time时报错")
        else:
            logging_End_Time = ''  # '1900-01-01 00:00'
        self.lineEdit_121.setText(logging_End_Time)
        # 测井日期
        logging_Date = str(logging_End_Time).split(' ')[0]
        self.lineEdit_40.setText(logging_Date)
        ######################################################################## 推测日期
        # 将日期字符串解析为日期对象
        parsed_date = datetime.strptime(logging_End_Time, '%Y-%m-%d %H:%M:%S')

        # 添加几天
        raw_Data_Recieve_Date = parsed_date + timedelta(days=1)
        raw_Data_Evaluate_Date = parsed_date + timedelta(days=1.5)
        interpretation_Start_Date = parsed_date + timedelta(days=2)
        interpretation_Complete_Date = parsed_date + timedelta(days=3)
        data_Archive_Date = parsed_date + timedelta(days=4)

        # 格式化新日期为字符串
        raw_Data_Recieve_Date = raw_Data_Recieve_Date.strftime('%Y-%m-%d %H:%M:%S')
        raw_Data_Evaluate_Date = raw_Data_Evaluate_Date.strftime('%Y-%m-%d %H:%M:%S')
        interpretation_Start_Date = interpretation_Start_Date.strftime('%Y-%m-%d %H:%M:%S')
        interpretation_Complete_Date = interpretation_Complete_Date.strftime('%Y-%m-%d %H:%M:%S')
        data_Archive_Date = data_Archive_Date.strftime('%Y-%m-%d %H:%M:%S')

        self.lineEdit_33.setText(raw_Data_Recieve_Date)
        self.lineEdit_34.setText(raw_Data_Evaluate_Date)
        self.lineEdit_30.setText(interpretation_Start_Date)
        self.lineEdit_55.setText(interpretation_Complete_Date)
        self.lineEdit_32.setText(data_Archive_Date)

        ######################################################################## 最大井斜斜度
        try:
            max_Well_Deviation = document.tables[2].cell(1, 2).text
            max_Well_Deviation = max_Well_Deviation.replace(' ', '').replace('°', '')
            max_Well_Deviation = round(float(max_Well_Deviation), 2)
            max_Well_Deviation = str(max_Well_Deviation)
        except:
            # QMessageBox.information(self, "提示", "请检查max_Well_Deviation（最大井斜）是否为空")
            max_Well_Deviation = '-99999'
        self.lineEdit_28.setText(max_Well_Deviation)

        ######################################################################## 最大井斜深度
        try:
            max_Well_Deviation_Depth = document.tables[2].cell(1, 11).text
            max_Well_Deviation_Depth = max_Well_Deviation_Depth.replace(' ', '').replace('m', '')
            max_Well_Deviation_Depth = round(float(max_Well_Deviation_Depth), 2)
            max_Well_Deviation_Depth = str(max_Well_Deviation_Depth)
        except:
            # QMessageBox.information(self, "提示", "请检查max_Well_Deviation_Depth（最大井斜深度）是否为空")
            max_Well_Deviation_Depth = '-99999'
        self.lineEdit_24.setText(max_Well_Deviation_Depth)

        if max_Well_Deviation != '' and max_Well_Deviation_Depth != '' and max_Well_Deviation != '-99999' and max_Well_Deviation_Depth != '-99999':
            dev_Depth_Ratio = ''.join([max_Well_Deviation, '/', max_Well_Deviation_Depth])
        else:
            dev_Depth_Ratio = ''
        self.lineEdit_26.setText(dev_Depth_Ratio)

        ######################################################################## 人工井底arti_Bottom
        try:
            artificial_Bottom_of_Well = document.tables[2].cell(2, 2).text.strip()
            artificial_Bottom_of_Well = artificial_Bottom_of_Well.replace(' ', '')
            artificial_Bottom_of_Well = artificial_Bottom_of_Well.replace('m', '')
            artificial_Bottom_of_Well = round(float(artificial_Bottom_of_Well), 2)
            artificial_Bottom_of_Well = str(artificial_Bottom_of_Well)
            # 确保整数后面也有小数，为了好看
            if '.' in artificial_Bottom_of_Well:
                arti_Bottom = artificial_Bottom_of_Well
            else:
                if artificial_Bottom_of_Well != '':
                    arti_Bottom = ''.join([artificial_Bottom_of_Well, '.00'])
                else:
                    arti_Bottom = ''
        except:
            # QMessageBox.information(self, "提示", "请检查arti_Bottom（人工井底）是否为空")
            arti_Bottom = '-99999'
        self.lineEdit_65.setText(arti_Bottom)
        ######################################################################## 已注入水泥量cement_Quantity
        try:
            cement_Quantity = document.tables[2].cell(3, 11).text
            cement_Quantity = cement_Quantity.replace(' ', '')
            cement_Quantity = cement_Quantity.replace('T', '')
            cement_Quantity = cement_Quantity.replace('t', '')
            cement_Quantity = cement_Quantity.replace('m3', '')
            cement_Quantity = round(float(cement_Quantity), 2)
            cement_Quantity = str(cement_Quantity)
        except:
            # QMessageBox.information(self, "提示", "请检查cement_Quantity（水泥量）是否为空")
            cement_Quantity = '-99999'
        self.lineEdit_71.setText(cement_Quantity)

        ######################################################################## 水泥密度cement_Density
        cement_Density = ''
        slow_Cement_Density = document.tables[2].cell(7, 11).text.strip()
        fast_Cement_Density = document.tables[2].cell(8, 11).text.strip()
        if slow_Cement_Density == '':
            cement_Density = fast_Cement_Density
        elif fast_Cement_Density == '':
            cement_Density = slow_Cement_Density
        elif eval(str(slow_Cement_Density)) == eval(str(fast_Cement_Density)):
            cement_Density = fast_Cement_Density
        elif eval(str(slow_Cement_Density)) > eval(str(fast_Cement_Density)):
            cement_Density = ''.join([fast_Cement_Density, '~', slow_Cement_Density])
        elif eval(str(slow_Cement_Density)) < eval(str(fast_Cement_Density)):
            cement_Density = ''.join([slow_Cement_Density, '~', fast_Cement_Density])
        other_Cement_Density = document.tables[2].cell(9, 11).text.strip()
        if other_Cement_Density != '' and cement_Density == '':
            cement_Density = other_Cement_Density
        self.lineEdit_72.setText(cement_Density)

        # 密度大于1.75产生警告
        try:
            try:
                slow_Density = float(slow_Cement_Density)
            except:
                pass
            try:
                fast_Density = float(fast_Cement_Density)
            except:
                pass
            if slow_Density >= 1.75 or fast_Density >= 1.75:
                self.label_134.setText('注意按照15/30标准处理')
                self.label_134.setStyleSheet("font: 12pt")
                self.label_134.setStyleSheet("color: rgb(255, 0, 0)")
            else:
                pass
        except:
            pass

        ######################################################################## 水泥设计返高design_Depth
        try:
            design_Depth = document.tables[2].cell(5, 2).text.strip()
            design_Depth = design_Depth.replace(' ', '')
            design_Depth = design_Depth.replace('m', '').replace('.0', '').replace('.00', '')
            if design_Depth in ['井口', '地面']:
                design_Depth = '0'
        except:
            # QMessageBox.information(self, "提示", "请检查design_Depth（水泥设计返高）是否为空")
            design_Depth = '-99999'
        self.lineEdit_69.setText(design_Depth)
        if design_Depth == '':
            self.lineEdit_69.setText('0')

        ######################################################################## 水泥实际返高actual_Depth
        try:
            actual_Depth = document.tables[2].cell(5, 11).text.strip()
            actual_Depth = actual_Depth.replace(' ', '')
            actual_Depth = actual_Depth.replace('m', '')
            actual_Depth = actual_Depth.replace('以上', '')
            actual_Depth = actual_Depth.replace('（', '')
            actual_Depth = actual_Depth.replace('）', '')
            actual_Depth = actual_Depth.replace('(', '')
            actual_Depth = actual_Depth.replace(')', '')
            actual_Depth = actual_Depth.replace('地面', '0')
            actual_Depth = actual_Depth.replace('井口', '0')
            actual_Depth = actual_Depth.replace('.00', '.0')
            if '.' in actual_Depth:
                actual_Depth = actual_Depth
            else:
                if actual_Depth == '0':
                    actual_Depth = actual_Depth
                elif actual_Depth != '':
                    actual_Depth = ''.join([actual_Depth, '.0'])
                else:
                    actual_Depth = '空'
        except:
            # QMessageBox.information(self, "提示", "请检查actual_Depth（水泥实际返高）是否为空")
            actual_Depth = '-99999'
        self.lineEdit_70.setText(actual_Depth)
        try:
            if float(actual_Depth) < 200:
                self.lineEdit_107.setText('0')
            else:
                pass
        except:
            pass

        # 利用相对定位找寻套管数据起始坐标
        table = document.tables[2]
        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                # table.cell(row, col).text += '({0},{1})'.format(row, col)  # 给文本中的单元格添加表格坐标
                if '钢级' in table.cell(row, col).text:
                    # print('(', str(row), ',', str(col), '):', table.cell(row, col).text)
                    row_reference = row
                    ######################################################################## 套管数据
        # 套管外径
        casing1_Dia = document.tables[2].cell(row_reference + 1, 5).text.strip()
        casing2_Dia = document.tables[2].cell(row_reference + 2, 5).text.strip()
        casing3_Dia = document.tables[2].cell(row_reference + 3, 5).text.strip()
        casing4_Dia = document.tables[2].cell(row_reference + 4, 5).text.strip()
        casing5_Dia = document.tables[2].cell(row_reference + 5, 5).text.strip()
        casing6_Dia = document.tables[2].cell(row_reference + 6, 5).text.strip()
        casing7_Dia = document.tables[2].cell(row_reference + 7, 5).text.strip()
        casing8_Dia = document.tables[2].cell(row_reference + 8, 5).text.strip()
        casing9_Dia = document.tables[2].cell(row_reference + 9, 5).text.strip()
        casing10_Dia = document.tables[2].cell(row_reference + 10, 5).text.strip()
        casing11_Dia = document.tables[2].cell(row_reference + 11, 5).text.strip()
        casing12_Dia = document.tables[2].cell(row_reference + 12, 5).text.strip()
        self.tableWidget_7.setItem(0, 1, QTableWidgetItem(str(casing1_Dia)))
        self.tableWidget_7.setItem(1, 1, QTableWidgetItem(str(casing2_Dia)))
        self.tableWidget_7.setItem(2, 1, QTableWidgetItem(str(casing3_Dia)))
        self.tableWidget_7.setItem(3, 1, QTableWidgetItem(str(casing4_Dia)))
        self.tableWidget_7.setItem(4, 1, QTableWidgetItem(str(casing5_Dia)))

        # 套管内径
        casing1_Inner_Dia = document.tables[2].cell(row_reference + 1, 7).text.strip()
        casing2_Inner_Dia = document.tables[2].cell(row_reference + 2, 7).text.strip()
        casing3_Inner_Dia = document.tables[2].cell(row_reference + 3, 7).text.strip()
        casing4_Inner_Dia = document.tables[2].cell(row_reference + 4, 7).text.strip()
        casing5_Inner_Dia = document.tables[2].cell(row_reference + 5, 7).text.strip()
        casing6_Inner_Dia = document.tables[2].cell(row_reference + 6, 7).text.strip()
        casing7_Inner_Dia = document.tables[2].cell(row_reference + 7, 7).text.strip()
        casing8_Inner_Dia = document.tables[2].cell(row_reference + 8, 7).text.strip()
        casing9_Inner_Dia = document.tables[2].cell(row_reference + 9, 7).text.strip()
        casing10_Inner_Dia = document.tables[2].cell(row_reference + 10, 7).text.strip()
        casing11_Inner_Dia = document.tables[2].cell(row_reference + 11, 7).text.strip()
        casing12_Inner_Dia = document.tables[2].cell(row_reference + 12, 7).text.strip()
        self.tableWidget_7.setItem(0, 0, QTableWidgetItem(str(casing1_Inner_Dia)))
        self.tableWidget_7.setItem(1, 0, QTableWidgetItem(str(casing2_Inner_Dia)))
        self.tableWidget_7.setItem(2, 0, QTableWidgetItem(str(casing3_Inner_Dia)))
        self.tableWidget_7.setItem(3, 0, QTableWidgetItem(str(casing4_Inner_Dia)))
        self.tableWidget_7.setItem(4, 0, QTableWidgetItem(str(casing5_Inner_Dia)))

        # 套管壁厚
        casing1_Thickness = document.tables[2].cell(row_reference + 1, 8).text.strip().replace('尺寸（mm）', '')
        casing2_Thickness = document.tables[2].cell(row_reference + 2, 8).text.strip().replace('尺寸（mm）', '')
        casing3_Thickness = document.tables[2].cell(row_reference + 3, 8).text.strip().replace('尺寸（mm）', '')
        casing4_Thickness = document.tables[2].cell(row_reference + 4, 8).text.strip().replace('尺寸（mm）', '')
        casing5_Thickness = document.tables[2].cell(row_reference + 5, 8).text.strip().replace('尺寸（mm）', '')
        casing6_Thickness = document.tables[2].cell(row_reference + 6, 8).text.strip().replace('尺寸（mm）', '')
        casing7_Thickness = document.tables[2].cell(row_reference + 7, 8).text.strip().replace('尺寸（mm）', '')
        casing8_Thickness = document.tables[2].cell(row_reference + 8, 8).text.strip().replace('尺寸（mm）', '')
        casing9_Thickness = document.tables[2].cell(row_reference + 9, 8).text.strip().replace('尺寸（mm）', '')
        casing10_Thickness = document.tables[2].cell(row_reference + 10, 8).text.strip().replace('尺寸（mm）', '')
        casing11_Thickness = document.tables[2].cell(row_reference + 11, 8).text.strip().replace('尺寸（mm）', '')
        casing12_Thickness = document.tables[2].cell(row_reference + 12, 8).text.strip().replace('尺寸（mm）', '')
        self.tableWidget_7.setItem(0, 2, QTableWidgetItem(str(casing1_Thickness)))
        self.tableWidget_7.setItem(1, 2, QTableWidgetItem(str(casing2_Thickness)))
        self.tableWidget_7.setItem(2, 2, QTableWidgetItem(str(casing3_Thickness)))
        self.tableWidget_7.setItem(3, 2, QTableWidgetItem(str(casing4_Thickness)))
        self.tableWidget_7.setItem(4, 2, QTableWidgetItem(str(casing5_Thickness)))

        # 避免套管下深井段为单数字而不为井段
        casing1_bottom = ''
        casing2_bottom = ''
        casing3_bottom = ''
        casing4_bottom = ''
        casing5_bottom = ''
        casing6_bottom = ''
        casing7_bottom = ''
        casing8_bottom = ''
        casing9_bottom = ''
        casing10_bottom = ''
        casing11_bottom = ''
        casing12_bottom = ''

        casing1_interval = document.tables[2].cell(row_reference + 1, 9).text.strip().replace('测量井段（m）',
                                                                                              '').replace('m', '')
        casing1_interval = casing1_interval.replace(' ', '')
        casing1_interval = casing1_interval.replace('～', '-')
        casing1_interval = casing1_interval.replace('~', '-')
        if '~' not in casing1_interval and '～' not in casing1_interval and \
                '-' not in casing1_interval and casing1_interval != '':
            casing1_interval = ''.join(['0', '-', casing1_interval])
        if casing1_interval != '':
            casing1_bottom = casing1_interval.split('-')[1]
        self.tableWidget_7.setItem(0, 3, QTableWidgetItem(str(casing1_bottom)))
        self.tableWidget_7.setItem(0, 4, QTableWidgetItem(str(casing1_interval)))

        casing2_interval = document.tables[2].cell(row_reference + 2, 9).text.strip().replace('测量井段（m）',
                                                                                              '').replace('m', '')
        casing2_interval = casing2_interval.replace(' ', '')
        casing2_interval = casing2_interval.replace('～', '-')
        casing2_interval = casing2_interval.replace('~', '-')
        if '~' not in casing2_interval and '～' not in casing2_interval and \
                '-' not in casing2_interval and casing2_interval != '':
            casing2_interval = ''.join(['0', '-', casing2_interval])
        if casing2_interval != '':
            casing2_bottom = casing2_interval.split('-')[1]
        self.tableWidget_7.setItem(1, 3, QTableWidgetItem(str(casing2_bottom)))
        self.tableWidget_7.setItem(1, 4, QTableWidgetItem(str(casing2_interval)))

        casing3_interval = document.tables[2].cell(row_reference + 3, 9).text.strip().replace('测量井段（m）',
                                                                                              '').replace('m', '')
        casing3_interval = casing3_interval.replace(' ', '')
        casing3_interval = casing3_interval.replace('～', '-')
        casing3_interval = casing3_interval.replace('~', '-')
        if '~' not in casing3_interval and '～' not in casing3_interval and \
                '-' not in casing3_interval and casing3_interval != '':
            casing3_interval = ''.join(['0', '-', casing3_interval])
        if casing3_interval != '':
            casing3_bottom = casing3_interval.split('-')[1]
        self.tableWidget_7.setItem(2, 3, QTableWidgetItem(str(casing3_bottom)))
        self.tableWidget_7.setItem(2, 4, QTableWidgetItem(str(casing3_interval)))

        casing4_interval = document.tables[2].cell(row_reference + 4, 9).text.strip().replace('测量井段（m）',
                                                                                              '').replace('m', '')
        casing4_interval = casing4_interval.replace(' ', '')
        casing4_interval = casing4_interval.replace('～', '-')
        casing4_interval = casing4_interval.replace('~', '-')
        if '~' not in casing4_interval and '～' not in casing4_interval and \
                '-' not in casing4_interval and casing4_interval != '':
            casing4_interval = ''.join(['0', '-', casing4_interval])
        if casing4_interval != '':
            casing4_bottom = casing4_interval.split('-')[1]
        self.tableWidget_7.setItem(3, 3, QTableWidgetItem(str(casing4_bottom)))
        self.tableWidget_7.setItem(3, 4, QTableWidgetItem(str(casing4_interval)))

        casing5_interval = document.tables[2].cell(row_reference + 5, 9).text.strip().replace('测量井段（m）',
                                                                                              '').replace('m', '')
        casing5_interval = casing5_interval.replace(' ', '')
        casing5_interval = casing5_interval.replace('～', '-')
        casing5_interval = casing5_interval.replace('~', '-')
        if '~' not in casing5_interval and '～' not in casing5_interval and \
                '-' not in casing5_interval and casing5_interval != '':
            casing5_interval = ''.join(['0', '-', casing5_interval])
        if casing5_interval != '':
            casing5_bottom = casing5_interval.split('-')[1]
        self.tableWidget_7.setItem(4, 3, QTableWidgetItem(str(casing5_bottom)))
        self.tableWidget_7.setItem(4, 4, QTableWidgetItem(str(casing5_interval)))

        casing6_interval = document.tables[2].cell(row_reference + 6, 9).text.strip().replace('测量井段（m）',
                                                                                              '').replace('m', '')
        casing6_interval = casing6_interval.replace(' ', '')
        casing6_interval = casing6_interval.replace('～', '-')
        casing6_interval = casing6_interval.replace('~', '-')
        if '~' not in casing6_interval and '～' not in casing6_interval and \
                '-' not in casing6_interval and casing6_interval != '':
            casing6_interval = ''.join(['0', '-', casing6_interval])
        if casing6_interval != '':
            casing6_bottom = casing6_interval.split('-')[1]

        casing7_interval = document.tables[2].cell(row_reference + 7, 9).text.strip().replace('测量井段（m）',
                                                                                              '').replace('m', '')
        casing7_interval = casing7_interval.replace(' ', '')
        casing7_interval = casing7_interval.replace('～', '-')
        casing7_interval = casing7_interval.replace('~', '-')
        if '~' not in casing7_interval and '～' not in casing7_interval and \
                '-' not in casing7_interval and casing7_interval != '':
            casing7_interval = ''.join(['0', '-', casing7_interval])
        if casing7_interval != '':
            casing7_bottom = casing7_interval.split('-')[1]

        casing8_interval = document.tables[2].cell(row_reference + 8, 9).text.strip().replace('测量井段（m）',
                                                                                              '').replace('m', '')
        casing8_interval = casing8_interval.replace(' ', '')
        casing8_interval = casing8_interval.replace('～', '-')
        casing8_interval = casing8_interval.replace('~', '-')
        if '~' not in casing8_interval and '～' not in casing8_interval and \
                '-' not in casing8_interval and casing8_interval != '':
            casing8_interval = ''.join(['0', '-', casing8_interval])
        if casing8_interval != '':
            casing8_bottom = casing8_interval.split('-')[1]

        casing9_interval = document.tables[2].cell(row_reference + 9, 9).text.strip().replace('测量井段（m）',
                                                                                              '').replace('m', '')
        casing9_interval = casing9_interval.replace(' ', '')
        casing9_interval = casing9_interval.replace('～', '-')
        casing9_interval = casing9_interval.replace('~', '-')
        if '~' not in casing9_interval and '～' not in casing9_interval and \
                '-' not in casing9_interval and casing9_interval != '':
            casing9_interval = ''.join(['0', '-', casing9_interval])
        if casing9_interval != '':
            casing9_bottom = casing9_interval.split('-')[1]

        casing10_interval = document.tables[2].cell(row_reference + 10, 9).text.strip().replace('测量井段（m）',
                                                                                                '').replace('m', '')
        casing10_interval = casing10_interval.replace(' ', '')
        casing10_interval = casing10_interval.replace('～', '-')
        casing10_interval = casing10_interval.replace('~', '-')
        if '~' not in casing10_interval and '～' not in casing10_interval and \
                '-' not in casing10_interval and casing10_interval != '':
            casing10_interval = ''.join(['0', '-', casing10_interval])
        if casing10_interval != '':
            casing10_bottom = casing10_interval.split('-')[1]

        casing11_interval = document.tables[2].cell(row_reference + 11, 9).text.strip().replace('测量井段（m）',
                                                                                                '').replace('m', '')
        casing11_interval = casing11_interval.replace(' ', '')
        casing11_interval = casing11_interval.replace('～', '-')
        casing11_interval = casing11_interval.replace('~', '-')
        if '~' not in casing11_interval and '～' not in casing11_interval and \
                '-' not in casing11_interval and casing11_interval != '':
            casing11_interval = ''.join(['0', '-', casing11_interval])
        if casing11_interval != '':
            casing11_bottom = casing11_interval.split('-')[1]

        casing12_interval = document.tables[2].cell(row_reference + 12, 9).text.strip().replace('测量井段（m）',
                                                                                                '').replace('m', '')
        casing12_interval = casing12_interval.replace(' ', '')
        casing12_interval = casing12_interval.replace('～', '-')
        casing12_interval = casing12_interval.replace('~', '-')
        if '~' not in casing12_interval and '～' not in casing12_interval and \
                '-' not in casing12_interval and casing12_interval != '':
            casing12_interval = ''.join(['0', '-', casing12_interval])
        if casing12_interval != '':
            casing12_bottom = casing12_interval.split('-')[1]

        # 目标套管尺寸casing_Goal
        temp_list = ['', '类  型', '硬塑料扶正器', '弹簧扶正器', '橡胶扶正器', '橡胶', '外径（mm）']
        if casing12_Dia not in temp_list:
            casing_Goal = casing12_Dia
        elif casing11_Dia not in temp_list:
            casing_Goal = casing11_Dia
        elif casing10_Dia not in temp_list:
            casing_Goal = casing10_Dia
        elif casing9_Dia not in temp_list:
            casing_Goal = casing9_Dia
        elif casing8_Dia not in temp_list:
            casing_Goal = casing8_Dia
        elif casing7_Dia not in temp_list:
            casing_Goal = casing7_Dia
        elif casing6_Dia not in temp_list:
            casing_Goal = casing6_Dia
        elif casing5_Dia not in temp_list:
            casing_Goal = casing5_Dia
        elif casing4_Dia not in temp_list:
            casing_Goal = casing4_Dia
        elif casing3_Dia not in temp_list:
            casing_Goal = casing3_Dia
        elif casing2_Dia not in temp_list:
            casing_Goal = casing2_Dia
        elif casing1_Dia not in temp_list:
            casing_Goal = casing1_Dia
        self.lineEdit_108.setText(casing_Goal)

        ######################################################################## 目标套管下深casing_Goal_Depth
        if casing12_Dia == casing_Goal:
            if '～' in document.tables[2].cell(row_reference + 12, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 12, 9).text.strip().split('～')[1]
            elif '~' in document.tables[2].cell(row_reference + 12, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 12, 9).text.strip().split('~')[1]
            elif '-' in document.tables[2].cell(row_reference + 12, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 12, 9).text.strip().split('-')[1]
            elif '-' not in document.tables[2].cell(row_reference + 12, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 12, 9).text.strip()
        elif casing11_Dia == casing_Goal:
            if '～' in document.tables[2].cell(row_reference + 11, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 11, 9).text.strip().split('～')[1]
            elif '~' in document.tables[2].cell(row_reference + 11, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 11, 9).text.strip().split('~')[1]
            elif '-' in document.tables[2].cell(row_reference + 11, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 11, 9).text.strip().split('-')[1]
            elif '-' not in document.tables[2].cell(row_reference + 11, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 11, 9).text.strip()
        elif casing10_Dia == casing_Goal:
            if '～' in document.tables[2].cell(row_reference + 10, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 10, 9).text.strip().split('～')[1]
            elif '~' in document.tables[2].cell(row_reference + 10, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 10, 9).text.strip().split('~')[1]
            elif '-' in document.tables[2].cell(row_reference + 10, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 10, 9).text.strip().split('-')[1]
            elif '-' not in document.tables[2].cell(row_reference + 10, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 10, 9).text.strip()
        elif casing9_Dia == casing_Goal:
            if '～' in document.tables[2].cell(row_reference + 9, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 9, 9).text.strip().split('～')[1]
            elif '~' in document.tables[2].cell(row_reference + 9, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 9, 9).text.strip().split('~')[1]
            elif '-' in document.tables[2].cell(row_reference + 9, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 9, 9).text.strip().split('-')[1]
            elif '-' not in document.tables[2].cell(row_reference + 9, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 9, 9).text.strip()
        elif casing8_Dia == casing_Goal:
            if '～' in document.tables[2].cell(row_reference + 8, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 8, 9).text.strip().split('～')[1]
            elif '~' in document.tables[2].cell(row_reference + 8, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 8, 9).text.strip().split('~')[1]
            elif '-' in document.tables[2].cell(row_reference + 8, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 8, 9).text.strip().split('-')[1]
            elif '-' not in document.tables[2].cell(row_reference + 8, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 8, 9).text.strip()
        elif casing7_Dia == casing_Goal:
            if '～' in document.tables[2].cell(row_reference + 7, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 7, 9).text.strip().split('～')[1]
            elif '~' in document.tables[2].cell(row_reference + 7, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 7, 9).text.strip().split('~')[1]
            elif '-' in document.tables[2].cell(row_reference + 7, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 7, 9).text.strip().split('-')[1]
            elif '-' not in document.tables[2].cell(row_reference + 7, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 7, 9).text.strip()
        elif casing6_Dia == casing_Goal:
            if '～' in document.tables[2].cell(row_reference + 6, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 6, 9).text.strip().split('～')[1]
            elif '~' in document.tables[2].cell(row_reference + 6, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 6, 9).text.strip().split('~')[1]
            elif '-' in document.tables[2].cell(row_reference + 6, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 6, 9).text.strip().split('-')[1]
            elif '-' not in document.tables[2].cell(row_reference + 6, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 6, 9).text.strip()
        elif casing5_Dia == casing_Goal:
            if '～' in document.tables[2].cell(row_reference + 5, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 5, 9).text.strip().split('～')[1]
            elif '~' in document.tables[2].cell(row_reference + 5, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 5, 9).text.strip().split('~')[1]
            elif '-' in document.tables[2].cell(row_reference + 5, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 5, 9).text.strip().split('-')[1]
            elif '-' not in document.tables[2].cell(row_reference + 5, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 5, 9).text.strip()
        elif casing4_Dia == casing_Goal:
            if '～' in document.tables[2].cell(row_reference + 4, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 4, 9).text.strip().split('～')[1]
            elif '~' in document.tables[2].cell(row_reference + 4, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 4, 9).text.strip().split('~')[1]
            elif '-' in document.tables[2].cell(row_reference + 4, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 4, 9).text.strip().split('-')[1]
            elif '-' not in document.tables[2].cell(row_reference + 4, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 4, 9).text.strip()
        elif casing3_Dia == casing_Goal:
            if '～' in document.tables[2].cell(row_reference + 3, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 3, 9).text.strip().split('～')[1]
            elif '~' in document.tables[2].cell(row_reference + 3, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 3, 9).text.strip().split('~')[1]
            elif '-' in document.tables[2].cell(row_reference + 3, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 3, 9).text.strip().split('-')[1]
            elif '-' not in document.tables[2].cell(row_reference + 3, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 3, 9).text.strip()
        elif casing2_Dia == casing_Goal:
            if '～' in document.tables[2].cell(row_reference + 2, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 2, 9).text.strip().split('～')[1]
            elif '~' in document.tables[2].cell(row_reference + 2, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 2, 9).text.strip().split('~')[1]
            elif '-' in document.tables[2].cell(row_reference + 2, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 2, 9).text.strip().split('-')[1]
            elif '-' not in document.tables[2].cell(row_reference + 2, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 2, 9).text.strip()
        elif casing1_Dia == casing_Goal:
            if '～' in document.tables[2].cell(row_reference + 1, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 1, 9).text.strip().split('～')[1]
            elif '~' in document.tables[2].cell(row_reference + 1, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 1, 9).text.strip().split('~')[1]
            elif '-' in document.tables[2].cell(row_reference + 1, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 1, 9).text.strip().split('-')[1]
            elif '-' not in document.tables[2].cell(row_reference + 1, 9).text.strip():
                casing_Goal_Depth = document.tables[2].cell(row_reference + 1, 9).text.strip()
        self.lineEdit_109.setText(casing_Goal_Depth)

        ######################################################################## 获取测量井段
        for row in range(3, 26):
            if document.tables[3].cell(row, 5).text.replace(' ', '').replace('点测', '') != '':
                measure_Interval = document.tables[3].cell(row, 5).text.replace(' ', '').replace('点测', '')
                measure_Interval = measure_Interval.replace('~', '-')
                measure_Interval = measure_Interval.replace('～', '-')
                measure_Interval = measure_Interval.replace('m', '')
                measure_Interval_Start_Depth = measure_Interval.split('-')[0]
                measure_Interval_End_Depth = measure_Interval.split('-')[1]
                pass
        try:  # 强制类型转换后增加小数位数
            measure_Interval_Start_Depth = str(round(float(measure_Interval_Start_Depth), 1))
            measure_Interval_End_Depth = str(round(float(measure_Interval_End_Depth), 1))
        except:
            pass
        self.lineEdit_110.setText(measure_Interval_Start_Depth)
        self.lineEdit_111.setText(measure_Interval_End_Depth)

    def automate_table_helper(self):
        # 读取单层统计表
        PATH = ".\\WorkSpace\\报告生成工区\\成果表"
        fileName = ''
        for fileName in os.listdir(PATH):
            if '1单' in fileName and '.xls' in fileName and '$' not in fileName:
                fileDir = PATH + "\\" + fileName
                try:
                    # 表格表头字段规范
                    self.xls_formatting_first_layer(fileDir)
                finally:
                    pass
                # fileDir = ''.join([fileDir.replace('.' + fileDir.split('.')[-1], ''), '(已规范化).xls'])
                workbook = xlrd.open_workbook(fileDir)

        if fileName != '':
            sheet = workbook.sheets()[0]

            # 获得表单的行数及列数
            nrow = sheet.nrows
            ncol = sheet.ncols
            # 处理评价井段
            start_Evaluation = str(sheet.cell_value(3, 1)).strip()
            start_Evaluation = start_Evaluation.split('-')[0]
            end_Evaluation = str(sheet.cell_value(nrow - 1, 1)).strip('')
            end_Evaluation = ''.join(end_Evaluation.split())  # 去除所有空格
            end_Evaluation = end_Evaluation.split('-')[1]

            self.lineEdit_103.setText(start_Evaluation)
            self.lineEdit_105.setText(end_Evaluation)

            # 自动补充井次名
            self.set_well_detail_name()
        elif fileName == '':
            QMessageBox.information(self, "提示", "成果表放置后该按钮才有用噢")
        else:
            QMessageBox.information(self, "提示", "遇到了一点问题，请联系开发者")

    def generate_txt_file(self):
        if self.lineEdit_5.text() == '' or '路径' in self.lineEdit_5.text():
            self.error_animation()
            self.lineEdit_5.setText('您还没有选中登记表路径！')
        else:
            if self.run_on_net == True:
                Supervisor.lead_txt_usage_supervisor()
            else:
                pass

            well_Name = self.lineEdit.text()
            well_Category = self.comboBox_7.currentText()
            well_Type = self.comboBox_6.currentText()
            well_Depth = self.lineEdit_4.text()
            x_Coordinate = self.lineEdit_9.text()
            y_Coordinate = self.lineEdit_10.text()
            magnetic_Declination = self.lineEdit_14.text()
            ground_Elevation = self.lineEdit_12.text()
            bushing_Height = self.lineEdit_13.text()
            kelly_Bushing = self.lineEdit_11.text()
            oil_Field = self.lineEdit_17.text()
            client_Name = self.lineEdit_50.text()
            geo_Position = self.lineEdit_16.text()
            stru_Position = self.lineEdit_18.text()
            completion_Date = self.lineEdit_39.text()
            spud_Date = self.lineEdit_20.text()
            end_Drilling_Date = self.lineEdit_21.text()
            deepest_bit = self.lineEdit_102.text()
            drilling_Unit = self.lineEdit_74.text()
            max_Well_Deviation_Depth = self.lineEdit_24.text()
            max_Well_Deviation = self.lineEdit_28.text()
            bit1_Diameter = self.tableWidget_6.item(0, 0).text()
            bit1_Depth = self.tableWidget_6.item(0, 1).text()
            bit2_Diameter = self.tableWidget_6.item(1, 0).text()
            bit2_Depth = self.tableWidget_6.item(1, 1).text()
            bit3_Diameter = self.tableWidget_6.item(2, 0).text()
            bit3_Depth = self.tableWidget_6.item(2, 1).text()
            bit4_Diameter = self.tableWidget_6.item(3, 0).text()
            bit4_Depth = self.tableWidget_6.item(3, 1).text()
            bit5_Diameter = self.tableWidget_6.item(4, 0).text()
            bit5_Depth = self.tableWidget_6.item(4, 1).text()
            well_Times_Name = self.lineEdit_3.text()
            well_Times_Type = self.lineEdit_49.text()
            logging_Date = self.lineEdit_121.text()
            measure_Interval_Start_Depth = self.lineEdit_110.text()
            measure_Interval_End_Depth = self.lineEdit_111.text()
            logging_Equipment = self.comboBox_5.currentText()
            arti_Bottom = self.lineEdit_65.text()
            evaluation_start_depth = self.lineEdit_103.text()
            evaluation_end_depth = self.lineEdit_105.text()
            task_Number = self.lineEdit_19.text()
            logging_Group = self.lineEdit_43.text()
            logging_Leader = self.lineEdit_54.text()
            logging_Operator = self.lineEdit_44.text()
            report_Writer = self.comboBox_2.currentText()
            report_Checker = self.comboBox.currentText()
            report_Supervisor = self.comboBox_4.currentText()
            flu_Property = self.lineEdit_60.text()
            flu_Density = self.lineEdit_61.text()
            flu_Viscosity = self.lineEdit_62.text()
            flu_Fluid_RT = self.lineEdit_66.text()
            flu_Fluid_TEMP = self.lineEdit_67.text()
            flu_Fluid_PH = self.lineEdit_68.text()
            flu_Fluid_Lost_Water = self.lineEdit_63.text()
            design_Depth = self.lineEdit_69.text()
            actual_Depth = self.lineEdit_70.text()
            cement_Quantity = self.lineEdit_71.text()
            cement_Density = self.lineEdit_72.text()
            cement_End_Date = self.lineEdit_101.text()
            cement_Unit = self.lineEdit_25.text()
            encounter_obstacle_depth =self.lineEdit_115.text().replace(' ', '')
            raw_Data_Recieve_Date = self.lineEdit_33.text()
            raw_Data_Evaluate_Date = self.lineEdit_34.text()
            interpretation_Start_Date = self.lineEdit_30.text()
            interpretation_Complete_Date = self.lineEdit_55.text()
            data_Archive_Date = self.lineEdit_32.text()



            casing1_Inner_Dia = self.tableWidget_7.item(0, 0).text()
            casing1_Dia = self.tableWidget_7.item(0, 1).text()
            casing1_Thickness = self.tableWidget_7.item(0, 2).text()
            casing1_bottom = self.tableWidget_7.item(0, 3).text()

            casing2_Inner_Dia = self.tableWidget_7.item(1, 0).text()
            casing2_Dia = self.tableWidget_7.item(1, 1).text()
            casing2_Thickness = self.tableWidget_7.item(1, 2).text()
            casing2_bottom = self.tableWidget_7.item(1, 3).text()

            casing3_Inner_Dia = self.tableWidget_7.item(2, 0).text()
            casing3_Dia = self.tableWidget_7.item(2, 1).text()
            casing3_Thickness = self.tableWidget_7.item(2, 2).text()
            casing3_bottom = self.tableWidget_7.item(2, 3).text()

            casing4_Inner_Dia = self.tableWidget_7.item(3, 0).text()
            casing4_Dia = self.tableWidget_7.item(3, 1).text()
            casing4_Thickness = self.tableWidget_7.item(3, 2).text()
            casing4_bottom = self.tableWidget_7.item(3, 3).text()

            casing5_Inner_Dia = self.tableWidget_7.item(4, 0).text()
            casing5_Dia = self.tableWidget_7.item(4, 1).text()
            casing5_Thickness = self.tableWidget_7.item(4, 2).text()
            casing5_bottom = self.tableWidget_7.item(4, 3).text()

            cement_End_Time = self.lineEdit_101.text()
            logging_Start_Time = self.lineEdit_104.text()
            logging_End_Time = self.lineEdit_121.text()
            logging_Method = self.lineEdit_100.text()

            logging_Company = self.lineEdit_15.text()
            up_company = self.lineEdit_29.text()

            # 避免报告中出现'-99999'
            if cement_Quantity == '-99999' or '':
                cement_Quantity = '/'
            else:
                pass

            if arti_Bottom == '-99999' or '':
                arti_Bottom = '/'
            else:
                pass

            # 新规定，起始评价深度从0米开始
            # if float(evaluation_start_depth) < 200:
            #     evaluation_start_depth = '0'

            DICT_TXT = {
                "井名": well_Name,
                "井别": well_Category,
                "井型": well_Type,
                "井深": well_Depth,
                "X坐标": x_Coordinate,
                "Y坐标": y_Coordinate,
                "经度": '-99999',
                "纬度": '-99999',
                "磁偏角": magnetic_Declination,
                "地面海拔": ground_Elevation,
                "补心高度": bushing_Height,
                "补心海拔": kelly_Bushing,
                "油田": oil_Field,
                "甲方单位": client_Name,
                "地理位置": geo_Position,
                "构造位置": stru_Position,
                "完井日期": completion_Date,
                "开钻日期": spud_Date,
                "完钻日期": end_Drilling_Date,
                "设计井深": '-99999',
                "完钻井深": deepest_bit,
                "钻井单位": drilling_Unit,
                "最大井斜深度": max_Well_Deviation_Depth,
                "最大井斜斜度": max_Well_Deviation,
                "钻头1直径": bit1_Diameter,
                "钻头1深度": bit1_Depth,
                "钻头2直径": bit2_Diameter,
                "钻头2深度": bit2_Depth,
                "钻头3直径": bit3_Diameter,
                "钻头3深度": bit3_Depth,
                "钻头4直径": bit4_Diameter,
                "钻头4深度": bit4_Depth,
                "钻头5直径": bit5_Diameter,
                "钻头5深度": bit5_Depth,
                "井次名称": well_Times_Name,
                "井次类别": well_Times_Type,
                "测井日期": logging_Date,
                "测井顶部深度": measure_Interval_Start_Depth,
                "测井底部深度": measure_Interval_End_Depth,
                "标准段顶部深度": measure_Interval_Start_Depth,
                "标准段底部深度": measure_Interval_End_Depth,
                "测井公司": logging_Company,
                "测井类型": '工程测井',
                "解释项目": '固井质量检测',
                "测井装备": logging_Equipment,
                "测井仪器": '三组合变密度',
                "测时井深": arti_Bottom,
                "资料接收日期": raw_Data_Recieve_Date,
                "曲线评定完成日期": raw_Data_Evaluate_Date,
                "解释开始日期": interpretation_Start_Date,
                "解释完成日期": interpretation_Complete_Date,
                "资料归档日期": data_Archive_Date,
                "解释段顶部深度": evaluation_start_depth,
                "解释段底部深度": evaluation_end_depth,
                "任务单号": task_Number,
                "测井小队": logging_Group,
                "小队长": logging_Leader,
                "操作员": logging_Operator,
                "监督员": report_Writer,
                "处理员": report_Writer,
                "校对员": report_Checker,
                "审核员": report_Supervisor,
                "泥浆类型": flu_Property,
                "泥浆密度": flu_Density,
                "泥浆粘度": flu_Viscosity,
                "泥浆电阻率": flu_Fluid_RT,
                "泥浆温度": flu_Fluid_TEMP,
                "泥浆滤液电阻率": -99999,
                "泥浆滤液温度": -99999,
                "泥浆PH值": flu_Fluid_PH,
                "泥浆失水": flu_Fluid_Lost_Water,
                "井内液体": flu_Property,
                "人工井底": arti_Bottom,
                "预计水泥返高": design_Depth,
                "实际水泥返高": actual_Depth,
                "注水泥量": cement_Quantity,
                "水泥密度": cement_Density,
                "固井日期": cement_End_Date,
                "固井单位": cement_Unit,
                "套管1内径": casing1_Inner_Dia,
                "套管1外径": casing1_Dia,
                "套管1壁厚": casing1_Thickness,
                "套管1终深": casing1_bottom,
                "套管2内径": casing2_Inner_Dia,
                "套管2外径": casing2_Dia,
                "套管2壁厚": casing2_Thickness,
                "套管2终深": casing2_bottom,
                "套管3内径": casing3_Inner_Dia,
                "套管3外径": casing3_Dia,
                "套管3壁厚": casing3_Thickness,
                "套管3终深": casing3_bottom,
                "套管4内径": casing4_Inner_Dia,
                "套管4外径": casing4_Dia,
                "套管4壁厚": casing4_Thickness,
                "套管4终深": casing4_bottom,
                "套管5内径": casing5_Inner_Dia,
                "套管5外径": casing5_Dia,
                "套管5壁厚": casing5_Thickness,
                "套管5终深": casing5_bottom,
                "MINTHICK": 1,
                "钻井深度": deepest_bit,
                "固井结束时间": cement_End_Time,
                "测井开始时间": logging_Start_Time,
                "测井结束时间": logging_End_Time,
                "测井工艺": logging_Method,
                "测井软件": 'CIFLog-LEAD',
                "上传单位": up_company,
                "仪器遇阻位置": encounter_obstacle_depth
            }

            # 去除字典中的斜杠
            for key, value in DICT_TXT.items():
                if value == '/':
                    DICT_TXT[key] = ''

            temp = self.lineEdit_5.text().split('/')[-1]
            temp2 = self.lineEdit_5.text().replace(temp, '')

            f = open(temp2 + well_Name + '-井信息(LEAD4.0).txt', 'w', encoding='UTF-8')
            for key, value in DICT_TXT.items():
                f.write(key + '=' + str(value) + '\n')
            f.close()
            QMessageBox.information(self, "提示", "井信息文件已生成，和原始资料登记表在同一目录")

    def bit_info_table(self):
        self.tableWidget_6.setColumnCount(2)
        self.tableWidget_6.setRowCount(5)
        row = 0  # 第几行（从0开始）
        col = 0  # 第几列（从0开始）
        # self.tableWidget_6.horizontalHeader().setStretchLastSection(True)  # 设置最后一列拉伸至最大
        self.tableWidget_6.horizontalHeader().setSectionsClickable(False)  # 禁止点击表头的列
        self.headers = ['直径', '深度']
        self.tableWidget_6.setHorizontalHeaderLabels(self.headers)

        self.tableWidget_6.setColumnWidth(0, 70)
        self.tableWidget_6.setColumnWidth(1, 80)

        self.tableWidget_6.setRowHeight(0, 30)
        self.tableWidget_6.setRowHeight(1, 30)
        self.tableWidget_6.setRowHeight(2, 30)
        self.tableWidget_6.setRowHeight(3, 30)
        self.tableWidget_6.setRowHeight(4, 30)

        # 值初始化，否则会出现AttributeError: 'NoneType' object has no attribute 'text'
        self.tableWidget_6.setItem(0, 0, QTableWidgetItem(str(' ')))
        self.tableWidget_6.setItem(0, 1, QTableWidgetItem(str(' ')))
        self.tableWidget_6.setItem(1, 0, QTableWidgetItem(str(' ')))
        self.tableWidget_6.setItem(1, 1, QTableWidgetItem(str(' ')))
        self.tableWidget_6.setItem(2, 0, QTableWidgetItem(str(' ')))
        self.tableWidget_6.setItem(2, 1, QTableWidgetItem(str(' ')))
        self.tableWidget_6.setItem(3, 0, QTableWidgetItem(str(' ')))
        self.tableWidget_6.setItem(3, 1, QTableWidgetItem(str(' ')))
        self.tableWidget_6.setItem(4, 0, QTableWidgetItem(str(' ')))
        self.tableWidget_6.setItem(4, 1, QTableWidgetItem(str(' ')))

        # self.tableWidget_6.setRowHeight(0, 50)
        # self.tableWidget_6.verticalHeader().setVisible(False)  # 隐藏垂直表头
        # self.tableWidget_6.horizontalHeader().setVisible(False)  # 隐藏水平表头

    def casing_info_table(self):
        self.tableWidget_7.setColumnCount(5)
        self.tableWidget_7.setRowCount(5)
        row = 0  # 第几行（从0开始）
        col = 0  # 第几列（从0开始）
        # self.tableWidget_7.horizontalHeader().setStretchLastSection(True)  # 设置最后一列拉伸至最大
        self.tableWidget_7.horizontalHeader().setSectionsClickable(False)  # 禁止点击表头的列
        self.headers = ['内径', '外径', '壁厚', '终深', '深度段']
        self.tableWidget_7.setHorizontalHeaderLabels(self.headers)

        self.tableWidget_7.setColumnWidth(0, 80)
        self.tableWidget_7.setColumnWidth(1, 80)
        self.tableWidget_7.setColumnWidth(2, 80)
        self.tableWidget_7.setColumnWidth(3, 80)
        self.tableWidget_7.setColumnWidth(4, 150)

        self.tableWidget_7.setRowHeight(0, 30)
        self.tableWidget_7.setRowHeight(1, 30)
        self.tableWidget_7.setRowHeight(2, 30)
        self.tableWidget_7.setRowHeight(3, 30)
        self.tableWidget_7.setRowHeight(4, 30)

        # 值初始化
        self.tableWidget_7.setItem(0, 0, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(0, 1, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(0, 2, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(0, 3, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(0, 4, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(1, 0, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(1, 1, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(1, 2, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(1, 3, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(1, 4, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(2, 0, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(2, 1, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(2, 2, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(2, 3, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(2, 4, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(3, 0, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(3, 1, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(3, 2, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(3, 3, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(3, 4, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(4, 0, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(4, 1, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(4, 2, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(4, 3, QTableWidgetItem(str(' ')))
        self.tableWidget_7.setItem(4, 4, QTableWidgetItem(str(' ')))

        # self.tableWidget_7.setRowHeight(0, 50)
        # self.tableWidget_7.verticalHeader().setVisible(False)  # 隐藏垂直表头
        # self.tableWidget_7.horizontalHeader().setVisible(False)  # 隐藏水平表头

    ################################################################################
    # 函数定义集结地

    def mkdir(self, path):
        path = path.strip()  # 去除首位空格
        path = path.rstrip("\\")  # 去除尾部 \ 符号
        isExists = os.path.exists(path)
        if not isExists:
            os.makedirs(path)
            # print(path + ' 创建成功')
            return True
        else:
            # print(path + ' 目录已存在')
            return False

    # 定义一个函数，增加重新计算后的厚度列
    def get_thickness(self, x):
        thickness = x['井段End'] - x['井段Start']
        return thickness

    # 定义进度条函数，用作进度展示
    def view_bar(self, num, total):
        rate = float(num) / float(total)
        rate_num = int(rate * 100)
        r = '\r[%s%s]%d%%' % ("*" * rate_num, " " * (100 - rate_num), rate_num)
        sys.stdout.write(r)
        sys.stdout.flush()

    # 函数，获取文件路径、文件名、后缀名
    def get_filePath_fileName_fileExt(self, filename):
        (filepath, tempfilename) = os.path.split(filename)
        (shotname, extension) = os.path.splitext(tempfilename)
        return filepath, shotname, extension

    # 文档替换主程序
    def document_replace(self):
        # 补充信息
        well_Name = self.lineEdit.text()
        casing_Goal = self.lineEdit_108.text()
        logging_Date = self.lineEdit_40.text()
        try:
            year = logging_Date.split('-')[0]
            month = logging_Date.split('-')[1]
            day = logging_Date.split('-')[2]
        except:
            year = ''
            month = ''
            day = ''
        first_Pro_Interval = ''.join([self.lineEdit_103.text(), '-', self.lineEdit_105.text()])
        # 新规定，生成一个从0开始的处理深度， 为了多计费
        # TODO
        if float(self.lineEdit_103.text()) < 200:
            measure_start = '0'
            self.measure_from_Pro = ''.join([measure_start, '-', self.lineEdit_105.text()])
        else:
            self.measure_from_Pro = ''.join([self.lineEdit_103.text(), '-', self.lineEdit_105.text()])

        PATH = '.\\WorkSpace\\报告生成工区\\'
        TEMPLATE_PATH = '.\\resources\\模板\\'
        # newFile = PATH + well_Name + '_' + year + month + \
        #           day + '_(' + casing_Goal + 'mm套,VDL_' + first_Pro_Interval + 'm)固井报告' + '.docx'
        newFile = PATH + well_Name + '_固井质量测井评价报告_' + year + month + \
                  day + '_' + casing_Goal + 'mm套' + '.docx'
        if self.radioButton_3.isChecked() == True:
            document = Document(TEMPLATE_PATH + 'template-of-cbl-vdl-report.docx')
        elif self.radioButton_4.isChecked() == True:
            document = Document(TEMPLATE_PATH + 'template-of-cbl-vdl-report-interval.docx')
        elif self.radioButton_7.isChecked() == True:
            document = Document(TEMPLATE_PATH + 'template-of-cbl-vdl-report.docx')

        if self.checkBox_9.isChecked():
            document = self.check(document)  # 调用替换函数
        else:
            pass
        # 全文档表格内容居中
        for table in document.tables:
            for row in range(len(table.rows)):
                for col in range(len(table.columns)):
                    table.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # 整体设置，未起作用
            # table.style.font.color.rgb = RGBColor(255, 0, 0)
            # table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.save(newFile)

    def check(self, document):
        # tables
        for table in document.tables:
            for row in range(len(table.rows)):
                self.view_bar(row, len(table.rows))  # 进度条
                for col in range(len(table.columns)):
                    for key, value in self.DICT.items():
                        if key in table.cell(row, col).text:
                            # print(key + " = " + value)
                            table.cell(row, col).text = table.cell(row, col).text.replace(key, value)

        # paragraphs
        for para in document.paragraphs:
            for i in range(len(para.runs)):
                # self.view_bar(i, len(para.runs) - 1)
                for key, value in self.DICT.items():
                    if key in para.runs[i].text:
                        # print(key + " = " + value)
                        para.runs[i].text = para.runs[i].text.replace(key, value)

        # sections
        for sec in document.sections:
            for i in range(len(sec.header.paragraphs)):
                for key, value in self.DICT.items():
                    if key in sec.header.paragraphs[i].text:
                        # print(key + " = " + value)
                        sec.header.paragraphs[i].text = sec.header.paragraphs[i].text.replace(key, value)
        return document

    ################################################################################
    # 一界面某井段评价函数
    def layer_evaluation1(self, df, start, end):
        df1 = df
        df1 = df1.reset_index()
        formation_Start = start
        formation_End = end
        # # 越界警告，不用加了，后面有考虑到越界情况
        # if df1.iloc[-1, df1.columns.get_loc('井段End')] < formation_End or df1.iloc[0, df1.columns.get_loc('井段Start')] > formation_Start:
        #     QMessageBox.information(self, '注意', '储层边界超过了单层评价表的范围。')
        # 截取我们想要的目标数据体
        df_temp = df1.loc[(df1['井段Start'] >= formation_Start) & (df1['井段Start'] <= formation_End), :]
        # 获取起始深度到第一层井段底界的结论
        df_temp_start_to_first_layer = df1.loc[(df1['井段Start'] <= formation_Start), :]
        if len(df_temp_start_to_first_layer) != 0:
            start_to_upper_result = df_temp_start_to_first_layer.loc[len(df_temp_start_to_first_layer) - 1, '结论']
        elif len(df_temp_start_to_first_layer) == 0:
            start_to_upper_result = df1.loc[0, '结论']
        # 获取calculation_Start所在段的声幅值
        df_temp_formation_Start = df1.loc[(df1['井段Start'] <= formation_Start) & (
                df1['井段End'] >= formation_Start), :]
        df_temp_formation_Start.reset_index(drop=True, inplace=True)  # 重新设置列索引#防止若截取中段，index不从0开始的bug
        # 补充储层界到井段的深度
        x, y = df_temp.shape
        df_temp = df_temp.reset_index()
        df_temp.drop(['index'], axis=1, inplace=True)
        if x != 0:  # 防止df_temp为空时，loc报错的bug
            first_layer_start = df_temp.loc[0, '井段Start']
        if x > 0 and first_layer_start != formation_Start:
            upper = pd.DataFrame({'井段(m)': ''.join([str(formation_Start), '-', str(first_layer_start)]),
                                  '厚度(m)': first_layer_start - formation_Start,
                                  '最大声幅(%)': df_temp_formation_Start.loc[0, '最大声幅(%)'],
                                  '最小声幅(%)': df_temp_formation_Start.loc[0, '最小声幅(%)'],
                                  '平均声幅(%)': df_temp_formation_Start.loc[0, '平均声幅(%)'],
                                  '结论': start_to_upper_result,
                                  '井段Start': formation_Start,
                                  '井段End': first_layer_start},
                                 index=[1])  # 自定义索引为：1 ，这里也可以不设置index
            df_temp.loc[len(df_temp) - 1, '井段End'] = formation_End
            df_temp = pd.concat([upper, df_temp], ignore_index=True)
            # df_temp = df_temp.append(new, ignore_index=True)  # ignore_index=True,表示不按原来的索引，从0开始自动递增
            df_temp.loc[:, "重计算厚度"] = df_temp.apply(self.get_thickness, axis=1)
            # 修改df_temp的最末一行
            df_temp.loc[len(df_temp) - 1, '井段(m)'] = ''.join([str(df_temp.loc[len(df_temp) - 1, '井段Start']), \
                                                                '-', str(df_temp.loc[len(df_temp) - 1, '井段End'])])
            df_temp.loc[len(df_temp) - 1, '厚度(m)'] = df_temp.loc[len(df_temp) - 1, '重计算厚度']
        elif x > 0 and first_layer_start == formation_Start:
            df_temp.loc[len(df_temp) - 1, '井段End'] = formation_End
            df_temp.loc[:, "重计算厚度"] = df_temp.apply(self.get_thickness, axis=1)
            # 修改df_temp的最末一行
            df_temp.loc[len(df_temp) - 1, '井段(m)'] = ''.join([str(df_temp.loc[len(df_temp) - 1, '井段Start']), \
                                                                '-', str(df_temp.loc[len(df_temp) - 1, '井段End'])])
            df_temp.loc[len(df_temp) - 1, '厚度(m)'] = df_temp.loc[len(df_temp) - 1, '重计算厚度']
        else:  # 储层包含在一个井段内的情况
            df_temp = pd.DataFrame({'井段(m)': ''.join([str(formation_Start), '-', str(formation_End)]),
                                    '厚度(m)': formation_End - formation_Start,
                                    '最大声幅(%)': df_temp_formation_Start.loc[0, '最大声幅(%)'],
                                    '最小声幅(%)': df_temp_formation_Start.loc[0, '最小声幅(%)'],
                                    '平均声幅(%)': df_temp_formation_Start.loc[0, '平均声幅(%)'],
                                    '结论': start_to_upper_result,
                                    '井段Start': formation_Start,
                                    '井段End': formation_End},
                                   index=[1])  # 自定义索引为：1 ，这里也可以不设置index
            df_temp.loc[:, "重计算厚度"] = df_temp.apply(self.get_thickness, axis=1)
            # 修改df_temp的最末一行
            df_temp.loc[len(df_temp), '井段(m)'] = ''.join([str(df_temp.loc[len(df_temp), '井段Start']),
                                                            '-', str(df_temp.loc[len(df_temp), '井段End'])])
            df_temp.loc[len(df_temp), '厚度(m)'] = df_temp.loc[len(df_temp), '重计算厚度']
        # print(df_temp)
        ratio_Series = df_temp.groupby(by=['结论'])['重计算厚度'].sum() / df_temp['重计算厚度'].sum() * 100

        if ratio_Series.__len__() == 2:
            if '胶结好' not in ratio_Series:
                ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
            elif '胶结中等' not in ratio_Series:
                ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
            elif '胶结差' not in ratio_Series:
                ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
        elif ratio_Series.__len__() == 1:
            if ('胶结好' not in ratio_Series) & ('胶结中等' not in ratio_Series):
                ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
            elif ('胶结好' not in ratio_Series) & ('胶结差' not in ratio_Series):
                ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
            elif ('胶结中等' not in ratio_Series) & ('胶结差' not in ratio_Series):
                ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))

        # 条件判断，参数需要研究
        if ratio_Series['胶结好'] >= 95:
            evaluation_of_formation = '胶结好'
        elif ratio_Series['胶结中等'] >= 95:
            evaluation_of_formation = '胶结中等'
        elif ratio_Series['胶结差'] >= 95:
            evaluation_of_formation = '胶结差'
        elif (95 >= ratio_Series['胶结好'] >= 5) & (95 >= ratio_Series['胶结中等'] >= 5) & (
                5 >= ratio_Series['胶结差']):
            if ratio_Series['胶结好'] >= ratio_Series['胶结中等']:
                evaluation_of_formation = '胶结中到好，以好为主'
            elif ratio_Series['胶结好'] <= ratio_Series['胶结中等']:
                evaluation_of_formation = '胶结中到好，以中等为主'
        elif (95 >= ratio_Series['胶结差'] >= 5) & (95 >= ratio_Series['胶结中等'] >= 5) & (
                5 >= ratio_Series['胶结好']):
            if ratio_Series['胶结差'] >= ratio_Series['胶结中等']:
                evaluation_of_formation = '胶结中到差，以差为主'
            elif ratio_Series['胶结差'] <= ratio_Series['胶结中等']:
                evaluation_of_formation = '胶结中到差，以中等为主'
        elif (95 >= ratio_Series['胶结好'] >= 5) & (95 >= ratio_Series['胶结差'] >= 5) & (
                5 >= ratio_Series['胶结中等']):
            if ratio_Series['胶结好'] >= ratio_Series['胶结差']:
                evaluation_of_formation = '胶结好到差，以好为主'
            elif ratio_Series['胶结好'] <= ratio_Series['胶结差']:
                evaluation_of_formation = '胶结好到差，以差为主'
        elif (95 > ratio_Series['胶结好'] > 5) & (95 > ratio_Series['胶结差'] > 5) & (
                95 > ratio_Series['胶结中等'] > 5):
            evaluation_of_formation = '胶结好到中等到差'
        elif (95 > ratio_Series['胶结好'] > 5) & (5 >= ratio_Series['胶结差']) & (5 >= ratio_Series['胶结中等']):
            evaluation_of_formation = '胶结好到中到差，以好为主'
        elif (5 >= ratio_Series['胶结好']) & (5 >= ratio_Series['胶结差']) & (95 > ratio_Series['胶结中等'] > 5):
            evaluation_of_formation = '胶结好到中到差，以中等为主'
        elif (5 >= ratio_Series['胶结好']) & (95 > ratio_Series['胶结差'] > 5) & (5 >= ratio_Series['胶结中等']):
            evaluation_of_formation = '胶结好到中到差，以差为主'
        return ratio_Series, evaluation_of_formation

    ################################################################################
    # 二界面某井段评价函数
    def layer_evaluation2(self, df, start, end):
        df1 = df
        df1 = df1.reset_index()
        formation_Start = start
        formation_End = end
        df_temp = df1.loc[(df1['井段Start'] >= formation_Start) & (df1['井段Start'] <= formation_End), :]
        # 获取起始深度到第一层井段底界的结论
        df_temp_start_to_first_layer = df1.loc[(df1['井段Start'] <= formation_Start), :]
        if len(df_temp_start_to_first_layer) != 0:
            start_to_upper_result = df_temp_start_to_first_layer.loc[len(df_temp_start_to_first_layer) - 1, '结论']
        elif len(df_temp_start_to_first_layer) == 0:
            start_to_upper_result = df1.loc[0, '结论']
        # 获取calculation_Start所在段的声幅值
        df_temp_formation_Start = df1.loc[(df1['井段Start'] <= formation_Start) & (
                df1['井段End'] >= formation_Start), :]
        df_temp_formation_Start.reset_index(drop=True, inplace=True)  # 重新设置列索引#防止若截取中段，index不从0开始的bug
        # 补充储层界到井段的深度
        x, y = df_temp.shape
        df_temp = df_temp.reset_index()
        df_temp.drop(['index'], axis=1, inplace=True)
        if x != 0:  # 防止df_temp为空时，loc报错的bug
            first_layer_start = df_temp.loc[0, '井段Start']
        if x > 0 and first_layer_start != formation_Start:
            upper = pd.DataFrame({'井段(m)': ''.join([str(formation_Start), '-', str(first_layer_start)]),
                                  '厚度(m)': first_layer_start - formation_Start,
                                  '最大指数': df_temp_formation_Start.loc[0, '最大指数'],
                                  '最小指数': df_temp_formation_Start.loc[0, '最小指数'],
                                  '平均指数': df_temp_formation_Start.loc[0, '平均指数'],
                                  '结论': start_to_upper_result,
                                  '井段Start': formation_Start,
                                  '井段End': first_layer_start},
                                 index=[1])  # 自定义索引为：1 ，这里也可以不设置index
            df_temp.loc[len(df_temp) - 1, '井段End'] = formation_End
            df_temp = pd.concat([upper, df_temp], ignore_index=True)
            # df_temp = df_temp.append(new, ignore_index=True)  # ignore_index=True,表示不按原来的索引，从0开始自动递增
            df_temp.loc[:, "重计算厚度"] = df_temp.apply(self.get_thickness, axis=1)
            # 修改df_temp的最末一行
            df_temp.loc[len(df_temp) - 1, '井段(m)'] = ''.join([str(df_temp.loc[len(df_temp) - 1, '井段Start']), \
                                                                '-', str(df_temp.loc[len(df_temp) - 1, '井段End'])])
            df_temp.loc[len(df_temp) - 1, '厚度(m)'] = df_temp.loc[len(df_temp) - 1, '重计算厚度']
        elif x > 0 and first_layer_start == formation_Start:
            df_temp.loc[len(df_temp) - 1, '井段End'] = formation_End
            df_temp.loc[:, "重计算厚度"] = df_temp.apply(self.get_thickness, axis=1)
            # 修改df_temp的最末一行
            df_temp.loc[len(df_temp) - 1, '井段(m)'] = ''.join([str(df_temp.loc[len(df_temp) - 1, '井段Start']), \
                                                                '-', str(df_temp.loc[len(df_temp) - 1, '井段End'])])
            df_temp.loc[len(df_temp) - 1, '厚度(m)'] = df_temp.loc[len(df_temp) - 1, '重计算厚度']
        else:  # 储层包含在一个井段内的情况
            df_temp = pd.DataFrame({'井段(m)': ''.join([str(formation_Start), '-', str(formation_End)]),
                                    '厚度(m)': formation_End - formation_Start,
                                    '最大指数': df_temp_formation_Start.loc[0, '最大指数'],
                                    '最小指数': df_temp_formation_Start.loc[0, '最小指数'],
                                    '平均指数': df_temp_formation_Start.loc[0, '平均指数'],
                                    '结论': start_to_upper_result,
                                    '井段Start': formation_Start,
                                    '井段End': formation_End},
                                   index=[1])  # 自定义索引为：1 ，这里也可以不设置index
            df_temp.loc[:, "重计算厚度"] = df_temp.apply(self.get_thickness, axis=1)
            # 修改df_temp的最末一行
            df_temp.loc[len(df_temp), '井段(m)'] = ''.join([str(df_temp.loc[len(df_temp), '井段Start']),
                                                            '-', str(df_temp.loc[len(df_temp), '井段End'])])
            df_temp.loc[len(df_temp), '厚度(m)'] = df_temp.loc[len(df_temp), '重计算厚度']
        # print(df_temp)
        ratio_Series = df_temp.groupby(by=['结论'])['重计算厚度'].sum() / df_temp['重计算厚度'].sum() * 100
        if '不确定' not in ratio_Series:
            if ratio_Series.__len__() == 2:
                if '胶结好' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                elif '胶结中等' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                elif '胶结差' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
            elif ratio_Series.__len__() == 1:
                if ('胶结好' not in ratio_Series) & ('胶结中等' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                elif ('胶结好' not in ratio_Series) & ('胶结差' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
                elif ('胶结中等' not in ratio_Series) & ('胶结差' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
        else:
            if ratio_Series.__len__() == 1:
                ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
            elif ratio_Series.__len__() == 2:
                if ('胶结好' not in ratio_Series) & ('胶结中等' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                elif ('胶结好' not in ratio_Series) & ('胶结差' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
                elif ('胶结中等' not in ratio_Series) & ('胶结差' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
            elif ratio_Series.__len__() == 3:
                if '胶结好' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                elif '胶结中等' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                elif '胶结差' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))

        # 条件判断，参数需要研究
        if ratio_Series['胶结好'] >= 95:
            evaluation_of_formation = '胶结好'
        elif ratio_Series['胶结中等'] >= 95:
            evaluation_of_formation = '胶结中等'
        elif ratio_Series['胶结差'] >= 95:
            evaluation_of_formation = '胶结差'
        elif (95 >= ratio_Series['胶结好'] >= 5) & (95 >= ratio_Series['胶结中等'] >= 5) & (
                5 >= ratio_Series['胶结差']):
            if ratio_Series['胶结好'] >= ratio_Series['胶结中等']:
                evaluation_of_formation = '胶结中到好，以好为主'
            elif ratio_Series['胶结好'] <= ratio_Series['胶结中等']:
                evaluation_of_formation = '胶结中到好，以中等为主'
        elif (95 >= ratio_Series['胶结差'] >= 5) & (95 >= ratio_Series['胶结中等'] >= 5) & (
                5 >= ratio_Series['胶结好']):
            if ratio_Series['胶结差'] >= ratio_Series['胶结中等']:
                evaluation_of_formation = '胶结中到差，以差为主'
            elif ratio_Series['胶结差'] <= ratio_Series['胶结中等']:
                evaluation_of_formation = '胶结中到差，以中等为主'
        elif (95 >= ratio_Series['胶结好'] >= 5) & (95 >= ratio_Series['胶结差'] >= 5) & (
                5 >= ratio_Series['胶结中等']):
            if ratio_Series['胶结好'] >= ratio_Series['胶结差']:
                evaluation_of_formation = '胶结好到差，以好为主'
            elif ratio_Series['胶结好'] <= ratio_Series['胶结差']:
                evaluation_of_formation = '胶结好到差，以差为主'
        elif (95 > ratio_Series['胶结好'] > 5) & (95 > ratio_Series['胶结差'] > 5) & (
                95 > ratio_Series['胶结中等'] > 5):
            evaluation_of_formation = '胶结好到中等到差'
        elif (95 > ratio_Series['胶结好'] > 5) & (5 >= ratio_Series['胶结差']) & (5 >= ratio_Series['胶结中等']):
            evaluation_of_formation = '胶结好到中到差，以好为主'
        elif (5 >= ratio_Series['胶结好']) & (5 >= ratio_Series['胶结差']) & (95 > ratio_Series['胶结中等'] > 5):
            evaluation_of_formation = '胶结好到中到差，以中等为主'
        elif (5 >= ratio_Series['胶结好']) & (95 > ratio_Series['胶结差'] > 5) & (5 >= ratio_Series['胶结中等']):
            evaluation_of_formation = '胶结好到中到差，以差为主'
        elif ratio_Series['不确定'] >= 1:
            evaluation_of_formation = '不确定'
        return ratio_Series, evaluation_of_formation

    ################################################################################
    # 清理所有非空文件夹和文件
    def clean_dir_of_all(self, path):
        list = os.listdir(path)
        if len(list) != 0:
            for i in range(0, len(list)):
                path_to_clean = os.path.join(path, list[i])
                if '.' not in list[i]:
                    shutil.rmtree(path_to_clean)  # 清理文件夹，可非空
                else:
                    os.remove(path_to_clean)  # 清理文件
        else:
            pass

    # 只清理文件，跳过文件夹
    def clean_dir_just_files(self, path):
        list = os.listdir(path)
        if len(list) != 0:
            for i in range(0, len(list)):
                path_to_clean = os.path.join(path, list[i])
                if '.' not in list[i]:
                    pass
                else:
                    os.remove(path_to_clean)  # 清理文件
        else:
            pass

    def clean_report_workspace(self):
        dir1_path = '.\\WorkSpace\\报告生成工区\\原始资料'
        dir2_path = '.\\WorkSpace\\报告生成工区\\成果表'
        dir4_path = '.\\WorkSpace\\报告生成工区\\储层图'
        dir5_path = '.\\WorkSpace\\报告生成工区\\胶结差图'
        self.clean_dir_of_all(dir1_path)
        self.clean_dir_of_all(dir2_path)
        self.clean_dir_of_all(dir4_path)
        self.clean_dir_of_all(dir5_path)
        QMessageBox.information(self, "提示", "报告生成工区清理完毕\n（除了储层表文件夹）")

    def clean_report_workspace_all(self):
        dir1_path = '.\\WorkSpace\\报告生成工区\\原始资料'
        dir2_path = '.\\WorkSpace\\报告生成工区\\成果表'
        dir3_path = '.\\WorkSpace\\报告生成工区\\储层表'
        dir4_path = '.\\WorkSpace\\报告生成工区\\储层图'
        dir5_path = '.\\WorkSpace\\报告生成工区\\胶结差图'
        self.clean_dir_of_all(dir1_path)
        self.clean_dir_of_all(dir2_path)
        self.clean_dir_of_all(dir3_path)
        self.clean_dir_of_all(dir4_path)
        self.clean_dir_of_all(dir5_path)
        QMessageBox.information(self, "提示", "报告生成工区全部清理完毕")

    def clean_workspace_all(self):
        dir1_path = '.\\WorkSpace\\报告生成工区\\原始资料'
        dir2_path = '.\\WorkSpace\\报告生成工区\\成果表'
        dir3_path = '.\\WorkSpace\\报告生成工区\\储层表'
        dir4_path = '.\\WorkSpace\\报告生成工区\\储层图'
        dir5_path = '.\\WorkSpace\\报告生成工区\\胶结差图'
        dir6_path = '.\\WorkSpace\\分层和成果表工区'
        dir7_path = '.\\WorkSpace\\合并统计工区'
        dir8_path = '.\\WorkSpace'
        self.clean_dir_of_all(dir1_path)
        self.clean_dir_of_all(dir2_path)
        self.clean_dir_of_all(dir3_path)
        self.clean_dir_of_all(dir4_path)
        self.clean_dir_of_all(dir5_path)
        self.clean_dir_of_all(dir6_path)
        self.clean_dir_of_all(dir7_path)
        self.clean_dir_just_files(dir8_path)
        QMessageBox.information(self, "提示", "WorkSpace工区全部清理完毕")

    def open_report_workspace_directory(self):
        path = '.\\WorkSpace\\报告生成工区'
        os.startfile(path)

    def open_result_table_directory(self):
        path = '.\\WorkSpace\\报告生成工区\\成果表'
        os.startfile(path)

    def open_formation_table_directory(self):
        path = '.\\WorkSpace\\报告生成工区\\储层表'
        os.startfile(path)

    def open_formation_pictures_directory(self):
        path = '.\\WorkSpace\\报告生成工区\\储层图'
        os.startfile(path)

    def open_bad_cement_pictures_directory(self):
        path = '.\\WorkSpace\\报告生成工区\\胶结差图'
        os.startfile(path)

    def flush_on_textEdits(self):
        dir1_path = '.\\WorkSpace\\报告生成工区\\原始资料'
        dir2_path = '.\\WorkSpace\\报告生成工区\\成果表'
        dir3_path = '.\\WorkSpace\\报告生成工区\\储层表'
        dir4_path = '.\\WorkSpace\\报告生成工区\\储层图'
        dir5_path = '.\\WorkSpace\\报告生成工区\\胶结差图'
        ###################
        fileNames = []
        if os.listdir(dir2_path) != []:
            for fileName in os.listdir(dir2_path):
                fileNames.append(fileName)
        else:
            fileNames = ['空']

        text = ''
        for item in fileNames:
            text = text + item + '\n'
        self.textEdit_4.setText(text)

        ###################
        fileNames = []
        if os.listdir(dir3_path) != []:
            for fileName in os.listdir(dir3_path):
                fileNames.append(fileName)
        else:
            fileNames = ['空']

        text = ''
        for item in fileNames:
            text = text + item + '\n'
        self.textEdit_5.setText(text)
        ###################
        fileNames = []
        if os.listdir(dir4_path) != []:
            for fileName in os.listdir(dir4_path):
                fileNames.append(fileName)
        else:
            fileNames = ['空']
        # 利用lambda表达式排序
        if fileNames != ['空']:
            fileNames.sort(key=lambda x: int(x.split('#')[0].split('-')[0]))
        text = ''
        for item in fileNames:
            text = text + item + '\n'
        self.textEdit_6.setText(text)
        ###################
        fileNames = []
        if os.listdir(dir5_path) != []:
            for fileName in os.listdir(dir5_path):
                fileNames.append(fileName)
        else:
            fileNames = ['空']
        # 利用lambda表达式排序
        if fileNames != ['空']:
            fileNames.sort(key=lambda x: int(x.split('-')[0]))
        text = ''
        for item in fileNames:
            text = text + item + '\n'
        self.textEdit_8.setText(text)

    def result_table_process_in_report_module(self):
        # 读取单层统计表
        PATH = ".\\WorkSpace\\报告生成工区\\成果表"
        for fileName in os.listdir(PATH):
            if '1单' in fileName and '.xls' in fileName and '$' not in fileName:
                fileDir = PATH + "\\" + fileName
                try:
                    # 表格表头字段规范
                    self.xls_formatting_first_layer(fileDir)
                finally:
                    pass
                QMessageBox.information(self, "提示", "一界面表格数据规范化完毕")
        for fileName in os.listdir(PATH):
            if '2单' in fileName and '.xls' in fileName and '$' not in fileName:
                fileDir = PATH + "\\" + fileName
                try:
                    # 表格表头字段规范
                    self.xls_formatting_second_layer(fileDir)
                finally:
                    pass
                QMessageBox.information(self, "提示", "二界面表格数据规范化完毕")

    def generate_report_thread(self):
        generate_report = threading.Thread(target=self.generate_report)
        generate_report.start()

    def generate_report(self):
        if self.run_on_net == True:
            Supervisor.generate_report_usage_supervisor()
        else:
            pass
        self.lock.acquire()  # 上锁

        # 长宁和威远的井报告模板较特殊，需要提示
        # QMessageBox会导致死机，因此改用setText方式进行提醒
        well_Name = self.lineEdit.text()
        if '长宁' in well_Name or '宁' in well_Name:
            self.label_132.setText('请注意长宁的井报告模板特殊格式')
            self.label_132.setStyleSheet("font: 12pt")
            self.label_132.setStyleSheet("color: rgb(255, 0, 0)")
        elif '威' in well_Name:
            self.label_132.setText('请注意威远的井报告模板特殊格式')
            self.label_132.setStyleSheet("font: 12pt")
            self.label_132.setStyleSheet("color: rgb(255, 0, 0)")
        else:
            pass

        # 重庆气矿报告模板较特殊，需要提示
        client_Name = self.lineEdit_50.text()
        if '重庆' in client_Name:
            self.label_132.setText('请注意重庆气矿报告模板特殊格式')
            self.label_132.setStyleSheet("font: 12pt")
            self.label_132.setStyleSheet("color: rgb(255, 0, 0)")
        else:
            pass

        # 提取成果表中的内容
        PATH = ".\\WorkSpace\\报告生成工区\\成果表"
        for fileName in os.listdir(PATH):
            if '1统' in fileName:
                fileDir = PATH + "\\" + fileName
                workbook1 = xlrd.open_workbook(fileDir)
            elif '2统' in fileName:
                fileDir = PATH + "\\" + fileName
                workbook2 = xlrd.open_workbook(fileDir)

        ##########################
        # 解析解释成果表-1统
        sheet1 = workbook1.sheets()[0]

        nrow1 = sheet1.nrows
        ncol1 = sheet1.ncols

        # 统计结论
        first_GLength = str(sheet1.cell_value(3, 2))
        first_GRatio = str(sheet1.cell_value(3, 3))

        first_MLength = str(sheet1.cell_value(4, 2))
        first_MRatio = str(sheet1.cell_value(4, 3))

        first_BLength = str(sheet1.cell_value(5, 2))
        first_BRatio = str(sheet1.cell_value(5, 3))

        # 合格率
        first_Pass_Percent = str(round((float(sheet1.cell_value(3, 3)) + float(sheet1.cell_value(4, 3))), 2))

        ##########################
        # 解析解释成果表-2统
        sheet2 = workbook2.sheets()[0]

        nrow2 = sheet2.nrows
        ncol2 = sheet2.ncols

        # 统计结论
        second_GLength = str(sheet2.cell_value(3, 2))
        second_GRatio = str(sheet2.cell_value(3, 3))

        second_MLength = str(sheet2.cell_value(4, 2))
        second_MRatio = str(sheet2.cell_value(4, 3))

        second_BLength = str(sheet2.cell_value(5, 2))
        second_BRatio = str(sheet2.cell_value(5, 3))

        # 合格率
        second_Pass_Percent = str(round((float(sheet2.cell_value(3, 3)) + float(sheet2.cell_value(4, 3))), 2))
        print('统计表解析完成')

        # 整体评价
        if eval(first_Pass_Percent) >= 70:
            first_Eval_Result = '合格'
        else:
            first_Eval_Result = '不合格'

        if eval(second_Pass_Percent) >= 70:
            second_Eval_Result = '合格'
        else:
            second_Eval_Result = '不合格'
        ################################################################################
        # 读取单层统计表
        PATH = ".\\WorkSpace\\报告生成工区\\成果表"
        for fileName in os.listdir(PATH):
            if '1单' in fileName and '.xls' in fileName and '$' not in fileName:
                fileDir = PATH + "\\" + fileName
                workbook = xlrd.open_workbook(fileDir)

        sheet = workbook.sheets()[0]

        # 获得表单的行数及列数
        nrow = sheet.nrows
        ncol = sheet.ncols
        # 处理评价井段
        start_Evaluation = str(sheet.cell_value(3, 1)).strip()
        start_Evaluation = start_Evaluation.split('-')[0]
        end_Evaluation = str(sheet.cell_value(nrow - 1, 1)).strip('')
        end_Evaluation = ''.join(end_Evaluation.split())  # 去除所有空格
        end_Evaluation = end_Evaluation.split('-')[1]
        first_Pro_Interval = ''.join([start_Evaluation, '-', end_Evaluation])
        # 新规定，生成一个从0开始的处理深度， 为了多计费
        # TODO
        if float(start_Evaluation) < 200:
            measure_start = '0'
            measure_from_Pro = ''.join([measure_start, '-', end_Evaluation])
        else:
            measure_from_Pro = ''.join([start_Evaluation, '-', end_Evaluation])

        PATH = ".\\WorkSpace\\报告生成工区\\成果表"
        for fileName in os.listdir(PATH):
            if '2单' in fileName and '.xls' in fileName and '$' not in fileName:
                fileDir = PATH + "\\" + fileName
                workbook = xlrd.open_workbook(fileDir)

        sheet = workbook.sheets()[0]

        # 获得表单的行数及列数
        nrow = sheet.nrows
        ncol = sheet.ncols
        # 处理评价井段
        start_Evaluation = str(sheet.cell_value(3, 1)).strip()
        start_Evaluation = start_Evaluation.split('-')[0]
        end_Evaluation = str(sheet.cell_value(nrow - 1, 1)).strip('')
        end_Evaluation = ''.join(end_Evaluation.split())  # 去除所有空格
        end_Evaluation = end_Evaluation.split('-')[1]
        second_Pro_Interval = ''.join([start_Evaluation, '-', end_Evaluation])

        # 液面高度的获取
        fluid_Height = start_Evaluation

        ################################################################################
        # 判断是否有储层
        PATH = ".\\WorkSpace\\报告生成工区\\储层表\\"
        if os.listdir(PATH) != []:
            for fileName in os.listdir(PATH):
                fileName = fileName
        else:
            fileName = ''

        f_path = PATH + fileName

        if os.path.isdir(f_path):
            formation_be_or_not = '无储层'
        else:
            formation_be_or_not = '有储层'
        ################################################################################
        # 储层表解析
        if formation_be_or_not == '有储层':
            PATH = ".\\WorkSpace\\报告生成工区\\储层表"
            for fileName in os.listdir(PATH):
                fileDir = PATH + "\\" + fileName
                workbook = xlrd.open_workbook(fileDir)

            sheet = workbook.sheets()[0]

            # 通过xlrd的接口获得表单的行数及列数
            nrow = sheet.nrows
            ncol = sheet.ncols

            if nrow >= 3:
                formation_Number = str(nrow - 2)
            else:
                formation_Number = '[待确定]'
            print('储层表解析完成')
        else:
            print('未发现储层表')
        ################################################################################
        # 储层表和单层统计表的联动数据分析
        all_evaluation_of_formation1 = []
        all_evaluation_of_formation2 = []
        if formation_be_or_not == '有储层':
            # 单层统计表
            PATH = ".\\WorkSpace\\报告生成工区\\成果表"
            for fileName in os.listdir(PATH):
                if '1单' in fileName and '$' not in fileName:
                    fileDir1 = PATH + "\\" + fileName
                elif '2单' in fileName and '$' not in fileName:
                    fileDir2 = PATH + "\\" + fileName
            df1 = pd.read_excel(fileDir1, header=2)
            # df1.drop([0], inplace=True)
            df1.loc[:, '井段(m)'] = df1['井段(m)'].str.replace(' ', '')  # 消除数据中空格
            df1['井段Start'] = df1['井段(m)'].map(lambda x: x.split("-")[0])
            df1['井段End'] = df1['井段(m)'].map(lambda x: x.split("-")[1])

            df2 = pd.read_excel(fileDir2, header=2)
            # df2.drop([0], inplace=True)
            df2.loc[:, '井段(m)'] = df2['井段(m)'].str.replace(' ', '')  # 消除数据中空格
            df2['井段Start'] = df2['井段(m)'].map(lambda x: x.split("-")[0])
            df2['井段End'] = df2['井段(m)'].map(lambda x: x.split("-")[1])

            # 储层表
            PATH = ".\\WorkSpace\\报告生成工区\\储层表"
            for fileName in os.listdir(PATH):
                fileDir3 = PATH + "\\" + fileName
            df3 = pd.read_excel(fileDir3, header=0)
            df3.drop([0], inplace=True)
            df3.drop(['层位', '解释结论'], axis=1, inplace=True)
            df3.loc[:, '井        段'] = df3['井        段'].str.replace(' ', '')  # 消除数据中空格
            df3['储层Start'] = df3['井        段'].map(lambda x: x.split("--")[0])
            df3['储层End'] = df3['井        段'].map(lambda x: x.split("--")[1])

            # 表格数据清洗
            df1.loc[:, "井段Start"] = df1["井段Start"].str.replace(" ", "").astype('float')
            df1.loc[:, "井段End"] = df1["井段End"].str.replace(" ", "").astype('float')
            df2.loc[:, "井段Start"] = df2["井段Start"].str.replace(" ", "").astype('float')
            df2.loc[:, "井段End"] = df2["井段End"].str.replace(" ", "").astype('float')
            df3.loc[:, "储层Start"] = df3["储层Start"].str.replace(" ", "").astype('float')
            df3.loc[:, "储层End"] = df3["储层End"].str.replace(" ", "").astype('float')
            rows1, cols1 = df1.shape
            rows2, cols2 = df2.shape
            rows3, cols3 = df3.shape

            # 针对每个储层在单层评价表中得出好中差比例
            for row in range(1, rows3 + 1):
                formation_Start = df3.loc[row, '储层Start']
                formation_End = df3.loc[row, '储层End']
                # print('----------------第', row, '个储层内的井段----------------')
                if (formation_End <= float(end_Evaluation)) & (formation_Start >= float(start_Evaluation)):
                    ratio_Series1 = self.layer_evaluation1(df1, formation_Start, formation_End)[0]  # 调取一界面评价函数
                    evaluation_of_formation1 = self.layer_evaluation1(df1, formation_Start, formation_End)[
                        1]  # 调取一界面评价函数
                    all_evaluation_of_formation1.append(evaluation_of_formation1)

                    ratio_Series2 = self.layer_evaluation2(df2, formation_Start, formation_End)[0]  # 调取二界面评价函数
                    evaluation_of_formation2 = self.layer_evaluation2(df2, formation_Start, formation_End)[
                        1]  # 调取二界面评价函数
                    all_evaluation_of_formation2.append(evaluation_of_formation2)

                elif (formation_End > float(end_Evaluation)) & (formation_Start < float(end_Evaluation)) & (
                        formation_Start >= float(start_Evaluation)):
                    ratio_Series1 = self.layer_evaluation1(df1, formation_Start, float(end_Evaluation))[0]  # 调取一界面评价函数
                    evaluation_of_formation1 = self.layer_evaluation1(df1, formation_Start, float(end_Evaluation))[
                        1]  # 调取一界面评价函数
                    ratio_Series2 = self.layer_evaluation2(df2, formation_Start, float(end_Evaluation))[0]  # 调取二界面评价函数
                    evaluation_of_formation2 = self.layer_evaluation2(df2, formation_Start, float(end_Evaluation))[
                        1]  # 调取二界面评价函数
                else:
                    print('储层界超出了测量范围，请检查')
                    pass

        ################################################################################
        # 基于文本替换方案的文档生成(和生成LEAD txt文件不同的是，报告水泥返高和图件水泥返高不一样)
        well_Name = self.lineEdit.text()
        well_Category = self.comboBox_7.currentText()
        well_Type = self.comboBox_6.currentText()
        well_Depth = self.lineEdit_4.text()
        x_Coordinate = self.lineEdit_9.text()
        y_Coordinate = self.lineEdit_10.text()
        magnetic_Declination = self.lineEdit_14.text()
        ground_Elevation = self.lineEdit_12.text()
        bushing_Height = self.lineEdit_13.text()
        kelly_Bushing = self.lineEdit_11.text()
        oil_Field = self.lineEdit_17.text()
        client_Name = self.lineEdit_50.text()
        geo_Position = self.lineEdit_16.text()
        stru_Position = self.lineEdit_18.text()
        completion_Date = self.lineEdit_39.text()
        spud_Date = self.lineEdit_20.text()
        end_Drilling_Date = self.lineEdit_21.text()
        deepest_bit = self.lineEdit_102.text()
        drilling_Unit = self.lineEdit_74.text()
        bit1_Diameter = self.tableWidget_6.item(0, 0).text()
        bit1_Depth = self.tableWidget_6.item(0, 1).text()
        bit2_Diameter = self.tableWidget_6.item(1, 0).text()
        bit2_Depth = self.tableWidget_6.item(1, 1).text()
        bit3_Diameter = self.tableWidget_6.item(2, 0).text()
        bit3_Depth = self.tableWidget_6.item(2, 1).text()
        bit4_Diameter = self.tableWidget_6.item(3, 0).text()
        bit4_Depth = self.tableWidget_6.item(3, 1).text()
        bit5_Diameter = self.tableWidget_6.item(4, 0).text()
        bit5_Depth = self.tableWidget_6.item(4, 1).text()
        well_Times_Name = self.lineEdit_3.text()
        well_Times_Type = self.lineEdit_49.text()
        logging_Date = self.lineEdit_40.text()
        measure_Interval_Start_Depth = self.lineEdit_110.text()
        measure_Interval_End_Depth = self.lineEdit_111.text()
        logging_Equipment = self.comboBox_5.currentText()
        arti_Bottom = self.lineEdit_65.text()
        interpretation_Complete_Date = self.lineEdit_55.text()
        evaluation_start_depth = self.lineEdit_103.text()
        evaluation_end_depth = self.lineEdit_105.text()
        task_Number = self.lineEdit_19.text()
        logging_Group = self.lineEdit_43.text()
        logging_Leader = self.lineEdit_54.text()
        logging_Operator = self.lineEdit_44.text()
        report_Writer = self.comboBox_2.currentText()
        report_Checker = self.comboBox.currentText()
        report_Supervisor = self.comboBox_4.currentText()
        flu_Property = self.lineEdit_60.text()
        flu_Density = self.lineEdit_61.text()
        flu_Viscosity = self.lineEdit_62.text()
        design_Depth = self.lineEdit_69.text()
        actual_Depth = self.lineEdit_107.text()
        cement_Quantity = self.lineEdit_71.text()
        cement_Density = self.lineEdit_72.text()
        cement_End_Date = self.lineEdit_22.text()
        cement_Unit = self.lineEdit_25.text()

        casing1_Inner_Dia = self.tableWidget_7.item(0, 0).text()
        casing1_Dia = self.tableWidget_7.item(0, 1).text()
        casing1_Thickness = self.tableWidget_7.item(0, 2).text()
        casing1_bottom = self.tableWidget_7.item(0, 3).text()

        casing2_Inner_Dia = self.tableWidget_7.item(1, 0).text()
        casing2_Dia = self.tableWidget_7.item(1, 1).text()
        casing2_Thickness = self.tableWidget_7.item(1, 2).text()
        casing2_bottom = self.tableWidget_7.item(1, 3).text()

        casing3_Inner_Dia = self.tableWidget_7.item(2, 0).text()
        casing3_Dia = self.tableWidget_7.item(2, 1).text()
        casing3_Thickness = self.tableWidget_7.item(2, 2).text()
        casing3_bottom = self.tableWidget_7.item(2, 3).text()

        casing4_Inner_Dia = self.tableWidget_7.item(3, 0).text()
        casing4_Dia = self.tableWidget_7.item(3, 1).text()
        casing4_Thickness = self.tableWidget_7.item(3, 2).text()
        casing4_bottom = self.tableWidget_7.item(3, 3).text()

        casing5_Inner_Dia = self.tableWidget_7.item(4, 0).text()
        casing5_Dia = self.tableWidget_7.item(4, 1).text()
        casing5_Thickness = self.tableWidget_7.item(4, 2).text()
        casing5_bottom = self.tableWidget_7.item(4, 3).text()

        cement_End_Time = self.lineEdit_101.text()
        cement_End_Date = self.lineEdit_22.text()
        logging_Start_Time = self.lineEdit_104.text()
        logging_Method = self.lineEdit_100.text()

        # 补充信息
        casing_Goal = self.lineEdit_108.text()
        measure_Interval = ''.join([measure_Interval_Start_Depth, '-', measure_Interval_End_Depth])
        log_End_Time = logging_Date
        try:
            year = logging_Date.split('-')[0]
            month = logging_Date.split('-')[1]
            day = logging_Date.split('-')[2]
        except:
            year = ''
            month = ''
            day = ''
        dev_Depth_Ratio = self.lineEdit_26.text()
        casing1_interval = self.tableWidget_7.item(0, 4).text()
        casing2_interval = self.tableWidget_7.item(1, 4).text()
        casing3_interval = self.tableWidget_7.item(2, 4).text()
        casing4_interval = self.tableWidget_7.item(3, 4).text()
        casing5_interval = self.tableWidget_7.item(4, 4).text()

        # 报告中的未知字段补充
        second_Start = self.lineEdit_106.text()
        evaluation_Start = self.lineEdit_103.text()
        evaluation_Bottom = self.lineEdit_105.text()
        liner_overlap_section = self.comboBox_8.currentText()
        upper_casing_shoe = self.comboBox_9.currentText()
        goal_layer_name = self.lineEdit_118.text()
        goal_layer_first = self.comboBox_11.currentText()
        goal_layer_second = self.comboBox_12.currentText()
        encounter_obstacle_depth = self.lineEdit_115.text()
        max_Well_Deviation_Depth = self.lineEdit_24.text()
        max_Well_Deviation = self.lineEdit_28.text()

        TEMPLATE_PATH = ".\\resources\\模板"
        PATH = "."

        if cement_Quantity == '-99999':
            cement_Quantity = '/'

        try:
            if float(design_Depth) <= 200:
                design_Depth = '地面'
        except:
            pass

        try:
            if float(actual_Depth) <= 200:
                actual_Depth = '地面'
        except:
            pass

        ##############################################################################
        # 定性描述补充

        if self.lineEdit_119.text() != '':
            start_depth = ['', '', '', '']
            end_depth = ['', '', '', '']

            start_depth[0] = self.lineEdit_103.text()
            end_depth[0] = self.lineEdit_106.text()

            start_depth[1] = self.lineEdit_106.text()
            end_depth[1] = self.lineEdit_120.text()

            start_depth[2] = self.lineEdit_120.text()
            end_depth[2] = self.lineEdit_119.text()

            start_depth[3] = self.lineEdit_119.text()
            end_depth[3] = self.lineEdit_105.text()

            PATH = ".\\WorkSpace\\报告生成工区\\成果表"
            # 1单
            for fileName in os.listdir(PATH):
                if '1单' in fileName and '.xls' in fileName and '$' not in fileName:
                    fileDir1 = PATH + "\\" + fileName
            # 2单
            for fileName in os.listdir(PATH):
                if '2单' in fileName and '.xls' in fileName and '$' not in fileName:
                    fileDir2 = PATH + "\\" + fileName

            first_layer_result = [0, 0, 0, 0]
            second_layer_result = [0, 0, 0, 0]

            for i in range(4):
                # 获取一界面单层评价表的深度界限
                df1 = pd.read_excel(fileDir1, header=2)
                # df1.drop([0], inplace=True)
                df1 = df1.dropna(axis=0, how='any')  # 删除dataframe里NaN的所有行
                df1.loc[:, '井段(m)'] = df1['井段(m)'].str.replace(' ', '')  # 消除数据中空格
                # if len(df1) % 2 == 0:#如果len(df1)为偶数需要删除最后一行NaN，一行的情况不用删
                #     df1.drop([len(df1)], inplace=True)
                df1['井段Start'] = df1['井段(m)'].map(lambda x: x.split("-")[0])
                df1['井段End'] = df1['井段(m)'].map(lambda x: x.split("-")[1])
                # 表格数据清洗
                df1.loc[:, "井段Start"] = df1["井段Start"].str.replace(" ", "").astype('float')
                df1.loc[:, "井段End"] = df1["井段End"].str.replace(" ", "").astype('float')
                evaluation_of_formation1 = self.layer_evaluation1(df1, float(start_depth[i]), float(end_depth[i]))[
                    0]  # 调取一界面评价函数
                first_layer_result[0] = evaluation_of_formation1['胶结好']
                first_layer_result[1] = evaluation_of_formation1['胶结中等']
                first_layer_result[2] = evaluation_of_formation1['胶结差']
                not_sure = 100 - evaluation_of_formation1['胶结好'] - evaluation_of_formation1['胶结中等'] - \
                           evaluation_of_formation1['胶结差']
                first_layer_result[3] = not_sure

                # 获取二界面单层评价表的深度界限
                df1 = pd.read_excel(fileDir2, header=2)
                # df1.drop([0], inplace=True)
                df1 = df1.dropna(axis=0, how='any')  # 删除dataframe里NaN的所有行
                df1.loc[:, '井段(m)'] = df1['井段(m)'].str.replace(' ', '')  # 消除数据中空格
                # if len(df1) % 2 == 0:#如果len(df1)为偶数需要删除最后一行NaN，一行的情况不用删
                #     df1.drop([len(df1)], inplace=True)
                df1['井段Start'] = df1['井段(m)'].map(lambda x: x.split("-")[0])
                df1['井段End'] = df1['井段(m)'].map(lambda x: x.split("-")[1])
                # 表格数据清洗
                df1.loc[:, "井段Start"] = df1["井段Start"].str.replace(" ", "").astype('float')
                df1.loc[:, "井段End"] = df1["井段End"].str.replace(" ", "").astype('float')
                evaluation_of_formation2 = self.layer_evaluation2(df1, float(start_depth[i]), float(end_depth[i]))[
                    0]  # 调取二界面评价函数
                second_layer_result[0] = evaluation_of_formation2['胶结好']
                second_layer_result[1] = evaluation_of_formation2['胶结中等']
                second_layer_result[2] = evaluation_of_formation2['胶结差']
                not_sure = 100 - evaluation_of_formation2['胶结好'] - evaluation_of_formation2['胶结中等'] - \
                           evaluation_of_formation2['胶结差']
                second_layer_result[3] = not_sure

                ######################################################## 一界面描述
                if first_layer_result[0] >= 95:
                    evaluation_of_formation1 = '胶结好'
                    cbl_amplitude = '低'
                elif first_layer_result[1] >= 95:
                    evaluation_of_formation1 = '胶结中等'
                    cbl_amplitude = '胶结中等'
                elif first_layer_result[2] >= 95:
                    evaluation_of_formation1 = '胶结差'
                    cbl_amplitude = '高'
                elif (95 >= first_layer_result[0] >= 5) & (95 >= first_layer_result[1] >= 5) & (
                        5 >= first_layer_result[2]):
                    if first_layer_result[0] >= first_layer_result[1]:
                        evaluation_of_formation1 = '胶结中到好，以好为主'
                        cbl_amplitude = '中到低，以低为主'
                    elif first_layer_result[0] <= first_layer_result[1]:
                        evaluation_of_formation1 = '胶结中到好，以中等为主'
                        cbl_amplitude = '中到低，以低为主'
                elif (95 >= first_layer_result[2] >= 5) & (95 >= first_layer_result[1] >= 5) & (
                        5 >= first_layer_result[0]):
                    if first_layer_result[2] >= first_layer_result[1]:
                        evaluation_of_formation1 = '胶结中到差，以差为主'
                        cbl_amplitude = '中到高，以高为主'
                    elif first_layer_result[2] <= first_layer_result[1]:
                        evaluation_of_formation1 = '胶结中到差，以中等为主'
                        cbl_amplitude = '中到高，以中为主'
                elif (95 >= first_layer_result[0] >= 5) & (95 >= first_layer_result[2] >= 5) & (
                        5 >= first_layer_result[1]):
                    if first_layer_result[0] >= first_layer_result[2]:
                        evaluation_of_formation1 = '胶结好到差，以好为主'
                        cbl_amplitude = '低到高，以低为主'
                    elif first_layer_result[0] <= first_layer_result[2]:
                        evaluation_of_formation1 = '胶结好到差，以差为主'
                        cbl_amplitude = '低到高，以高为主'
                elif (95 > first_layer_result[0] > 5) & (95 > first_layer_result[2] > 5) & (
                        95 > first_layer_result[1] > 5):
                    evaluation_of_formation1 = '胶结好到中等到差'
                    cbl_amplitude = '低到中到高'
                elif (95 > first_layer_result[0] > 5) & (5 >= first_layer_result[2]) & (5 >= first_layer_result[1]):
                    evaluation_of_formation1 = '胶结好到中到差，以好为主'
                    cbl_amplitude = '低到中到高，以低为主'
                elif (5 >= first_layer_result[0]) & (5 >= first_layer_result[2]) & (95 > first_layer_result[1] > 5):
                    evaluation_of_formation1 = '胶结好到中到差，以中等为主'
                    cbl_amplitude = '低到中到高，以中为主'
                elif (5 >= first_layer_result[0]) & (95 > first_layer_result[2] > 5) & (5 >= first_layer_result[1]):
                    evaluation_of_formation1 = '胶结好到中到差，以差为主'
                    cbl_amplitude = '低到中到高，以高为主'

                ######################################################## 二界面描述
                if second_layer_result[3] >= 95:
                    evaluation_of_formation2 = '不确定'
                    formation_wave_amplitude = '弱'
                elif second_layer_result[0] >= 95:
                    evaluation_of_formation2 = '胶结好'
                    formation_wave_amplitude = '强'
                elif second_layer_result[1] >= 95:
                    evaluation_of_formation2 = '胶结中等'
                    formation_wave_amplitude = '胶结中等'
                elif second_layer_result[2] >= 95:
                    evaluation_of_formation2 = '胶结差'
                    formation_wave_amplitude = '弱'
                elif (95 >= second_layer_result[0] >= 5) & (95 >= second_layer_result[1] >= 5) & (
                        5 >= second_layer_result[2]):
                    if second_layer_result[0] >= second_layer_result[1]:
                        evaluation_of_formation2 = '胶结中到好，以好为主'
                        formation_wave_amplitude = '中到强，以强为主'
                    elif second_layer_result[0] <= second_layer_result[1]:
                        evaluation_of_formation2 = '胶结中到好，以中等为主'
                        formation_wave_amplitude = '中到强，以中等为主'
                elif (95 >= second_layer_result[2] >= 5) & (95 >= second_layer_result[1] >= 5) & (
                        5 >= second_layer_result[0]):
                    if second_layer_result[2] >= second_layer_result[1]:
                        evaluation_of_formation2 = '胶结中到差，以差为主'
                        formation_wave_amplitude = '中到弱，以弱为主'
                    elif second_layer_result[2] <= second_layer_result[1]:
                        evaluation_of_formation2 = '胶结中到差，以中等为主'
                        formation_wave_amplitude = '中到弱，以中等为主'
                elif (95 >= second_layer_result[0] >= 5) & (95 >= second_layer_result[2] >= 5) & (
                        5 >= second_layer_result[1]):
                    if second_layer_result[0] >= second_layer_result[2]:
                        evaluation_of_formation2 = '胶结好到差，以好为主'
                        formation_wave_amplitude = '强到弱，以强为主'
                    elif second_layer_result[0] <= second_layer_result[2]:
                        evaluation_of_formation2 = '胶结好到差，以差为主'
                        formation_wave_amplitude = '强到弱，以弱为主'
                elif (95 > second_layer_result[0] > 5) & (95 > second_layer_result[2] > 5) & (
                        95 > second_layer_result[1] > 5):
                    evaluation_of_formation2 = '胶结好到中等到差'
                    formation_wave_amplitude = '强到中到弱'
                elif (95 > second_layer_result[0] > 5) & (5 >= second_layer_result[2]) & (5 >= second_layer_result[1]):
                    evaluation_of_formation2 = '胶结好到中到差，以好为主'
                    formation_wave_amplitude = '强到中到弱，以强为主'
                elif (5 >= second_layer_result[0]) & (5 >= second_layer_result[2]) & (95 > second_layer_result[1] > 5):
                    evaluation_of_formation2 = '胶结好到中到差，以中等为主'
                    formation_wave_amplitude = '强到中到弱，以中等为主'
                elif (5 >= second_layer_result[0]) & (95 > second_layer_result[2] > 5) & (5 >= second_layer_result[1]):
                    evaluation_of_formation2 = '胶结好到中到差，以差为主'
                    formation_wave_amplitude = '强到中到弱，以弱为主'

                if i == 0:
                    Description_first = start_depth[i] + '-' + end_depth[
                        i] + 'm井段，声幅幅度' + cbl_amplitude + '，套管首波和接箍信号中；因双层套管，故地层波信号弱，第一界面固井质量' + evaluation_of_formation1 + '，第二界面固井质量不确定。该段固井质量综合评价为' + '?' + '。'
                elif i == 1:
                    Description_second = start_depth[i] + '-' + end_depth[
                        i] + 'm井段，声幅幅度' + cbl_amplitude + '，套管首波和接箍信号中；地层波信号' + formation_wave_amplitude + '，第一界面固井质量' + evaluation_of_formation1 + '，第二界面固井质量' + evaluation_of_formation2 + '。该段固井质量综合评价为' + '?' + '。'
                elif i == 2:
                    Description_third = start_depth[i] + '-' + end_depth[
                        i] + 'm井段，声幅幅度' + cbl_amplitude + '，套管首波和接箍信号中；地层波信号' + formation_wave_amplitude + '，第一界面固井质量' + evaluation_of_formation1 + '，第二界面固井质量' + evaluation_of_formation2 + '。该段固井质量综合评价为' + '?' + '。'
                elif i == 3:
                    Description_fourth = start_depth[i] + '-' + end_depth[
                        i] + 'm井段，声幅幅度' + cbl_amplitude + '，套管首波和接箍信号中；地层波信号' + formation_wave_amplitude + '，第一界面固井质量' + evaluation_of_formation1 + '，第二界面固井质量' + evaluation_of_formation2 + '。该段固井质量综合评价为' + '?' + '。'

        else:
            Description_first = 'xxxx-xxxxm井段，声幅幅度中到高，套管首波和接箍信号中；因双层套管，故地层波信号弱，第一界面固井质量中到差，局部为好，第二界面固井质量好到中到差，以差为主。该段固井质量综合评价为不合格。'
            Description_second = 'xxxx-xxxxm井段，声幅幅度为中到低，套管首波和接箍信号中到低；地层波信号中，第一界面固井质量中到差，以中为主，局部为好，第二界面固井质量好到中到差，以好为主。该段固井质量综合评价为合格。'
            Description_third = 'xxxx-xxxxm井段，声幅幅度中到高，套管首波和接箍信号中到高；第一界面固井质量以好为主、局部为中，第二界面固井质量好到中到差，以好为主。该段固井质量综合评价为合格。'
            Description_fourth = 'xxxx-xxxxm井段，声幅幅度中到高，套管首波和接箍信号中到高；第一界面固井质量以好为主、局部为中，第二界面固井质量好到中到差，以好为主。该段固井质量综合评价为合格。'

        self.DICT = {
            "well_Name": well_Name,
            "stru_Position": stru_Position,
            "casing_Goal": casing_Goal,
            "start_Evaluation": start_Evaluation,
            "end_Evaluation": end_Evaluation,
            "measure_Interval": measure_Interval,
            "first_Pro_Interval": first_Pro_Interval,
            "measure_from_Pro": measure_from_Pro,
            "second_Pro_Interval": second_Pro_Interval,
            "geo_Position": geo_Position,
            "deepest_bit": deepest_bit,
            "arti_Bottom": arti_Bottom,
            "max_Well_Deviation": max_Well_Deviation,
            "max_Well_Deviation_Depth": max_Well_Deviation_Depth,
            "casing1_Dia": casing1_Dia,
            "casing2_Dia": casing2_Dia,
            "casing3_Dia": casing3_Dia,
            "casing4_Dia": casing4_Dia,
            "casing5_Dia": casing5_Dia,
            "bit1_Diameter": bit1_Diameter,
            "bit2_Diameter": bit2_Diameter,
            "bit3_Diameter": bit3_Diameter,
            "bit4_Diameter": bit4_Diameter,
            "bit5_Diameter": bit5_Diameter,
            "flu_Property": flu_Property,
            "flu_Density": flu_Density,
            "flu_Viscosity": flu_Viscosity,
            "cement_Density": cement_Density,
            "cement_Quantity": cement_Quantity,
            "design_Depth": design_Depth,
            "cement_End_Time": cement_End_Time,
            "cement_End_Date": cement_End_Date,
            "log_End_Time": log_End_Time,
            "logging_Group": logging_Group,
            "logging_Leader": logging_Leader,
            "logging_Equipment": logging_Equipment,
            "actual_Depth": actual_Depth,
            "first_GRatio": first_GRatio,
            "first_MRatio": first_MRatio,
            "first_BRatio": first_BRatio,
            "first_Pass_Percent": first_Pass_Percent,
            "first_Eval_Result": first_Eval_Result,
            "first_GLength": first_GLength,
            "first_MLength": first_MLength,
            "first_BLength": first_BLength,
            "second_GRatio": second_GRatio,
            "second_MRatio": second_MRatio,
            "second_BRatio": second_BRatio,
            "second_Pass_Percent": second_Pass_Percent,
            "second_Eval_Result": second_Eval_Result,
            "second_GLength": second_GLength,
            "second_MLength": second_MLength,
            "second_BLength": second_BLength,
            "year": year,
            "month": month,
            "day": day,
            "casing1_interval": casing1_interval,
            "casing2_interval": casing2_interval,
            "casing3_interval": casing3_interval,
            "casing4_interval": casing4_interval,
            "casing5_interval": casing5_interval,
            "bit1_Depth": bit1_Depth,
            "bit2_Depth": bit2_Depth,
            "bit3_Depth": bit3_Depth,
            "bit4_Depth": bit4_Depth,
            "bit5_Depth": bit5_Depth,
            "fluid_Height": fluid_Height,
            "liner_overlap_section": liner_overlap_section,
            "upper_casing_shoe": upper_casing_shoe,
            "goal_layer_name": goal_layer_name,
            "goal_layer_first": goal_layer_first,
            "goal_layer_second": goal_layer_second,
            "second_Start": second_Start,
            "evaluation_Bottom": evaluation_Bottom,
            "evaluation_Start": evaluation_Start,
            "Description_first": Description_first,
            "Description_second": Description_second,
            "Description_third": Description_third,
            "Description_fourth": Description_fourth,
            "encounter_obstacle_depth": encounter_obstacle_depth
        }

        print('模板替换开始，请等待……')

        self.document_replace()  # 模板替换主程序

        print('\n模板替换完成')
        print('储层表添加中，请等待……')

        ################################################################################
        # 储层表的嵌入
        PATH = ".\\WorkSpace\\报告生成工区\\"
        for fileName in os.listdir(PATH):
            # newFile = PATH + well_Name + '_' + year + month + \
            #           day + '_(' + casing_Goal + 'mm套,VDL_' + self.measure_from_Pro + 'm)固井报告' + '.docx'
            newFile = PATH + well_Name + '_固井质量测井评价报告_' + year + month + \
                      day + '_' + casing_Goal + 'mm套' + '.docx'
        document = Document(newFile)

        if formation_be_or_not == '有储层':
            PATH = ".\\WorkSpace\\报告生成工区\\储层表"
            for fileName in os.listdir(PATH):
                fileDir = PATH + "\\" + fileName
                workbook = xlrd.open_workbook(fileDir)

            sheet = workbook.sheets()[0]

            # 获得表单的行数及列数
            nrow = sheet.nrows
            ncol = sheet.ncols

            formation_table = document.tables[9]
            for num in range(eval(formation_Number) - 1):
                row_cells = formation_table.add_row()

            for row in range(1, len(formation_table.rows)):
                print('已添加第', str(row), '个储层数据到报告')
                for col in range(len(formation_table.columns)):
                    formation_table.cell(row, col).text = str(sheet.cell_value(row + 1, col)).strip()
                    formation_table.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    # print(formation_table.cell(row, col).text)
                    formation_table.cell(row, col).paragraphs[
                        0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # 首列居中
            for row in range(len(formation_table.rows)):
                formation_table.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            print('储层表写入完成')
            document.save(newFile)  # 保存下

        ################################################################################
        print('单层统计表添加中……')
        PATH = ".\\WorkSpace\\报告生成工区\\成果表"
        # 添加1单
        for fileName in os.listdir(PATH):
            if '1单' in fileName and '.xls' in fileName and '$' not in fileName:
                fileDir = PATH + "\\" + fileName
                workbook = xlrd.open_workbook(fileDir)
        sheet = workbook.sheets()[0]

        # 通过xlrd的接口获得表单的行数及列数
        nrow = sheet.nrows
        ncol = sheet.ncols

        # 行数过多提醒
        # QMessageBox会导致崩溃，改为label提示的方式
        if nrow > 280:
            # QMessageBox.information(self, '提示', '单层评价表行数较多，表格调整耗时较长，建议取消勾选，自行手动调整:)')
            self.label_133.setText('单层评价表行数较多，格式优化耗时较长，可取消勾选，自行手动调整:)')
            self.label_133.setStyleSheet("font: 16pt")
            self.label_133.setStyleSheet("color: rgb(255, 0, 0)")
        else:
            pass

        document = Document(newFile)
        document.styles['Normal'].font.size = Pt(11)
        document.styles['Normal'].font.name = u'Times New Roman'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        table = document.tables[12]
        table.autofit = True
        for num in range(nrow - 3):
            table.add_row()

        # 设置整个表格字体属性
        table.style.font.color.rgb = RGBColor(0, 0, 0)
        table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table.cell(0, 0).width = Pt(30)
        table.cell(0, 1).width = Pt(100)

        # 单层评价表写入
        if self.checkBox_10.isChecked():
            for row in range(1, len(table.rows)):
                self.view_bar(row, len(table.rows) - 1)
                print(' @第', str(row), '行')
                for col in range(len(table.columns)):
                    table.cell(row, col).text = str(sheet.cell_value(row + 2, col))
                    table.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
            table.cell(0, 0).text = '解释序号'

            # for row in range(1, len(table.rows)):
            #     for col in range(len(table.columns)):
            #         table.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
        else:
            pass
        document.save(newFile)

        # 合并单元格
        if self.checkBox_11.isChecked():
            print('\n一界面单层统计表格式优化中……')
            for row in range(1, len(table.rows)):
                self.view_bar(row, len(table.rows) - 1)
                table.rows[row].height = Pt(20)
                for col in range(len(table.columns)):
                    # table.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # 首列居中
            # for row in range(len(table.rows)):
            #     table.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            pass
        document.save(newFile)
        print('\n一界面单层统计表添加完成')

        # 添加2单
        for fileName in os.listdir(PATH):
            if '2单' in fileName and '.xls' in fileName and '$' not in fileName:
                fileDir = PATH + "\\" + fileName
                workbook = xlrd.open_workbook(fileDir)
        sheet = workbook.sheets()[0]

        # 通过xlrd的接口获得表单的行数及列数
        nrow = sheet.nrows
        ncol = sheet.ncols

        document = Document(newFile)
        document.styles['Normal'].font.size = Pt(11)
        document.styles['Normal'].font.name = u'Times New Roman'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        table = document.tables[13]
        table.autofit = True
        for num in range(nrow - 3):
            table.add_row()

        # 设置整个表格字体属性
        table.style.font.color.rgb = RGBColor(0, 0, 0)
        table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table.cell(0, 0).width = Pt(30)
        table.cell(0, 1).width = Pt(100)

        # 单层评价表写入
        if self.checkBox_10.isChecked():
            for row in range(1, len(table.rows)):
                self.view_bar(row, len(table.rows) - 1)
                print(' @第', str(row), '行')
                for col in range(len(table.columns)):
                    table.cell(row, col).text = str(sheet.cell_value(row + 2, col))
                    table.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
            table.cell(0, 0).text = '解释序号'

            # for row in range(1, len(table.rows)):
            #     for col in range(len(table.columns)):
            #         table.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
        else:
            pass
        document.save(newFile)

        # 合并单元格
        if self.checkBox_11.isChecked():
            print('\n二界面单层统计表格式优化中……')
            for row in range(1, len(table.rows)):
                self.view_bar(row, len(table.rows) - 1)
                table.rows[row].height = Pt(20)
                for col in range(len(table.columns)):
                    # table.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # 首列居中
            # for row in range(len(table.rows)):
            #     table.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            pass
        document.save(newFile)
        print('\n二界面单层统计表添加完成')

        print('正在添加储层段落，请等待……')
        ################################################################################
        # 储层固井质量评价
        # p = document.add_paragraph()
        # p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # p.paragraph_format.line_spacing = Pt(24)
        # run = p.add_run(u"5．储层段固井质量分析")
        # run = document.add_heading('', level=2).add_run(u"5．储层段固井质量分析")
        para = document.add_heading(u'5．储层段固井质量分析', level=2)
        # run = document.add_paragraph().add_run(u"5．储层段固井质量分析")
        # run.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.runs[0].font.name = 'Times New Roman'  # 英文字体
        para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 中文字体
        para.runs[0].font.size = Pt(14)
        # run.bold = True
        para.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        evaluation_of_formation_upper1 = []
        all_evaluation_of_formation_upper1 = []
        evaluation_of_formation_upper2 = []
        all_evaluation_of_formation_upper2 = []
        if formation_be_or_not == '有储层':
            # 创建一个空的Dataframe
            formation_pic_DataFrame = pd.DataFrame(columns=('formation_StartNumber', 'formation_EndNumber', \
                                                            'formation_Start_Depth', 'formation_End_Depth'))
            # 添加段落
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run('该次测井井段有' + str(formation_Number) + '个解释储层。')
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)

            PATH = ".\\WorkSpace\\报告生成工区\\储层图"
            # 储层图片名
            all_Formation_Names = []
            # 储层图片后缀
            all_Formation_Extentions = []
            # 图片的数量
            count = 0
            for fileName in os.listdir(PATH):
                if '.db' not in fileName:
                    count += 1
                    all_Formation_Names.append(self.get_filePath_fileName_fileExt(fileName)[1])
                    all_Formation_Extentions.append(self.get_filePath_fileName_fileExt(fileName)[2])
                else:
                    pass
            # 利用lambda表达式排序
            all_Formation_Names.sort(key=lambda x: int(x.split('#')[0].split('-')[0]))
            ################################################################################
            # 得到储层上部固井质量评价深度DataFrame
            for pic_number in range(count):
                formation_Name_Split = all_Formation_Names[pic_number].split('#')
                if '-' in formation_Name_Split[0]:
                    formation_StartNumber = formation_Name_Split[0].split('-')[0]
                    formation_EndNumber = formation_Name_Split[0].split('-')[1]
                    formation_Start_End_Number = ''.join([formation_StartNumber, '-', formation_EndNumber])
                else:
                    formation_StartNumber = formation_Name_Split[0]
                    formation_EndNumber = formation_Name_Split[0]
                    formation_Start_End_Number = formation_Name_Split[0]
                formation_Start_Depth = formation_Name_Split[1].split('-')[0]
                formation_End_Depth = formation_Name_Split[1].split('-')[1]
                formation_Start_End = ''.join([formation_Start_Depth, '-', formation_End_Depth])
                ################################################################################
                # 单层统计表
                PATH = ".\\WorkSpace\\报告生成工区\\成果表"
                for fileName in os.listdir(PATH):
                    if '1单' in fileName and '$' not in fileName:
                        fileDir1 = PATH + "\\" + fileName
                    elif '2单' in fileName and '$' not in fileName:
                        fileDir2 = PATH + "\\" + fileName
                df1 = pd.read_excel(fileDir1, header=2)
                # df1.drop([0], inplace=True)
                df1.loc[:, '井段(m)'] = df1['井段(m)'].str.replace(' ', '')  # 消除数据中空格
                df1['井段Start'] = df1['井段(m)'].map(lambda x: x.split("-")[0])
                df1['井段End'] = df1['井段(m)'].map(lambda x: x.split("-")[1])

                df2 = pd.read_excel(fileDir2, header=2)
                # df2.drop([0], inplace=True)
                df2.loc[:, '井段(m)'] = df2['井段(m)'].str.replace(' ', '')  # 消除数据中空格
                df2['井段Start'] = df2['井段(m)'].map(lambda x: x.split("-")[0])
                df2['井段End'] = df2['井段(m)'].map(lambda x: x.split("-")[1])

                # 储层表
                PATH = ".\\WorkSpace\\报告生成工区\\储层表"
                for fileName in os.listdir(PATH):
                    fileDir3 = PATH + "\\" + fileName
                df3 = pd.read_excel(fileDir3, header=0)
                df3.drop([0], inplace=True)
                df3.drop(['层位', '解释结论'], axis=1, inplace=True)
                df3.loc[:, '井        段'] = df3['井        段'].str.replace(' ', '')  # 消除数据中空格
                df3['储层Start'] = df3['井        段'].map(lambda x: x.split("--")[0])
                df3['储层End'] = df3['井        段'].map(lambda x: x.split("--")[1])

                # 当前储层图片里第一个储层的上界深度
                formation_Start = df3.loc[int(formation_StartNumber), '储层Start']
                # 添加要分析的DataFrame
                formation_pic_DataFrame = formation_pic_DataFrame.append(
                    pd.DataFrame(
                        {'formation_StartNumber': [formation_StartNumber], 'formation_EndNumber': [formation_EndNumber], \
                         '当前储层图片Start': [formation_Start_Depth], '第一个储层start': [formation_Start]}), \
                    ignore_index=True)

            ################################################################################
            # 表格数据清洗
            df1.loc[:, "井段Start"] = df1["井段Start"].str.replace(" ", "").astype('float')
            df1.loc[:, "井段End"] = df1["井段End"].str.replace(" ", "").astype('float')
            df2.loc[:, "井段Start"] = df2["井段Start"].str.replace(" ", "").astype('float')
            df2.loc[:, "井段End"] = df2["井段End"].str.replace(" ", "").astype('float')
            df3.loc[:, "储层Start"] = df3["储层Start"].str.replace(" ", "").astype('float')
            df3.loc[:, "储层End"] = df3["储层End"].str.replace(" ", "").astype('float')
            rows1, cols1 = df1.shape
            rows2, cols2 = df2.shape
            rows3, cols3 = df3.shape

            rows4, cols4 = formation_pic_DataFrame.shape
            formation_pic_DataFrame.index = formation_pic_DataFrame.index + 1  # 调整序号

            # 针对每层在单层评价表中得出好中差比例
            for row in range(1, rows4 + 1):
                formation_Start_Depth = formation_pic_DataFrame.loc[row, '当前储层图片Start']  # 当前储层图片Start深度
                formation_Start = df3.loc[
                    int(formation_pic_DataFrame.loc[row, 'formation_EndNumber']), '储层Start']  # 对应的第一个储层开始深度
                # print('----------------第', row, '个储层上部的井段----------------')
                if (float(formation_Start) <= float(end_Evaluation)) & (
                        float(formation_Start_Depth) >= float(start_Evaluation)):
                    ratio_Series1 = self.layer_evaluation1(df1, float(formation_Start_Depth), float(formation_Start))[
                        0]  # 调取一界面评价函数
                    evaluation_of_formation_upper1 = \
                        self.layer_evaluation1(df1, float(formation_Start_Depth), float(formation_Start))[
                            1]  # 调取一界面评价函数
                    all_evaluation_of_formation_upper1.append(evaluation_of_formation_upper1)

                    ratio_Series2 = self.layer_evaluation2(df2, float(formation_Start_Depth), float(formation_Start))[
                        0]  # 调取二界面评价函数
                    evaluation_of_formation_upper2 = \
                        self.layer_evaluation2(df2, float(formation_Start_Depth), float(formation_Start))[
                            1]  # 调取二界面评价函数
                    all_evaluation_of_formation_upper2.append(evaluation_of_formation_upper2)
                else:
                    print('储层上部深度范围溢出，请检查')
                    pass

            ################################################################################
            # 储层上部描述输出
            # 图片的数量
            PATH = ".\\WorkSpace\\报告生成工区\\储层图"
            count = 0
            for fileName in os.listdir(PATH):
                count += 1
            for pic_number in range(count):
                formation_Name_Split = all_Formation_Names[pic_number].split('#')
                if '-' in formation_Name_Split[0]:
                    formation_StartNumber = formation_Name_Split[0].split('-')[0]
                    formation_EndNumber = formation_Name_Split[0].split('-')[1]
                    formation_Start_End_Number = ''.join([formation_StartNumber, '-', formation_EndNumber])
                else:
                    formation_StartNumber = formation_Name_Split[0]
                    formation_EndNumber = formation_Name_Split[0]
                    formation_Start_End_Number = formation_Name_Split[0]
                formation_Start_Depth = formation_Name_Split[1].split('-')[0]
                formation_End_Depth = formation_Name_Split[1].split('-')[1]
                try:  # 强制类型转换后取整
                    formation_Start_Depth = str(int(float(formation_Start_Depth)))
                    formation_End_Depth = str(int(float(formation_End_Depth)))
                except:
                    pass
                else:
                    pass
                formation_Start_End = ''.join([formation_Start_Depth, '-', formation_End_Depth])
                ###
                if all_evaluation_of_formation_upper1[pic_number] == '胶结好':
                    cbl_value = '低'
                    case_Wave_Energy = '弱'
                elif all_evaluation_of_formation_upper1[pic_number] == '胶结中等':
                    cbl_value = '中'
                    case_Wave_Energy = '较强'
                elif all_evaluation_of_formation_upper1[pic_number] == '胶结差':
                    cbl_value = '高'
                    case_Wave_Energy = '强'
                elif all_evaluation_of_formation_upper1[pic_number] == '胶结中到好，以好为主':
                    cbl_value = '中到低'
                    case_Wave_Energy = '较强到弱，以弱为主'
                elif all_evaluation_of_formation_upper1[pic_number] == '胶结中到好，以中等为主':
                    cbl_value = '中到低'
                    case_Wave_Energy = '较强到弱，以较强为主'
                elif all_evaluation_of_formation_upper1[pic_number] == '胶结中到差，以差为主':
                    cbl_value = '中到高'
                    case_Wave_Energy = '较强到强，以强为主'
                elif all_evaluation_of_formation_upper1[pic_number] == '胶结中到差，以中等为主':
                    cbl_value = '中到高'
                    case_Wave_Energy = '较强到强，以较强为主'
                elif all_evaluation_of_formation_upper1[pic_number] == '胶结好到差，以好为主':
                    cbl_value = '低到高'
                    case_Wave_Energy = '强到弱，以弱为主'
                elif all_evaluation_of_formation_upper1[pic_number] == '胶结好到差，以差为主':
                    cbl_value = '低到高'
                    case_Wave_Energy = '强到弱，以强为主'
                elif all_evaluation_of_formation_upper1[pic_number] == '胶结好到中等到差':
                    cbl_value = '低到高'
                    case_Wave_Energy = '强到较强到弱'
                elif all_evaluation_of_formation_upper1[pic_number] == '胶结好到中到差，以好为主':
                    cbl_value = '低到高'
                    case_Wave_Energy = '强到较强到弱，以弱为主'
                elif all_evaluation_of_formation_upper1[pic_number] == '胶结好到中到差，以中等为主':
                    cbl_value = '低到高'
                    case_Wave_Energy = '强到较强到弱，以较强为主'
                elif all_evaluation_of_formation_upper1[pic_number] == '胶结好到中到差，以差为主':
                    cbl_value = '低到高'
                    case_Wave_Energy = '强到较强到弱，以强为主'
                ###
                if all_evaluation_of_formation_upper2[pic_number] == '胶结好':
                    formation_Wave_Energy = '强'
                elif all_evaluation_of_formation_upper2[pic_number] == '胶结中等':
                    formation_Wave_Energy = '较强'
                elif all_evaluation_of_formation_upper2[pic_number] == '胶结差':
                    formation_Wave_Energy = '弱'
                elif all_evaluation_of_formation_upper2[pic_number] == '胶结中到好，以好为主':
                    formation_Wave_Energy = '较强到强，以强为主'
                elif all_evaluation_of_formation_upper2[pic_number] == '胶结中到好，以中等为主':
                    formation_Wave_Energy = '较强到强，以较强为主'
                elif all_evaluation_of_formation_upper2[pic_number] == '胶结中到差，以差为主':
                    formation_Wave_Energy = '较强到弱，以弱为主'
                elif all_evaluation_of_formation_upper2[pic_number] == '胶结中到差，以中等为主':
                    formation_Wave_Energy = '较强到弱，以较强为主'
                elif all_evaluation_of_formation_upper2[pic_number] == '胶结好到差，以好为主':
                    formation_Wave_Energy = '弱到强，以强为主'
                elif all_evaluation_of_formation_upper2[pic_number] == '胶结好到差，以差为主':
                    formation_Wave_Energy = '弱到强，以弱为主'
                elif all_evaluation_of_formation_upper2[pic_number] == '胶结好到中等到差':
                    formation_Wave_Energy = '强到较强到弱'
                elif all_evaluation_of_formation_upper2[pic_number] == '胶结好到中到差，以好为主':
                    formation_Wave_Energy = '强到较强到弱，以强为主'
                elif all_evaluation_of_formation_upper2[pic_number] == '胶结好到中到差，以中等为主':
                    formation_Wave_Energy = '强到较强到弱，以较强为主'
                elif all_evaluation_of_formation_upper2[pic_number] == '胶结好到中到差，以差为主':
                    formation_Wave_Energy = '强到较强到弱，以弱为主'
                elif all_evaluation_of_formation_upper2[pic_number] == '不确定':
                    formation_Wave_Energy = '受双层套管及外层介质干扰波影响，内层套管二界面水泥胶结质量不确定。'
                ###
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p.paragraph_format.line_spacing = Pt(24)
                p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
                r = p.add_run('（' + str(pic_number + 1) + '）' + formation_Start_End + 'm该封固井段上部声幅值')
                # r.bold = True
                r.font.name = 'Times New Roman'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(12)

                r = p.add_run(cbl_value)
                # r.bold = True
                r.font.name = 'Times New Roman'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(0, 0, 250)

                r = p.add_run('，一界面水泥')
                # r.bold = True
                r.font.name = 'Times New Roman'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(12)

                r = p.add_run(all_evaluation_of_formation_upper1[pic_number])
                # r.bold = True
                r.font.name = 'Times New Roman'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(0, 0, 250)

                r = p.add_run('；变密度曲线反映套管波能量')
                # r.bold = True
                r.font.name = 'Times New Roman'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(12)

                r = p.add_run(case_Wave_Energy)
                # r.bold = True
                r.font.name = 'Times New Roman'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(0, 0, 250)

                r = p.add_run('，地层波能量')
                # r.bold = True
                r.font.name = 'Times New Roman'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(12)

                r = p.add_run(formation_Wave_Energy)
                # r.bold = True
                r.font.name = 'Times New Roman'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(0, 0, 250)

                if formation_Wave_Energy != '受双层套管及外层介质干扰波影响，内层套管二界面水泥胶结质量不确定。':
                    r = p.add_run('，二界面水泥')
                    # r.bold = True
                    r.font.name = 'Times New Roman'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(12)

                    r = p.add_run(all_evaluation_of_formation_upper2[pic_number])
                    # r.bold = True
                    r.font.name = 'Times New Roman'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(12)
                    r.font.color.rgb = RGBColor(0, 0, 250)

                    r = p.add_run('。')
                    # r.bold = True
                    r.font.name = 'Times New Roman'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(12)
                ########################################################################################
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p.paragraph_format.line_spacing = Pt(24)
                p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符

                # 储层描述循环输出
                for formation_Number_Temp in range(int(formation_StartNumber), int(formation_EndNumber) + 1):
                    ###
                    if formation_Number_Temp <= len(all_evaluation_of_formation1):
                        if all_evaluation_of_formation1[formation_Number_Temp - 1] == '胶结好':
                            cbl_value = '低'
                            case_Wave_Energy = '弱'
                        elif all_evaluation_of_formation1[formation_Number_Temp - 1] == '胶结中等':
                            cbl_value = '中'
                            case_Wave_Energy = '较强'
                        elif all_evaluation_of_formation1[formation_Number_Temp - 1] == '胶结差':
                            cbl_value = '高'
                            case_Wave_Energy = '强'
                        elif all_evaluation_of_formation1[formation_Number_Temp - 1] == '胶结中到好，以好为主':
                            cbl_value = '中到低'
                            case_Wave_Energy = '较强到弱，以弱为主'
                        elif all_evaluation_of_formation1[formation_Number_Temp - 1] == '胶结中到好，以中等为主':
                            cbl_value = '中到低'
                            case_Wave_Energy = '较强到弱，以较强为主'
                        elif all_evaluation_of_formation1[formation_Number_Temp - 1] == '胶结中到差，以差为主':
                            cbl_value = '中到高'
                            case_Wave_Energy = '较强到强，以强为主'
                        elif all_evaluation_of_formation1[formation_Number_Temp - 1] == '胶结中到差，以中等为主':
                            cbl_value = '中到高'
                            case_Wave_Energy = '较强到强，以较强为主'
                        elif all_evaluation_of_formation1[formation_Number_Temp - 1] == '胶结好到差，以好为主':
                            cbl_value = '低到高'
                            case_Wave_Energy = '强到弱，以弱为主'
                        elif all_evaluation_of_formation1[formation_Number_Temp - 1] == '胶结好到差，以差为主':
                            cbl_value = '低到高'
                            case_Wave_Energy = '强到弱，以强为主'
                        elif all_evaluation_of_formation1[formation_Number_Temp - 1] == '胶结好到中等到差':
                            cbl_value = '低到高'
                            case_Wave_Energy = '弱到较强到强'
                        elif all_evaluation_of_formation1[formation_Number_Temp - 1] == '胶结好到中到差，以好为主':
                            cbl_value = '低到高'
                            case_Wave_Energy = '弱到较强到强，以弱为主'
                        elif all_evaluation_of_formation1[formation_Number_Temp - 1] == '胶结好到中到差，以中等为主':
                            cbl_value = '低到高'
                            case_Wave_Energy = '弱到较强到强，以较强为主'
                        elif all_evaluation_of_formation1[formation_Number_Temp - 1] == '胶结好到中到差，以差为主':
                            cbl_value = '低到高'
                            case_Wave_Energy = '弱到较强到强，以强为主'
                    else:
                        all_evaluation_of_formation1.append('[储层范围超出测量边界，待确定]')

                    if all_evaluation_of_formation1[formation_Number_Temp - 1] == '[储层范围超出测量边界，待确定]':
                        cbl_value = '[储层范围超出测量边界，待确定]'
                        case_Wave_Energy = '[储层范围超出测量边界，待确定]'
                    ###
                    if formation_Number_Temp <= len(all_evaluation_of_formation2):
                        if all_evaluation_of_formation2[formation_Number_Temp - 1] == '胶结好':
                            formation_Wave_Energy = '强'
                        elif all_evaluation_of_formation2[formation_Number_Temp - 1] == '胶结中等':
                            formation_Wave_Energy = '较强'
                        elif all_evaluation_of_formation2[formation_Number_Temp - 1] == '胶结差':
                            formation_Wave_Energy = '弱'
                        elif all_evaluation_of_formation2[formation_Number_Temp - 1] == '胶结中到好，以好为主':
                            formation_Wave_Energy = '较强到强，以强为主'
                        elif all_evaluation_of_formation2[formation_Number_Temp - 1] == '胶结中到好，以中等为主':
                            formation_Wave_Energy = '较强到强，以较强为主'
                        elif all_evaluation_of_formation2[formation_Number_Temp - 1] == '胶结中到差，以差为主':
                            formation_Wave_Energy = '较强到弱，以弱为主'
                        elif all_evaluation_of_formation2[formation_Number_Temp - 1] == '胶结中到差，以中等为主':
                            formation_Wave_Energy = '较强到弱，以较强为主'
                        elif all_evaluation_of_formation2[formation_Number_Temp - 1] == '胶结好到差，以好为主':
                            formation_Wave_Energy = '弱到强，以强为主'
                        elif all_evaluation_of_formation2[formation_Number_Temp - 1] == '胶结好到差，以差为主':
                            formation_Wave_Energy = '弱到强，以弱为主'
                        elif all_evaluation_of_formation2[formation_Number_Temp - 1] == '胶结好到中等到差':
                            formation_Wave_Energy = '强到较强到弱'
                        elif all_evaluation_of_formation2[formation_Number_Temp - 1] == '胶结好到中到差，以好为主':
                            formation_Wave_Energy = '强到较强到弱，以强为主'
                        elif all_evaluation_of_formation2[formation_Number_Temp - 1] == '胶结好到中到差，以中等为主':
                            formation_Wave_Energy = '强到较强到弱，以较强为主'
                        elif all_evaluation_of_formation2[formation_Number_Temp - 1] == '胶结好到中到差，以差为主':
                            formation_Wave_Energy = '强到较强到弱，以弱为主'
                        elif all_evaluation_of_formation2[formation_Number_Temp - 1] == '不确定':
                            formation_Wave_Energy = '受双层套管及外层介质干扰波影响，内层套管二界面水泥胶结质量不确定。'
                    else:
                        all_evaluation_of_formation2.append('[储层范围超出测量边界，待确定]')

                    if all_evaluation_of_formation2[formation_Number_Temp - 1] == '[储层范围超出测量边界，待确定]':
                        formation_Wave_Energy = '[储层范围超出测量边界，待确定]'
                    ###
                    r = p.add_run(str(formation_Number_Temp) + '#储层声幅值')
                    # r.bold = True
                    r.font.name = 'Times New Roman'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(12)

                    r = p.add_run(cbl_value)
                    # r.bold = True
                    r.font.name = 'Times New Roman'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(12)
                    r.font.color.rgb = RGBColor(0, 0, 250)

                    r = p.add_run('，一界面水泥')
                    # r.bold = True
                    r.font.name = 'Times New Roman'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(12)

                    r = p.add_run(all_evaluation_of_formation1[formation_Number_Temp - 1])
                    # r.bold = True
                    r.font.name = 'Times New Roman'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(12)
                    r.font.color.rgb = RGBColor(0, 0, 250)

                    r = p.add_run('；变密度曲线反映套管波能量')
                    # r.bold = True
                    r.font.name = 'Times New Roman'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(12)

                    r = p.add_run(case_Wave_Energy)
                    # r.bold = True
                    r.font.name = 'Times New Roman'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(12)
                    r.font.color.rgb = RGBColor(0, 0, 250)

                    r = p.add_run('、地层波能量')
                    # r.bold = True
                    r.font.name = 'Times New Roman'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(12)

                    r = p.add_run(formation_Wave_Energy)
                    # r.bold = True
                    r.font.name = 'Times New Roman'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(12)
                    r.font.color.rgb = RGBColor(0, 0, 250)

                    if formation_Wave_Energy != '受双层套管及外层介质干扰波影响，内层套管二界面水泥胶结质量不确定。':
                        r = p.add_run('，二界面水泥')
                        # r.bold = True
                        r.font.name = 'Times New Roman'
                        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        r.font.size = Pt(12)

                        r = p.add_run(all_evaluation_of_formation2[formation_Number_Temp - 1])
                        # r.bold = True
                        r.font.name = 'Times New Roman'
                        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        r.font.size = Pt(12)
                        r.font.color.rgb = RGBColor(0, 0, 250)

                        r = p.add_run('。')
                        # r.bold = True
                        r.font.name = 'Times New Roman'
                        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        r.font.size = Pt(12)

                r = p.add_run('（见图' + str(pic_number + 1) + '）')
                # r.bold = True
                r.font.name = 'Times New Roman'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(12)

                # 添加储层图片
                paragraph = document.add_paragraph()
                # 图片居中设置
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run("")
                PATH = ".\\WorkSpace\\报告生成工区\\储层图"
                if all_Formation_Extentions[0] == '.png':
                    run.add_picture(PATH + '\\' + all_Formation_Names[pic_number] + '.png', width=Inches(6.0))
                elif all_Formation_Extentions[0] == '.jpg':
                    run.add_picture(PATH + '\\' + all_Formation_Names[pic_number] + '.jpg', width=Inches(6.0))
                elif all_Formation_Extentions[0] == '.bmp':
                    run.add_picture(PATH + '\\' + all_Formation_Names[pic_number] + '.bmp', width=Inches(6.0))
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.paragraph_format.line_spacing = Pt(24)
                r = p.add_run(
                    '图' + str(pic_number + 1) + '  ' + well_Name + '井（' + formation_Start_End + 'm）固井处理成果图')
                # r.bold = True
                r.font.name = 'Times New Roman'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(12)
                # print('已添加第', str(pic_number + 1), '个储层的段落')
        else:
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            # r = p.add_run('该次测量井段内无储层解释。')
            r = p.add_run('本次测量井段内无裸眼测井综合解释数据。')
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
        ################################################################################
        # 判断是否有胶结为差
        PATH = ".\\WorkSpace\\报告生成工区\\胶结差图\\"
        bad_interval_be_or_not = ''
        if os.listdir(PATH) != []:
            for fileName in os.listdir(PATH):
                if '.db' not in fileName:
                    fileName = fileName
                else:
                    fileName = ''
        else:
            fileName = ''

        f_path = PATH + fileName

        if os.path.isdir(f_path):
            bad_interval_be_or_not = '无胶结差'
        else:
            bad_interval_be_or_not = '有胶结差'
        ################################################################################
        if bad_interval_be_or_not == '有胶结差':
            PATH = ".\\WorkSpace\\报告生成工区\\胶结差图"
            bad_Interval_Names = []
            bad_Interval_Extentions = []
            for fileName in os.listdir(PATH):
                if '.db' not in fileName:
                    bad_Interval_Names.append(self.get_filePath_fileName_fileExt(fileName)[1])
                    bad_Interval_Extentions.append(self.get_filePath_fileName_fileExt(fileName)[2])
                else:
                    pass

            # 利用lambda表达式排序
            bad_Interval_Names.sort(key=lambda x: int(x.split('-')[0]))

            bad_Start_Ends = []
            for bad_number in range(len(bad_Interval_Names)):
                bad_Name_Split = bad_Interval_Names[bad_number].split('-')
                bad_Serial_Number = bad_Name_Split[0]
                bad_Start_Depth = bad_Name_Split[1]
                bad_End_Depth = bad_Name_Split[2]
                try:  # 强制类型转换后取整
                    bad_Start_Depth = str(int(float(bad_Start_Depth)))
                    bad_End_Depth = str(int(float(bad_End_Depth)))
                except:
                    pass
                bad_Start_End = ''.join([bad_Start_Depth, '-', bad_End_Depth])
                bad_Start_Ends.append(bad_Start_End + 'm、')
            bad_Start_Ends = ''.join(bad_Start_Ends).rstrip('、')

            # p = document.add_paragraph()
            # p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            # p.paragraph_format.line_spacing = Pt(24)
            # run = p.add_run(u"三 建议及其它")
            # run.font.name = '黑体'  # 英文字体
            # run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 中文字体
            # run.font.size = Pt(16)
            # run.bold = True
            # run.font.color.rgb = RGBColor(0, 0, 0)

            para = document.add_heading(u"三 建议及其它", level=1)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.runs[0].font.name = '黑体'
            para.runs[0].font.size = Pt(16)
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            para.runs[0].bold = True
            para.runs[0].font.color.rgb = RGBColor(0, 0, 0)

            if formation_be_or_not == '无储层':
                pic_number = -1
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            if bad_number != 0:
                r = p.add_run(
                    bad_Start_Ends + '井段声幅值较高，套管接箍信号明显，一界面固井质量较差（见附图' + str(1) + '-' + str(
                        bad_number + 1) + '）。')
            elif bad_number == 0:
                r = p.add_run(
                    bad_Start_Ends + '井段声幅值较高，套管接箍信号明显，一界面固井质量较差（见附图1）。')
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)

            print('正在添加胶结为差段落，请等待……')
            # 添加固井质量差图片
            for bad_number in range(len(bad_Interval_Names)):
                bad_Name_Split = bad_Interval_Names[bad_number].split('-')
                bad_Serial_Number = bad_Name_Split[0]
                bad_Start_Depth = bad_Name_Split[1]
                bad_End_Depth = bad_Name_Split[2]
                try:  # 强制类型转换后取整
                    bad_Start_Depth = str(int(float(bad_Start_Depth)))
                    bad_End_Depth = str(int(float(bad_End_Depth)))
                except:
                    pass
                bad_Start_End = ''.join([bad_Start_Depth, '-', bad_End_Depth])

                paragraph = document.add_paragraph()
                # 图片居中设置
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run("")

                if bad_Interval_Extentions[0] == '.png':
                    run.add_picture(PATH + '\\' + bad_Interval_Names[bad_number] + '.png', width=Inches(6.0))
                elif bad_Interval_Extentions[0] == '.jpg':
                    run.add_picture(PATH + '\\' + bad_Interval_Names[bad_number] + '.jpg', width=Inches(6.0))
                elif bad_Interval_Extentions[0] == '.bmp':
                    run.add_picture(PATH + '\\' + bad_Interval_Names[bad_number] + '.bmp', width=Inches(6.0))

                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.paragraph_format.line_spacing = Pt(24)
                r = p.add_run(
                    '附图' + str(bad_number + 1) + ' ' + well_Name + '井（' + bad_Start_End + 'm）固井处理成果图')
                # r.bold = True
                r.font.name = 'Times New Roman'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(12)
                # print('已添加第', str(bad_number + 1), '个固井为差的段落')
        else:
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)

            para = document.add_heading(u"三 建议及其它", level=1)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.runs[0].font.name = '黑体'
            para.runs[0].font.size = Pt(16)
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            para.runs[0].bold = True
            para.runs[0].font.color.rgb = RGBColor(0, 0, 0)

            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run('本次测量井段内，一、二界面固井水泥胶结质量合格，且以好为主。')
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)

        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.line_spacing = Pt(24)
        p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
        r = p.add_run('附件：固井施工质量评价表（见附件1）')
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)

        p = document.add_paragraph()

        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.line_spacing = Pt(24)
        p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
        r = p.add_run('附件1 ' + well_Name + '井固井施工质量评价表')
        # r.bold = True
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)
        ################################################################################
        # 签名
        PATH = '.\\resources\\签名\\'
        report_Writer = self.comboBox_2.currentText()
        add1 = document.tables[0].cell(0, 1).paragraphs[0]
        if report_Writer == '李柯沁':
            add1.add_run().add_picture(PATH + '签名-李柯沁.jpg', width=Inches(1.0))
        elif report_Writer == '刘佳露':
            add1.add_run().add_picture(PATH + '签名-刘佳露.jpg', width=Inches(1.0))
        elif report_Writer == '杨晨曦':
            add1.add_run().add_picture(PATH + '签名-杨晨曦.jpg', width=Inches(1.0))
        elif report_Writer == '杨玉竹':
            add1.add_run().add_picture(PATH + '签名-杨玉竹.jpg', width=Inches(1.0))
        elif report_Writer == '闫跃星':
            add1.add_run().add_picture(PATH + '签名-闫跃星.jpg', width=Inches(1.0))
        elif report_Writer == '赵晓军':
            add1.add_run().add_picture(PATH + '签名-赵晓军.jpg', width=Inches(1.0))
        elif report_Writer == '王遂华':
            add1.add_run().add_picture(PATH + '签名-王遂华.jpg', width=Inches(1.0))
        elif report_Writer == '周政英':
            add1.add_run().add_picture(PATH + '签名-周政英.jpg', width=Inches(1.0))
        else:
            pass
        report_Supervisor = self.comboBox_4.currentText()
        add2 = document.tables[0].cell(1, 1).paragraphs[0]
        if report_Supervisor == '王参文':
            add2.add_run().add_picture(PATH + '签名-王参文.jpg', width=Inches(1.0))
        elif report_Supervisor == '刘静':
            add2.add_run().add_picture(PATH + '签名-刘静.jpg', width=Inches(1.0))
        elif report_Supervisor == '朱莉':
            add2.add_run().add_picture(PATH + '签名-朱莉.jpg', width=Inches(1.0))
        elif report_Supervisor == '王昌德':
            add2.add_run().add_picture(PATH + '签名-王昌德.jpg', width=Inches(1.0))
        else:
            pass
        document.save(newFile)
        print('报告生成完毕，请查看报告生成工区')
        self.pushButton_4.setText('生成解释报告')
        self.pushButton_4.setEnabled(True)
        self.timer.stop()
        self.lock.release()  # 解锁
        # 加上会死机，可能是线程冲突
        # QMessageBox.information(self, "提示", "水泥胶结评价报告生成完毕，请查看报告生成工区")

    ########################################################################################套损评价相关函数
    def generate_fast_report(self):
        if self.run_on_net == True:
            Supervisor.generate_CHL_result_usage_supervisor()
        else:
            pass

        if self.checkBox.isChecked() and self.checkBox_2.isChecked() and self.checkBox_3.isChecked():
            PATH = ".\\"
            for fileName in os.listdir(PATH):
                if '损伤' in fileName and '.xlsx' in fileName and '$' not in fileName:
                    fileDir = PATH + "\\" + fileName
                    workbook_penetration = xlrd.open_workbook(fileDir)

            PATH = ".\\"
            for fileName in os.listdir(PATH):
                if '结垢' in fileName and '.xlsx' in fileName and '$' not in fileName:
                    fileDir = PATH + "\\" + fileName
                    workbook_projection = xlrd.open_workbook(fileDir)

            PATH = ".\\"
            for fileName in os.listdir(PATH):
                if '变形' in fileName and '.xlsx' in fileName and '$' not in fileName:
                    fileDir = PATH + "\\" + fileName
                    workbook_trasformation = xlrd.open_workbook(fileDir)

            sheet1 = workbook_penetration.sheets()[0]
            sheet2 = workbook_projection.sheets()[0]
            sheet3 = workbook_trasformation.sheets()[0]

            ################################################损伤段落
            # 获得表单的行数及列数
            nrow1 = sheet1.nrows
            ncol1 = sheet1.ncols

            PATH = ".\\resources\\模板\\"
            for fileName in os.listdir(PATH):
                newFile = PATH + 'template-for-all.docx'
            document = Document(newFile)
            # table = document.add_table(rows=nrow, cols=ncol)
            table1 = document.tables[0]
            for num in range(nrow1 - 1):
                row_cells = table1.add_row()
            # 设置整个表格字体属性
            table1.style.font.color.rgb = RGBColor(0, 0, 0)
            table1.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table1.cell(0, 0).width = Pt(30)
            table1.cell(0, 1).width = Pt(100)
            # 写入表格
            for row in range(1, len(table1.rows)):
                print(' @第', str(row), '行')
                for col in range(len(table1.columns)):
                    table1.cell(row, col).text = str(sheet1.cell_value(row, col))

            # 格式优化
            for row in range(len(table1.rows)):
                table1.rows[row].height = Pt(20)
                for col in range(len(table1.columns)):
                    table1.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table1.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table1.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table1.cell(0, 0).text = '解释序号'

            # 首列居中
            for row in range(len(table1.rows)):
                table1.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            text1 = self.textEdit.toPlainText()

            p = document.paragraphs[1]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            # p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run(text1)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
            ################################################结垢段落
            # 获得表单的行数及列数
            nrow2 = sheet2.nrows
            ncol2 = sheet2.ncols

            # table = document.add_table(rows=nrow, cols=ncol)
            table2 = document.tables[1]
            for num in range(nrow2 - 1):
                row_cells = table2.add_row()
            # 设置整个表格字体属性
            table2.style.font.color.rgb = RGBColor(0, 0, 0)
            table2.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table2.cell(0, 0).width = Pt(30)
            table2.cell(0, 1).width = Pt(100)
            # 写入表格
            for row in range(1, len(table2.rows)):
                print(' @第', str(row), '行')
                for col in range(len(table2.columns)):
                    table2.cell(row, col).text = str(sheet2.cell_value(row, col))

            # 格式优化
            for row in range(len(table2.rows)):
                table2.rows[row].height = Pt(20)
                for col in range(len(table2.columns)):
                    table2.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table2.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table2.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table2.cell(0, 0).text = '解释序号'
            # 首列居中
            for row in range(len(table2.rows)):
                table2.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            text2 = self.textEdit_2.toPlainText()

            p = document.paragraphs[3]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            # p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run(text2)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
            ################################################变形段落
            # 获得表单的行数及列数
            nrow3 = sheet3.nrows
            ncol3 = sheet3.ncols

            # table = document.add_table(rows=nrow, cols=ncol)
            table3 = document.tables[2]
            for num in range(nrow3 - 1):
                row_cells = table3.add_row()
            # 设置整个表格字体属性
            table3.style.font.color.rgb = RGBColor(0, 0, 0)
            table3.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table3.cell(0, 0).width = Pt(30)
            table3.cell(0, 1).width = Pt(100)
            # 写入表格
            for row in range(1, len(table3.rows)):
                print(' @第', str(row), '行')
                for col in range(len(table3.columns)):
                    table3.cell(row, col).text = str(sheet3.cell_value(row, col))

            # 格式优化
            for row in range(len(table3.rows)):
                table3.rows[row].height = Pt(20)
                for col in range(len(table3.columns)):
                    table3.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table3.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table3.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table3.cell(0, 0).text = '解释序号'
            # 首列居中
            for row in range(len(table3.rows)):
                table3.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            text3 = self.textEdit_3.toPlainText()

            p = document.paragraphs[5]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            # p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run(text3)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
            ################################################
            document.save('.\\WorkSpace\\套损检测快速解释结论.docx')
        elif self.checkBox.isChecked() and self.checkBox_2.isChecked():  # 损伤和结垢 #############################################
            PATH = ".\\"
            for fileName in os.listdir(PATH):
                if '损伤' in fileName and '.xlsx' in fileName and '$' not in fileName:
                    fileDir = PATH + "\\" + fileName
                    workbook_penetration = xlrd.open_workbook(fileDir)

            PATH = ".\\"
            for fileName in os.listdir(PATH):
                if '结垢' in fileName and '.xlsx' in fileName and '$' not in fileName:
                    fileDir = PATH + "\\" + fileName
                    workbook_projection = xlrd.open_workbook(fileDir)

            sheet1 = workbook_penetration.sheets()[0]
            sheet2 = workbook_projection.sheets()[0]

            ################################################损伤段落
            # 获得表单的行数及列数
            nrow1 = sheet1.nrows
            ncol1 = sheet1.ncols

            PATH = ".\\resources\\模板\\"
            for fileName in os.listdir(PATH):
                newFile = PATH + 'template-for-pe+prj.docx'
            document = Document(newFile)
            # table = document.add_table(rows=nrow, cols=ncol)
            table1 = document.tables[0]
            for num in range(nrow1 - 1):
                row_cells = table1.add_row()
            # 设置整个表格字体属性
            table1.style.font.color.rgb = RGBColor(0, 0, 0)
            table1.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table1.cell(0, 0).width = Pt(30)
            table1.cell(0, 1).width = Pt(100)
            # 写入表格
            for row in range(1, len(table1.rows)):
                print(' @第', str(row), '行')
                for col in range(len(table1.columns)):
                    table1.cell(row, col).text = str(sheet1.cell_value(row, col))

            # 格式优化
            for row in range(len(table1.rows)):
                table1.rows[row].height = Pt(20)
                for col in range(len(table1.columns)):
                    table1.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table1.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table1.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table1.cell(0, 0).text = '解释序号'
            # 首列居中
            for row in range(len(table1.rows)):
                table1.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            text1 = self.textEdit.toPlainText()

            p = document.paragraphs[1]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            # p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run(text1)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
            ################################################结垢段落
            # 获得表单的行数及列数
            nrow2 = sheet2.nrows
            ncol2 = sheet2.ncols

            # table = document.add_table(rows=nrow, cols=ncol)
            table2 = document.tables[1]
            for num in range(nrow2 - 1):
                row_cells = table2.add_row()
            # 设置整个表格字体属性
            table2.style.font.color.rgb = RGBColor(0, 0, 0)
            table2.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table2.cell(0, 0).width = Pt(30)
            table2.cell(0, 1).width = Pt(100)
            # 写入表格
            for row in range(1, len(table2.rows)):
                print(' @第', str(row), '行')
                for col in range(len(table2.columns)):
                    table2.cell(row, col).text = str(sheet2.cell_value(row, col))

            # 格式优化
            for row in range(len(table2.rows)):
                table2.rows[row].height = Pt(20)
                for col in range(len(table2.columns)):
                    table2.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table2.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table2.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table2.cell(0, 0).text = '解释序号'
            # 首列居中
            for row in range(len(table2.rows)):
                table2.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            text2 = self.textEdit_2.toPlainText()

            p = document.paragraphs[3]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            # p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run(text2)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
            ################################################
            document.save('.\\WorkSpace\\套损检测快速解释结论.docx')
        elif self.checkBox_2.isChecked() and self.checkBox_3.isChecked():  # 结垢和变形 #############################################
            PATH = ".\\"
            for fileName in os.listdir(PATH):
                if '结垢' in fileName and '.xlsx' in fileName and '$' not in fileName:
                    fileDir = PATH + "\\" + fileName
                    workbook_projection = xlrd.open_workbook(fileDir)

            PATH = ".\\"
            for fileName in os.listdir(PATH):
                if '变形' in fileName and '.xlsx' in fileName and '$' not in fileName:
                    fileDir = PATH + "\\" + fileName
                    workbook_trasformation = xlrd.open_workbook(fileDir)

            sheet2 = workbook_projection.sheets()[0]
            sheet3 = workbook_trasformation.sheets()[0]

            PATH = ".\\resources\\模板\\"
            for fileName in os.listdir(PATH):
                newFile = PATH + 'template-for-prj+tr.docx'
            document = Document(newFile)

            ################################################结垢段落
            # 获得表单的行数及列数
            nrow2 = sheet2.nrows
            ncol2 = sheet2.ncols

            # table = document.add_table(rows=nrow, cols=ncol)
            table2 = document.tables[0]
            for num in range(nrow2 - 1):
                row_cells = table2.add_row()
            # 设置整个表格字体属性
            table2.style.font.color.rgb = RGBColor(0, 0, 0)
            table2.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table2.cell(0, 0).width = Pt(30)
            table2.cell(0, 1).width = Pt(100)
            # 写入表格
            for row in range(1, len(table2.rows)):
                print(' @第', str(row), '行')
                for col in range(len(table2.columns)):
                    table2.cell(row, col).text = str(sheet2.cell_value(row, col))

            # 格式优化
            for row in range(len(table2.rows)):
                table2.rows[row].height = Pt(20)
                for col in range(len(table2.columns)):
                    table2.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table2.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table2.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table2.cell(0, 0).text = '解释序号'
            # 首列居中
            for row in range(len(table2.rows)):
                table2.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            text2 = self.textEdit_2.toPlainText()

            p = document.paragraphs[1]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            # p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run(text2)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
            ################################################变形段落
            # 获得表单的行数及列数
            nrow3 = sheet3.nrows
            ncol3 = sheet3.ncols

            # table = document.add_table(rows=nrow, cols=ncol)
            table3 = document.tables[1]
            for num in range(nrow3 - 1):
                row_cells = table3.add_row()
            # 设置整个表格字体属性
            table3.style.font.color.rgb = RGBColor(0, 0, 0)
            table3.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table3.cell(0, 0).width = Pt(30)
            table3.cell(0, 1).width = Pt(100)
            # 写入表格
            for row in range(1, len(table3.rows)):
                print(' @第', str(row), '行')
                for col in range(len(table3.columns)):
                    table3.cell(row, col).text = str(sheet3.cell_value(row, col))

            # 格式优化
            for row in range(len(table3.rows)):
                table3.rows[row].height = Pt(20)
                for col in range(len(table3.columns)):
                    table3.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table3.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table3.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table3.cell(0, 0).text = '解释序号'
            # 首列居中
            for row in range(len(table3.rows)):
                table3.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            text3 = self.textEdit_3.toPlainText()

            p = document.paragraphs[3]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            # p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run(text3)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
            ################################################
            document.save('.\\WorkSpace\\套损检测快速解释结论.docx')
        elif self.checkBox.isChecked() and self.checkBox_3.isChecked():  # 损伤和变形 #############################################
            PATH = ".\\"
            for fileName in os.listdir(PATH):
                if '损伤' in fileName and '.xlsx' in fileName and '$' not in fileName:
                    fileDir = PATH + "\\" + fileName
                    workbook_penetration = xlrd.open_workbook(fileDir)

            PATH = ".\\"
            for fileName in os.listdir(PATH):
                if '变形' in fileName and '.xlsx' in fileName and '$' not in fileName:
                    fileDir = PATH + "\\" + fileName
                    workbook_trasformation = xlrd.open_workbook(fileDir)

            sheet1 = workbook_penetration.sheets()[0]
            sheet3 = workbook_trasformation.sheets()[0]

            ################################################损伤段落
            # 获得表单的行数及列数
            nrow1 = sheet1.nrows
            ncol1 = sheet1.ncols

            PATH = ".\\resources\\模板\\"
            for fileName in os.listdir(PATH):
                newFile = PATH + 'template-for-pe+tr.docx'
            document = Document(newFile)
            # table = document.add_table(rows=nrow, cols=ncol)
            table1 = document.tables[0]
            for num in range(nrow1 - 1):
                row_cells = table1.add_row()
            # 设置整个表格字体属性
            table1.style.font.color.rgb = RGBColor(0, 0, 0)
            table1.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table1.cell(0, 0).width = Pt(30)
            table1.cell(0, 1).width = Pt(100)
            # 写入表格
            for row in range(1, len(table1.rows)):
                print(' @第', str(row), '行')
                for col in range(len(table1.columns)):
                    table1.cell(row, col).text = str(sheet1.cell_value(row, col))

            # 格式优化
            for row in range(len(table1.rows)):
                table1.rows[row].height = Pt(20)
                for col in range(len(table1.columns)):
                    table1.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table1.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table1.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table1.cell(0, 0).text = '解释序号'
            # 首列居中
            for row in range(len(table1.rows)):
                table1.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            text1 = self.textEdit.toPlainText()

            p = document.paragraphs[1]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            # p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run(text1)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
            ################################################变形段落
            # 获得表单的行数及列数
            nrow3 = sheet3.nrows
            ncol3 = sheet3.ncols

            # table = document.add_table(rows=nrow, cols=ncol)
            table3 = document.tables[1]
            for num in range(nrow3 - 1):
                row_cells = table3.add_row()
            # 设置整个表格字体属性
            table3.style.font.color.rgb = RGBColor(0, 0, 0)
            table3.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table3.cell(0, 0).width = Pt(30)
            table3.cell(0, 1).width = Pt(100)
            # 写入表格
            for row in range(1, len(table3.rows)):
                print(' @第', str(row), '行')
                for col in range(len(table3.columns)):
                    table3.cell(row, col).text = str(sheet3.cell_value(row, col))

            # 格式优化
            for row in range(len(table3.rows)):
                table3.rows[row].height = Pt(20)
                for col in range(len(table3.columns)):
                    table3.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table3.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table3.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table3.cell(0, 0).text = '解释序号'
            # 首列居中
            for row in range(len(table3.rows)):
                table3.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            text3 = self.textEdit_3.toPlainText()

            p = document.paragraphs[3]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            # p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run(text3)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
            ################################################
            document.save('.\\WorkSpace\\套损检测快速解释结论.docx')
        elif self.checkBox.isChecked():  # 损伤
            PATH = ".\\"
            for fileName in os.listdir(PATH):
                if '损伤' in fileName and '.xlsx' in fileName and '$' not in fileName:
                    fileDir = PATH + "\\" + fileName
                    workbook_penetration = xlrd.open_workbook(fileDir)

            sheet = workbook_penetration.sheets()[0]
            # 获得表单的行数及列数
            nrow = sheet.nrows
            ncol = sheet.ncols

            PATH = ".\\resources\\模板\\"
            for fileName in os.listdir(PATH):
                newFile = PATH + 'template-for-pe.docx'
            document = Document(newFile)
            # table = document.add_table(rows=nrow, cols=ncol)
            table = document.tables[0]
            for num in range(nrow - 1):
                row_cells = table.add_row()
            # 设置整个表格字体属性
            table.style.font.color.rgb = RGBColor(0, 0, 0)
            table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table.cell(0, 0).width = Pt(30)
            table.cell(0, 1).width = Pt(100)
            # 写入表格
            for row in range(1, len(table.rows)):
                print(' @第', str(row), '行')
                for col in range(len(table.columns)):
                    table.cell(row, col).text = str(sheet.cell_value(row, col))

            # 格式优化
            for row in range(len(table.rows)):
                table.rows[row].height = Pt(20)
                for col in range(len(table.columns)):
                    table.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table.cell(0, 0).text = '解释序号'
            # 首列居中
            for row in range(len(table.rows)):
                table.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            text_all = self.textEdit.toPlainText()

            p = document.paragraphs[1]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            # p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run(text_all)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
            document.save('.\\WorkSpace\\套损检测快速解释结论.docx')
        elif self.checkBox_2.isChecked():  # 结垢
            PATH = ".\\"
            for fileName in os.listdir(PATH):
                if '结垢' in fileName and '.xlsx' in fileName and '$' not in fileName:
                    fileDir = PATH + "\\" + fileName
                    workbook_projection = xlrd.open_workbook(fileDir)

            sheet = workbook_projection.sheets()[0]
            # 获得表单的行数及列数
            nrow = sheet.nrows
            ncol = sheet.ncols

            PATH = ".\\resources\\模板\\"
            for fileName in os.listdir(PATH):
                newFile = PATH + 'template-for-prj.docx'
            document = Document(newFile)
            # table = document.add_table(rows=nrow, cols=ncol)
            table = document.tables[0]
            for num in range(nrow - 1):
                row_cells = table.add_row()
            # 设置整个表格字体属性
            table.style.font.color.rgb = RGBColor(0, 0, 0)
            table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table.cell(0, 0).width = Pt(30)
            table.cell(0, 1).width = Pt(100)
            # 写入表格
            for row in range(1, len(table.rows)):
                print(' @第', str(row), '行')
                for col in range(len(table.columns)):
                    table.cell(row, col).text = str(sheet.cell_value(row, col))

            # 格式优化
            for row in range(len(table.rows)):
                table.rows[row].height = Pt(20)
                for col in range(len(table.columns)):
                    table.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table.cell(0, 0).text = '解释序号'
            # 首列居中
            for row in range(len(table.rows)):
                table.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            text_all = self.textEdit_2.toPlainText()

            p = document.paragraphs[1]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            # p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run(text_all)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
            document.save('.\\WorkSpace\\套损检测快速解释结论.docx')
        elif self.checkBox_3.isChecked():  # 变形
            PATH = ".\\"
            for fileName in os.listdir(PATH):
                if '变形' in fileName and '.xlsx' in fileName and '$' not in fileName:
                    fileDir = PATH + "\\" + fileName
                    workbook_transformation = xlrd.open_workbook(fileDir)

            sheet = workbook_transformation.sheets()[0]
            # 获得表单的行数及列数
            nrow = sheet.nrows
            ncol = sheet.ncols

            PATH = ".\\resources\\模板\\"
            for fileName in os.listdir(PATH):
                newFile = PATH + 'template-for-tr.docx'
            document = Document(newFile)
            # table = document.add_table(rows=nrow, cols=ncol)
            table = document.tables[0]
            for num in range(nrow - 1):
                row_cells = table.add_row()
            # 设置整个表格字体属性
            table.style.font.color.rgb = RGBColor(0, 0, 0)
            table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table.cell(0, 0).width = Pt(30)
            table.cell(0, 1).width = Pt(100)
            # 写入表格
            for row in range(1, len(table.rows)):
                print(' @第', str(row), '行')
                for col in range(len(table.columns)):
                    table.cell(row, col).text = str(sheet.cell_value(row, col))

            # 格式优化
            for row in range(len(table.rows)):
                table.rows[row].height = Pt(20)
                for col in range(len(table.columns)):
                    table.cell(row, 0).text = str(row)  # 因为序号带小数，重新赋值
                    table.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table.cell(0, 0).text = '解释序号'
            # 首列居中
            for row in range(len(table.rows)):
                table.cell(row, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            text_all = self.textEdit_3.toPlainText()

            p = document.paragraphs[1]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            # p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进0.74厘米，即2个字符
            r = p.add_run(text_all)
            # r.bold = True
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
            document.save('.\\WorkSpace\\套损检测快速解释结论.docx')
        else:
            QMessageBox.information(self, '提示', '您似乎没有勾选')

        QMessageBox.information(self, "提示", "快速解释结论生成完毕，请查看WorkSpace")

    def changecb_type1(self):
        if self.checkBox_6.isChecked():
            self.checkBox_7.setCheckState(Qt.Unchecked)
            self.checkBox_8.setCheckState(Qt.Unchecked)
            self.checkBox_12.setCheckState(Qt.Unchecked)

    def changecb_type2(self):
        if self.checkBox_7.isChecked():
            self.checkBox_6.setCheckState(Qt.Unchecked)
            self.checkBox_8.setCheckState(Qt.Unchecked)
            self.checkBox_12.setCheckState(Qt.Unchecked)

    def changecb_type3(self):
        if self.checkBox_8.isChecked():
            self.checkBox_6.setCheckState(Qt.Unchecked)
            self.checkBox_7.setCheckState(Qt.Unchecked)
            self.checkBox_12.setCheckState(Qt.Unchecked)

    def changecb_type4(self):
        if self.checkBox_12.isChecked():
            self.checkBox_6.setCheckState(Qt.Unchecked)
            self.checkBox_7.setCheckState(Qt.Unchecked)
            self.checkBox_8.setCheckState(Qt.Unchecked)

    def changecb1(self):
        if self.checkBox_4.checkState() == Qt.Checked:
            self.checkBox.setChecked(True)
            self.checkBox_2.setChecked(True)
            self.checkBox_3.setChecked(True)
        elif self.checkBox_4.checkState() == Qt.Unchecked:
            self.checkBox.setChecked(False)
            self.checkBox_2.setChecked(False)
            self.checkBox_3.setChecked(False)

    def changecb2(self):
        if self.checkBox.isChecked() and self.checkBox_2.isChecked() and self.checkBox_3.isChecked():
            self.checkBox_4.setCheckState(Qt.Checked)
        elif self.checkBox.isChecked() or self.checkBox_2.isChecked() or self.checkBox_3.isChecked():
            self.checkBox_4.setTristate()
            self.checkBox_4.setCheckState(Qt.PartiallyChecked)
        else:
            self.checkBox_4.setTristate(False)
            self.checkBox_4.setCheckState(Qt.Unchecked)

    def clean_the_dir(self):
        my_files = ['.\\WorkSpace\\套损检测快速解释结论.docx', '.\\Penetration.xlsx', '.\\Projection.xlsx',
                    '.\\Transformation.xlsx', \
                    '.\\损伤评价表.xlsx', '.\\结垢评价表.xlsx', '.\\变形评价表.xlsx', '.\\casing_data.xls']
        for my_file in my_files:
            if os.path.exists(my_file):
                # 删除文件，可使用以下两种方法
                os.remove(my_file)
                # os.unlink(my_file)
            else:
                print('no such file:%s' % my_file)

        QMessageBox.information(self, "提示", "清理完毕")
        # self.echo(reply)

    def open_las_file(self):
        fnames = QFileDialog.getOpenFileNames(self, '打开LAS文件', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.textEdit_7.append(fname)

    def read_las_file(self):
        fileDir = self.textEdit_7.toPlainText()

        # 解决中文乱码问题
        # myfont = fm.FontProperties(fname=r"C:\Windows\Fonts\simsun.ttc", size=14)
        # matplotlib.rcParams["axes.unicode_minus"] = False
        # 读取las数据
        # log = las.LASReader(u'ning209H19-4_resample_jz.LAS', null_subs=np.nan)
        # log = las.LASReader(fileDir, null_subs=np.nan)
        if self.checkBox_6.checkState() == Qt.Checked:
            self.mat_view = MATPLOTLIB_MIT24(fileDir)  # 关键一步，创建Matplot_class类对象
        elif self.checkBox_7.checkState() == Qt.Checked:
            self.mat_view = MATPLOTLIB_MIT60(fileDir)
        elif self.checkBox_8.checkState() == Qt.Checked:
            self.mat_view = MATPLOTLIB_MFC40(fileDir)
        elif self.checkBox_12.checkState() == Qt.Checked:
            self.mat_view = MATPLOTLIB_MFC24(fileDir)

    def table_casing(self):
        self.tableWidget_5.setColumnCount(5)
        self.tableWidget_5.setRowCount(0)
        row = 0  # 第几行（从0开始）
        col = 0  # 第几列（从0开始）
        # self.tableWidget_5.horizontalHeader().setStretchLastSection(True)  # 设置最后一列拉伸至最大
        self.tableWidget_5.horizontalHeader().setSectionsClickable(False)  # 禁止点击表头的列
        self.headers = ['起始井深(m)', '结束井深(m)', '套管外径(mm)', '套管内径(mm)', '壁厚(mm)']
        self.tableWidget_5.setHorizontalHeaderLabels(self.headers)

        self.tableWidget_5.setColumnWidth(0, 120)
        self.tableWidget_5.setColumnWidth(1, 120)
        self.tableWidget_5.setColumnWidth(2, 120)
        self.tableWidget_5.setColumnWidth(3, 120)
        self.tableWidget_5.setColumnWidth(4, 100)

        # self.tableWidget_5.setRowHeight(0, 50)
        # self.tableWidget_5.verticalHeader().setVisible(False)  # 隐藏垂直表头
        # self.tableWidget_5.horizontalHeader().setVisible(False)  # 隐藏水平表头

    def add_line_for_tableWidget_5(self):
        self.row_tableWidget_5 = self.tableWidget_5.rowCount()
        self.tableWidget_5.setRowCount(self.row_tableWidget_5 + 1)

        # 添加默认表格数据
        self.tableWidget_5.setItem(self.row_tableWidget_5, 0, QTableWidgetItem('0'))
        self.tableWidget_5.setItem(self.row_tableWidget_5, 1, QTableWidgetItem('9999'))
        self.tableWidget_5.setItem(self.row_tableWidget_5, 2, QTableWidgetItem('139.7'))
        self.tableWidget_5.setItem(self.row_tableWidget_5, 3, QTableWidgetItem('114.3'))
        self.tableWidget_5.setItem(self.row_tableWidget_5, 4, QTableWidgetItem('12.7'))

    def delete_line_for_tableWidget_5(self):
        # self.table.cellChanged.disconnect()
        self.row = self.tableWidget_5.rowCount()
        self.tableWidget_5.setRowCount(self.row - 1)

    def casing_info_save(self):
        self.row_tableWidget_5 = self.tableWidget_5.rowCount()
        xls = xlwt.Workbook()
        sht1 = xls.add_sheet('Sheet1')
        # 添加字段
        sht1.write(0, 0, self.tableWidget_5.item(0, 0).text())
        sht1.write(0, 1, self.tableWidget_5.item(0, 1).text())
        sht1.write(0, 2, self.tableWidget_5.item(0, 2).text())
        sht1.write(0, 3, self.tableWidget_5.item(0, 3).text())
        sht1.write(0, 4, self.tableWidget_5.item(0, 4).text())
        xls.save('.\\casing_data.xls')
        QMessageBox.information(self, "提示", "套管数据已保存")

    ############################################################################################
    # 损伤评价
    ############################################################################################

    def table2(self):
        self.tableWidget_2.setColumnCount(9)
        self.tableWidget_2.setRowCount(0)
        row = 0  # 第几行（从0开始）
        col = 0  # 第几列（从0开始）
        # self.tableWidget_2.horizontalHeader().setStretchLastSection(True)  # 设置最后一列拉伸至最大
        self.tableWidget_2.horizontalHeader().setSectionsClickable(False)  # 禁止点击表头的列
        self.headers = ['起始深度(m)', '结束深度(m)', '最大损伤点深度(m)', '臂号', '单臂测量值(mm)', '正常测量值(mm)',
                        '最小内径(mm)', '平均内径(mm)',
                        '最大内径(mm)']
        self.tableWidget_2.setHorizontalHeaderLabels(self.headers)

        self.tableWidget_2.setColumnWidth(0, 110)
        self.tableWidget_2.setColumnWidth(1, 110)
        self.tableWidget_2.setColumnWidth(2, 145)
        self.tableWidget_2.setColumnWidth(3, 60)
        self.tableWidget_2.setColumnWidth(4, 130)
        self.tableWidget_2.setColumnWidth(5, 130)
        self.tableWidget_2.setColumnWidth(6, 120)
        self.tableWidget_2.setColumnWidth(7, 120)
        self.tableWidget_2.setColumnWidth(8, 120)

        # self.tableWidget_2.setRowHeight(0, 50)
        # self.tableWidget_2.verticalHeader().setVisible(False)  # 隐藏垂直表头
        # self.tableWidget_2.horizontalHeader().setVisible(False)  # 隐藏水平表头

    def add_blank_line_for_tableWidget_2(self):
        self.row_tableWidget_2 = self.tableWidget_2.rowCount()
        self.tableWidget_2.setRowCount(self.row_tableWidget_2 + 1)
        # 添加表格数据
        self.tableWidget_2.setItem(self.row_tableWidget_2, 0, QTableWidgetItem(''))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 1, QTableWidgetItem(''))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 2, QTableWidgetItem(''))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 3, QTableWidgetItem(''))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 4, QTableWidgetItem(''))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 5, QTableWidgetItem(''))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 6, QTableWidgetItem(''))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 7, QTableWidgetItem(''))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 8, QTableWidgetItem(''))

    def add_blank_line_for_tableWidget_3(self):
        self.row_tableWidget_3 = self.tableWidget_3.rowCount()
        self.tableWidget_3.setRowCount(self.row_tableWidget_3 + 1)
        # 添加表格数据
        self.tableWidget_3.setItem(self.row_tableWidget_3, 0, QTableWidgetItem(''))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 1, QTableWidgetItem(''))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 2, QTableWidgetItem(''))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 3, QTableWidgetItem(''))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 4, QTableWidgetItem(''))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 5, QTableWidgetItem(''))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 6, QTableWidgetItem(''))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 7, QTableWidgetItem(''))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 8, QTableWidgetItem(''))

    def add_blank_line_for_tableWidget_4(self):
        self.row_tableWidget_4 = self.tableWidget_4.rowCount()
        self.tableWidget_4.setRowCount(self.row_tableWidget_4 + 1)
        # 添加表格数据
        self.tableWidget_4.setItem(self.row_tableWidget_4, 0, QTableWidgetItem(''))
        self.tableWidget_4.setItem(self.row_tableWidget_4, 1, QTableWidgetItem(''))
        self.tableWidget_4.setItem(self.row_tableWidget_4, 2, QTableWidgetItem(''))
        self.tableWidget_4.setItem(self.row_tableWidget_4, 3, QTableWidgetItem(''))
        self.tableWidget_4.setItem(self.row_tableWidget_4, 4, QTableWidgetItem(''))
        self.tableWidget_4.setItem(self.row_tableWidget_4, 5, QTableWidgetItem(''))

    def add_line_for_tableWidget_2(self):
        # 读取文件中的数值
        data = pd.read_excel('Penetration.xlsx')
        self.start_Depth = data.loc[0, 'value']
        self.end_Depth = data.loc[1, 'value']
        self.critical_Depth = data.loc[2, 'value']
        self.finger_Number = data.loc[3, 'value']
        self.finger_Value = data.loc[4, 'value']
        self.normal_Value = data.loc[5, 'value']
        self.min_Diameter = data.loc[6, 'value']
        self.ave_Diameter = data.loc[7, 'value']
        self.max_Diameter = data.loc[8, 'value']
        # self.table.cellChanged.disconnect()
        self.row_tableWidget_2 = self.tableWidget_2.rowCount()
        self.tableWidget_2.setRowCount(self.row_tableWidget_2 + 1)
        # 添加表格数据
        self.tableWidget_2.setItem(self.row_tableWidget_2, 0, QTableWidgetItem(str(self.start_Depth)))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 1, QTableWidgetItem(str(self.end_Depth)))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 2, QTableWidgetItem(str(self.critical_Depth)))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 3, QTableWidgetItem(str(int(self.finger_Number))))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 4, QTableWidgetItem(str(self.finger_Value)))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 5, QTableWidgetItem(str(self.normal_Value)))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 6, QTableWidgetItem(str(self.min_Diameter)))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 7, QTableWidgetItem(str(self.ave_Diameter)))
        self.tableWidget_2.setItem(self.row_tableWidget_2, 8, QTableWidgetItem(str(self.max_Diameter)))

    def delete_line_for_tableWidget_2(self):
        # self.table.cellChanged.disconnect()
        self.row = self.tableWidget_2.rowCount()
        self.tableWidget_2.setRowCount(self.row - 1)

    def generate_results_from_tableWidget_2(self):
        self.row_tableWidget_2 = self.tableWidget_2.rowCount()
        text_all = ''
        data = pd.DataFrame({'损伤井段(m)': '',
                             '最大损伤点深度(m)': '',
                             '单臂测量值(mm)': '',
                             '最小内径(mm)': '',
                             '平均内径(mm)': '',
                             '最大内径(mm)': '',
                             '损伤量(mm)': '',
                             '损伤程度(%)': '',
                             '损伤级别': ''},
                            index=[1])
        for row in range(0, self.row_tableWidget_2):
            self.row_tableWidget_5 = self.tableWidget_5.rowCount()
            if self.row_tableWidget_5 == 1:
                inner_Diameter = self.tableWidget_5.item(0, 3).text()
                inner_Diameter = float(inner_Diameter)
                thickness = self.tableWidget_5.item(0, 4).text()
                thickness = float(thickness)
                print('最大损伤点深度落在套管第一段')
            elif self.row_tableWidget_5 == 2:
                if float(self.tableWidget_2.item(row, 2).text()) < float(
                        self.tableWidget_5.item(0, 1).text()):  # 最大损伤点深度落在套管第一段
                    inner_Diameter = self.tableWidget_5.item(0, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(0, 4).text()
                    thickness = float(thickness)
                    print('最大损伤点深度落在套管第一段')
                elif float(self.tableWidget_2.item(row, 2).text()) >= float(
                        self.tableWidget_5.item(1, 0).text()):  # 最大损伤点深度落在套管第二段
                    inner_Diameter = self.tableWidget_5.item(1, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(1, 4).text()
                    thickness = float(thickness)
                    print('最大损伤点深度落在套管第二段')
            elif self.row_tableWidget_5 == 3:
                if float(self.tableWidget_2.item(row, 2).text()) < float(
                        self.tableWidget_5.item(0, 1).text()):  # 最大损伤点深度落在套管第一段
                    inner_Diameter = self.tableWidget_5.item(0, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(0, 4).text()
                    thickness = float(thickness)
                    print('最大损伤点深度落在套管第一段')
                elif (float(self.tableWidget_2.item(row, 2).text()) >= float(
                        self.tableWidget_5.item(1, 0).text())) and \
                        (float(self.tableWidget_2.item(row, 2).text()) < float(
                            self.tableWidget_5.item(1, 1).text())):  # 最大损伤点深度落在套管第二段
                    inner_Diameter = self.tableWidget_5.item(1, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(1, 4).text()
                    thickness = float(thickness)
                    print('最大损伤点深度落在套管第二段')
                elif (float(self.tableWidget_2.item(row, 2).text()) >= float(
                        self.tableWidget_5.item(2, 0).text())) and \
                        (float(self.tableWidget_2.item(row, 2).text()) < float(
                            self.tableWidget_5.item(2, 1).text())):  # 最大损伤点深度落在套管第三段
                    inner_Diameter = self.tableWidget_5.item(2, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(2, 4).text()
                    thickness = float(thickness)
                    print('最大损伤点深度落在套管第三段')

            well_interval = ''.join(
                [self.tableWidget_2.item(row, 0).text(), '-', self.tableWidget_2.item(row, 1).text()])
            penetration_value = float(self.tableWidget_2.item(row, 4).text()) - float(
                self.tableWidget_2.item(row, 5).text())
            penetration_value = round(penetration_value, 2)
            penetration_degree = penetration_value / thickness * 100
            penetration_degree = round(penetration_degree, 2)

            # 损伤级别判定
            penetration_describe = ''
            if penetration_degree < 10:
                penetration_describe = '一级损伤'
            elif 10 <= penetration_degree < 20:
                penetration_describe = '二级损伤'
            elif 20 <= penetration_degree < 40:
                penetration_describe = '三级损伤'
            elif 40 <= penetration_degree < 85:
                penetration_describe = '四级损伤'
            elif penetration_degree >= 85:
                penetration_describe = '五级损伤'

            data_single_layer = pd.DataFrame({'损伤井段(m)': well_interval,
                                              '最大损伤点深度(m)': self.tableWidget_2.item(row, 2).text(),
                                              '单臂测量值(mm)': self.tableWidget_2.item(row, 4).text(),
                                              '最小内径(mm)': self.tableWidget_2.item(row, 6).text(),
                                              '平均内径(mm)': self.tableWidget_2.item(row, 7).text(),
                                              '最大内径(mm)': self.tableWidget_2.item(row, 8).text(),
                                              '损伤量(mm)': str(penetration_value),
                                              '损伤程度(%)': str(penetration_degree),
                                              '损伤级别': str(penetration_describe)},
                                             index=[1])
            data = pd.concat([data, data_single_layer], ignore_index=True)
            text = ''.join(
                ['井段', well_interval, 'm，存在套管损伤，最大损伤点深度为', self.tableWidget_2.item(row, 2).text(),
                 'm，第', \
                 self.tableWidget_2.item(row, 3).text(), '号臂测得的最大值为',
                 self.tableWidget_2.item(row, 4).text(), \
                 'mm，该臂在正常段的测量值为', self.tableWidget_2.item(row, 5).text(),
                 'mm，在最大损伤点深度处测量得到的最小内径为', \
                 self.tableWidget_2.item(row, 6).text(), 'mm，测量平均内径为',
                 self.tableWidget_2.item(row, 7).text(), \
                 'mm，测量最大内径为', self.tableWidget_2.item(row, 8).text(), 'mm，最大损伤量为',
                 str(penetration_value), \
                 'mm，损伤程度为', str(penetration_degree), '%，根据解释标准评价为', str(penetration_describe), '。\n'])
            text_all = ''.join([text_all, text])
        data.drop([0], inplace=True)
        print(data)
        writer = pd.ExcelWriter('损伤评价表.xlsx')
        data.to_excel(writer, 'Sheet1')
        writer.save()

        # 生成描述建议
        self.textEdit.setText(text_all)

    ############################################################################################
    # 结垢评价
    ############################################################################################

    def table3(self):
        self.tableWidget_3.setColumnCount(9)
        self.tableWidget_3.setRowCount(0)
        row = 0  # 第几行（从0开始）
        col = 0  # 第几列（从0开始）
        # self.tableWidget_3.horizontalHeader().setStretchLastSection(True)  # 设置最后一列拉伸至最大
        self.tableWidget_3.horizontalHeader().setSectionsClickable(False)  # 禁止点击表头的列
        self.headers = ['起始深度(m)', '结束深度(m)', '最大结垢点深度(m)', '臂号', '单臂测量值(mm)', '正常测量值(mm)',
                        '最小内径(mm)', '平均内径(mm)',
                        '最大内径(mm)']
        self.tableWidget_3.setHorizontalHeaderLabels(self.headers)

        self.tableWidget_3.setColumnWidth(0, 110)
        self.tableWidget_3.setColumnWidth(1, 110)
        self.tableWidget_3.setColumnWidth(2, 145)
        self.tableWidget_3.setColumnWidth(3, 60)
        self.tableWidget_3.setColumnWidth(4, 130)
        self.tableWidget_3.setColumnWidth(5, 130)
        self.tableWidget_3.setColumnWidth(6, 120)
        self.tableWidget_3.setColumnWidth(7, 120)
        self.tableWidget_3.setColumnWidth(8, 120)

        # self.tableWidget_3.setRowHeight(0, 50)
        # self.tableWidget_3.verticalHeader().setVisible(False)  # 隐藏垂直表头
        # self.tableWidget_3.horizontalHeader().setVisible(False)  # 隐藏水平表头

    def add_line_for_tableWidget_3(self):
        # 读取文件中的数值
        data = pd.read_excel('Projection.xlsx')
        self.start_Depth = data.loc[0, 'value']
        self.end_Depth = data.loc[1, 'value']
        self.critical_Depth = data.loc[2, 'value']
        self.finger_Number = data.loc[3, 'value']
        self.finger_Value = data.loc[4, 'value']
        self.normal_Value = data.loc[5, 'value']
        self.min_Diameter = data.loc[6, 'value']
        self.ave_Diameter = data.loc[7, 'value']
        self.max_Diameter = data.loc[8, 'value']
        # self.table.cellChanged.disconnect()
        self.row_tableWidget_3 = self.tableWidget_3.rowCount()
        self.tableWidget_3.setRowCount(self.row_tableWidget_3 + 1)
        # 添加表格数据
        self.tableWidget_3.setItem(self.row_tableWidget_3, 0, QTableWidgetItem(str(self.start_Depth)))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 1, QTableWidgetItem(str(self.end_Depth)))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 2, QTableWidgetItem(str(self.critical_Depth)))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 3, QTableWidgetItem(str(int(self.finger_Number))))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 4, QTableWidgetItem(str(self.finger_Value)))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 5, QTableWidgetItem(str(self.normal_Value)))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 6, QTableWidgetItem(str(self.min_Diameter)))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 7, QTableWidgetItem(str(self.ave_Diameter)))
        self.tableWidget_3.setItem(self.row_tableWidget_3, 8, QTableWidgetItem(str(self.max_Diameter)))

    def delete_line_for_tableWidget_3(self):
        # self.table.cellChanged.disconnect()
        self.row = self.tableWidget_3.rowCount()
        self.tableWidget_3.setRowCount(self.row - 1)

    def generate_results_from_tableWidget_3(self):
        self.row_tableWidget_3 = self.tableWidget_3.rowCount()
        text_all = ''
        data = pd.DataFrame({'结垢井段(m)': '',
                             '最大结垢点深度(m)': '',
                             '单臂测量值(mm)': '',
                             '最小内径(mm)': '',
                             '平均内径(mm)': '',
                             '最大内径(mm)': '',
                             '结垢量(mm)': '',
                             '结垢程度(%)': '',
                             '结垢级别': ''},
                            index=[1])
        for row in range(0, self.row_tableWidget_3):
            self.row_tableWidget_5 = self.tableWidget_5.rowCount()
            if self.row_tableWidget_5 == 1:
                inner_Diameter = self.tableWidget_5.item(0, 3).text()
                inner_Diameter = float(inner_Diameter)
                thickness = self.tableWidget_5.item(0, 4).text()
                thickness = float(thickness)
                print('最大结垢点深度落在套管第一段')
            elif self.row_tableWidget_5 == 2:
                if float(self.tableWidget_3.item(row, 2).text()) < float(
                        self.tableWidget_5.item(0, 1).text()):  # 最大结垢点深度落在套管第一段
                    inner_Diameter = self.tableWidget_5.item(0, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(0, 4).text()
                    thickness = float(thickness)
                    print('最大结垢点深度落在套管第一段')
                elif float(self.tableWidget_3.item(row, 2).text()) >= float(
                        self.tableWidget_5.item(1, 0).text()):  # 最大结垢点深度落在套管第二段
                    inner_Diameter = self.tableWidget_5.item(1, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(1, 4).text()
                    thickness = float(thickness)
                    print('最大结垢点深度落在套管第二段')
            elif self.row_tableWidget_5 == 3:
                if float(self.tableWidget_3.item(row, 2).text()) < float(
                        self.tableWidget_5.item(0, 1).text()):  # 最大结垢点深度落在套管第一段
                    inner_Diameter = self.tableWidget_5.item(0, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(0, 4).text()
                    thickness = float(thickness)
                    print('最大结垢点深度落在套管第一段')
                elif (float(self.tableWidget_3.item(row, 2).text()) >= float(
                        self.tableWidget_5.item(1, 0).text())) and \
                        (float(self.tableWidget_3.item(row, 2).text()) < float(
                            self.tableWidget_5.item(1, 1).text())):  # 最大结垢点深度落在套管第二段
                    inner_Diameter = self.tableWidget_5.item(1, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(1, 4).text()
                    thickness = float(thickness)
                    print('最大结垢点深度落在套管第二段')
                elif (float(self.tableWidget_3.item(row, 2).text()) >= float(
                        self.tableWidget_5.item(2, 0).text())) and \
                        (float(self.tableWidget_3.item(row, 2).text()) < float(
                            self.tableWidget_5.item(2, 1).text())):  # 最大损伤点深度落在套管第三段
                    inner_Diameter = self.tableWidget_5.item(2, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(2, 4).text()
                    thickness = float(thickness)
                    print('最大结垢点深度落在套管第三段')

            well_interval = ''.join(
                [self.tableWidget_3.item(row, 0).text(), '-', self.tableWidget_3.item(row, 1).text()])
            projection_value = float(self.tableWidget_3.item(row, 5).text()) - float(
                self.tableWidget_3.item(row, 4).text())
            projection_value = round(projection_value, 2)
            projection_degree = projection_value / (inner_Diameter / 2) * 100
            projection_degree = round(projection_degree, 2)

            # 结垢级别判定
            projection_describe = ''
            if projection_degree < 10:
                projection_describe = '一级结垢'
            elif 10 <= projection_degree < 20:
                projection_describe = '二级结垢'
            elif 20 <= projection_degree < 40:
                projection_describe = '三级结垢'
            elif 40 <= projection_degree < 85:
                projection_describe = '四级结垢'
            elif projection_degree >= 85:
                projection_describe = '五级结垢'

            data_single_layer = pd.DataFrame({'结垢井段(m)': well_interval,
                                              '最大结垢点深度(m)': self.tableWidget_3.item(row, 2).text(),
                                              '单臂测量值(mm)': self.tableWidget_3.item(row, 4).text(),
                                              '最小内径(mm)': self.tableWidget_3.item(row, 6).text(),
                                              '平均内径(mm)': self.tableWidget_3.item(row, 7).text(),
                                              '最大内径(mm)': self.tableWidget_3.item(row, 8).text(),
                                              '结垢量(mm)': str(projection_value),
                                              '结垢程度(%)': str(projection_degree),
                                              '结垢级别': str(projection_describe)},
                                             index=[1])
            data = pd.concat([data, data_single_layer], ignore_index=True)
            text = ''.join(
                ['井段', well_interval, 'm，存在套管结垢，最大结垢点深度为', self.tableWidget_3.item(row, 2).text(),
                 'm，第', \
                 self.tableWidget_3.item(row, 3).text(), '号臂测得的最小值为',
                 self.tableWidget_3.item(row, 4).text(), \
                 'mm，该臂在正常段的测量值为', self.tableWidget_3.item(row, 5).text(),
                 'mm，在最大结垢点深度处测量得到的最小内径为', \
                 self.tableWidget_3.item(row, 6).text(), 'mm，测量平均内径为',
                 self.tableWidget_3.item(row, 7).text(), \
                 'mm，测量最大内径为', self.tableWidget_3.item(row, 8).text(), 'mm，最大结垢量为',
                 str(projection_value), \
                 'mm，结垢程度为', str(projection_degree), '%，根据解释标准评价为', str(projection_describe), '。\n'])
            text_all = ''.join([text_all, text])
        data.drop([0], inplace=True)
        print(data)
        writer = pd.ExcelWriter('结垢评价表.xlsx')
        data.to_excel(writer, 'Sheet1')
        writer.save()

        # 生成描述建议
        self.textEdit_2.setText(text_all)

    ############################################################################################
    # 变形评价
    ############################################################################################
    def table4(self):
        self.tableWidget_4.setColumnCount(6)
        self.tableWidget_4.setRowCount(0)
        row = 0  # 第几行（从0开始）
        col = 0  # 第几列（从0开始）
        # self.tableWidget_4.horizontalHeader().setStretchLastSection(True)  # 设置最后一列拉伸至最大
        self.tableWidget_4.horizontalHeader().setSectionsClickable(False)  # 禁止点击表头的列
        self.headers = ['起始深度(m)', '结束深度(m)', '最大变形点深度(m)', '最小内径(mm)', '平均内径(mm)',
                        '最大内径(mm)']
        self.tableWidget_4.setHorizontalHeaderLabels(self.headers)

        self.tableWidget_4.setColumnWidth(0, 110)
        self.tableWidget_4.setColumnWidth(1, 110)
        self.tableWidget_4.setColumnWidth(2, 145)
        self.tableWidget_4.setColumnWidth(3, 110)
        self.tableWidget_4.setColumnWidth(4, 110)
        self.tableWidget_4.setColumnWidth(5, 110)

        # self.tableWidget_4.setRowHeight(0, 50)
        # self.tableWidget_4.verticalHeader().setVisible(False)  # 隐藏垂直表头
        # self.tableWidget_4.horizontalHeader().setVisible(False)  # 隐藏水平表头

    def add_line_for_tableWidget_4(self):
        # 读取文件中的数值
        data = pd.read_excel('Transformation.xlsx')
        self.start_Depth = data.loc[0, 'value']
        self.end_Depth = data.loc[1, 'value']
        self.critical_Depth = data.loc[2, 'value']
        self.min_Diameter = data.loc[3, 'value']
        self.ave_Diameter = data.loc[4, 'value']
        self.max_Diameter = data.loc[5, 'value']
        # self.table.cellChanged.disconnect()
        self.row_tableWidget_4 = self.tableWidget_4.rowCount()
        self.tableWidget_4.setRowCount(self.row_tableWidget_4 + 1)
        # 添加表格数据
        self.tableWidget_4.setItem(self.row_tableWidget_4, 0, QTableWidgetItem(str(self.start_Depth)))
        self.tableWidget_4.setItem(self.row_tableWidget_4, 1, QTableWidgetItem(str(self.end_Depth)))
        self.tableWidget_4.setItem(self.row_tableWidget_4, 2, QTableWidgetItem(str(self.critical_Depth)))
        self.tableWidget_4.setItem(self.row_tableWidget_4, 3, QTableWidgetItem(str(self.min_Diameter)))
        self.tableWidget_4.setItem(self.row_tableWidget_4, 4, QTableWidgetItem(str(self.ave_Diameter)))
        self.tableWidget_4.setItem(self.row_tableWidget_4, 5, QTableWidgetItem(str(self.max_Diameter)))

    def delete_line_for_tableWidget_4(self):
        # self.table.cellChanged.disconnect()
        self.row = self.tableWidget_4.rowCount()
        self.tableWidget_4.setRowCount(self.row - 1)

    def generate_results_from_tableWidget_4(self):
        self.row_tableWidget_4 = self.tableWidget_4.rowCount()

        text_all = ''
        data = pd.DataFrame({'变形井段(m)': '',
                             '变形长度(m)': '',
                             '最大变形点深度(m)': '',
                             '最小内径(mm)': '',
                             '平均内径(mm)': '',
                             '最大内径(mm)': '',
                             '变形量(mm)': '',
                             '变形程度(%)': '',
                             '变形级别': ''},
                            index=[1])
        for row in range(0, self.row_tableWidget_4):
            self.row_tableWidget_5 = self.tableWidget_5.rowCount()
            if self.row_tableWidget_5 == 1:
                inner_Diameter = self.tableWidget_5.item(0, 3).text()
                inner_Diameter = float(inner_Diameter)
                thickness = self.tableWidget_5.item(0, 4).text()
                thickness = float(thickness)
                print('最大损伤点深度落在套管第一段')
            elif self.row_tableWidget_5 == 2:
                if float(self.tableWidget_4.item(row, 2).text()) < float(
                        self.tableWidget_5.item(0, 1).text()):  # 最大损伤点深度落在套管第一段
                    inner_Diameter = self.tableWidget_5.item(0, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(0, 4).text()
                    thickness = float(thickness)
                    print('最大损伤点深度落在套管第一段')
                elif float(self.tableWidget_4.item(row, 2).text()) >= float(
                        self.tableWidget_5.item(1, 0).text()):  # 最大损伤点深度落在套管第二段
                    inner_Diameter = self.tableWidget_5.item(1, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(1, 4).text()
                    thickness = float(thickness)
                    print('最大损伤点深度落在套管第二段')
            elif self.row_tableWidget_5 == 3:
                if float(self.tableWidget_4.item(row, 2).text()) < float(
                        self.tableWidget_5.item(0, 1).text()):  # 最大损伤点深度落在套管第一段
                    inner_Diameter = self.tableWidget_5.item(0, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(0, 4).text()
                    thickness = float(thickness)
                    print('最大损伤点深度落在套管第一段')
                elif (float(self.tableWidget_4.item(row, 2).text()) >= float(
                        self.tableWidget_5.item(1, 0).text())) and \
                        (float(self.tableWidget_4.item(row, 2).text()) < float(
                            self.tableWidget_5.item(1, 1).text())):  # 最大损伤点深度落在套管第二段
                    inner_Diameter = self.tableWidget_5.item(1, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(1, 4).text()
                    thickness = float(thickness)
                    print('最大损伤点深度落在套管第二段')
                elif (float(self.tableWidget_4.item(row, 2).text()) >= float(
                        self.tableWidget_5.item(2, 0).text())) and \
                        (float(self.tableWidget_4.item(row, 2).text()) < float(
                            self.tableWidget_5.item(2, 1).text())):  # 最大损伤点深度落在套管第三段
                    inner_Diameter = self.tableWidget_5.item(2, 3).text()
                    inner_Diameter = float(inner_Diameter)
                    thickness = self.tableWidget_5.item(2, 4).text()
                    thickness = float(thickness)
                    print('最大损伤点深度落在套管第三段')

            well_interval = ''.join(
                [self.tableWidget_4.item(row, 0).text(), '-', self.tableWidget_4.item(row, 1).text()])
            transformation_length = round(
                float(self.tableWidget_4.item(row, 1).text()) - float(self.tableWidget_4.item(row, 0).text()), 2)
            transformation_value = max(float(inner_Diameter - float(self.tableWidget_4.item(row, 3).text())),
                                       float(float(self.tableWidget_4.item(row, 5).text()) - inner_Diameter))
            transformation_value = round(transformation_value, 2)
            transformation_degree = round(transformation_value / inner_Diameter * 100, 2)

            # 变形级别判定
            if transformation_length > 10:
                if transformation_degree <= 5:
                    transformation_describe = '一级变形'
                elif 5 < transformation_degree <= 10:
                    transformation_describe = '二级变形'
                elif 10 < transformation_degree <= 20:
                    transformation_describe = '三级变形'
                elif 20 < transformation_degree <= 40:
                    transformation_describe = '四级变形'
                elif 40 < transformation_degree:
                    transformation_describe = '五级变形'
            elif transformation_length <= 10:
                if transformation_degree <= 10:
                    transformation_describe = '一级变形'
                elif 10 < transformation_degree <= 20:
                    transformation_describe = '二级变形'
                elif 20 < transformation_degree <= 40:
                    transformation_describe = '三级变形'
                elif 40 < transformation_degree <= 60:
                    transformation_describe = '四级变形'
                elif 60 < transformation_degree:
                    transformation_describe = '五级变形'

            data_single_layer = pd.DataFrame({'变形井段(m)': well_interval,
                                              '变形长度(m)': str(transformation_length),
                                              '最大变形点深度(m)': self.tableWidget_4.item(row, 2).text(),
                                              '最小内径(mm)': self.tableWidget_4.item(row, 3).text(),
                                              '平均内径(mm)': self.tableWidget_4.item(row, 4).text(),
                                              '最大内径(mm)': self.tableWidget_4.item(row, 5).text(),
                                              '变形量(mm)': str(transformation_value),
                                              '变形程度(%)': str(transformation_degree),
                                              '变形级别': str(transformation_describe)},
                                             index=[1])
            data = pd.concat([data, data_single_layer], ignore_index=True)
            text = ''.join(['井段', well_interval, 'm，存在套管变形特征，变形长度为', str(transformation_length), \
                            'm，最大变形点深度为', self.tableWidget_4.item(row, 2).text(), 'm，测量最小内径为', \
                            self.tableWidget_4.item(row, 3).text(), 'mm，测量平均内径为',
                            self.tableWidget_4.item(row, 4).text(), \
                            'mm，测量最大内径为', self.tableWidget_4.item(row, 5).text(), 'mm，最大变形量为',
                            str(transformation_value), \
                            'mm，变形程度为', str(transformation_degree), '%，根据解释标准评价为',
                            str(transformation_describe),
                            '。\n'])
            text_all = ''.join([text_all, text])
        data.drop([0], inplace=True)
        print(data)
        writer = pd.ExcelWriter('变形评价表.xlsx')
        data.to_excel(writer, 'Sheet1')
        writer.save()

        # 生成描述建议
        self.textEdit_3.setText(text_all)

    def now(self):
        return time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))

    def open_list_file1(self):
        fnames = QFileDialog.getOpenFileNames(self, '打开第一个文件', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.lineEdit_112.setText(fname)

    def open_list_file2(self):
        fnames = QFileDialog.getOpenFileNames(self, '打开第二个文件', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.lineEdit_114.setText(fname)

    def assemble_list_together(self):

        fileDir1 = self.lineEdit_112.text()
        fileDir2 = self.lineEdit_114.text()
        assembled_point = float(self.lineEdit_113.text())

        with open(fileDir1, "r") as f:  # 打开文件
            data1 = f.readlines()  # 读取文件

        with open(fileDir2, "r") as f:  # 打开文件
            data2 = f.readlines()  # 读取文件

        for item in data1[2:]:
            if float(item.split(',')[1].split('-')[0]) <= assembled_point <= float(item.split(',')[1].split('-')[1]):
                item_number1 = int(item.split(',')[0])
                data1 = data1[0:(item_number1 + 2)]
                break
            else:
                pass
        # print(item_number1)
        # print(data1)

        for item in data2[2:]:
            if float(item.split(',')[1].split('-')[0]) <= assembled_point <= float(item.split(',')[1].split('-')[1]):
                item_number2 = int(item.split(',')[0])
                data2 = data2[(item_number2 + 1):]
                break
            else:
                pass
        # print(item_number2)
        # print(data2)

        count = 0
        data3 = []
        for item in data2:
            item = ','.join(
                [str(item_number1 + count), item.split(',')[1], item.split(',')[2], item.split(',')[3],
                 item.split(',')[4],
                 item.split(',')[5], item.split(',')[6]])
            count = count + 1
            data3.append(item)

        # print(data3)
        first = data3[0].split(',')[0]
        second = ''.join([data1[-1].split(',')[1].split('-')[0], '-', data3[0].split(',')[1].split('-')[1]])
        third = round(float(data3[0].split(',')[1].split('-')[1]) - float(data1[-1].split(',')[1].split('-')[0]), 2)
        fouth = max(float(data1[-1].split(',')[3]), float(data3[0].split(',')[3]))
        fifth = min(float(data1[-1].split(',')[4]), float(data3[0].split(',')[4]))
        sixth = round((float(data1[-1].split(',')[5]) + float(data3[0].split(',')[5])) / 2, 2)
        seventh = data3[0].split(',')[6]
        data3[0] = ','.join([first, str(second), str(third), str(fouth), str(fifth), str(sixth), seventh])

        temp = fileDir1.split('/')[-1]
        path = fileDir1.replace(temp, '')
        f = open(path + '组合后的数据.list', 'w', encoding='UTF-8')
        for item1 in data1[0:-1]:
            f.write(item1)
        for item3 in data3:
            f.write(item3)
        f.close()

        with open(path + "组合后的数据.list", "r", encoding='UTF-8') as f:  # 打开文件
            data4 = f.read()  # 读取文件
        # print(data4)
        QMessageBox.information(self, "提示", "list文件拼接完毕，请在源数据同级目录查看")

    ##############################
    # 成果表工具集
    ##############################
    # 单层评价表拼接和分段统计界面
    def open_file1(self):
        fnames = QFileDialog.getOpenFileNames(None, '打开第一个文件', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.lineEdit_48.setText(fname)

    def open_file2(self):
        fnames = QFileDialog.getOpenFileNames(None, '打开第二个文件', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.lineEdit_52.setText(fname)

    # 固井质量综合评价界面
    def open_file3(self):
        fnames = QFileDialog.getOpenFileNames(None, '打开第一个文件', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.lineEdit_75.setText(fname)

    def open_file4(self):
        fnames = QFileDialog.getOpenFileNames(None, '打开第二个文件', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.lineEdit_73.setText(fname)

    # 查询优中差比例界面
    def open_file5(self):
        fnames = QFileDialog.getOpenFileNames(None, '打开第一个文件', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.lineEdit_77.setText(fname)

    def open_file6(self):
        fnames = QFileDialog.getOpenFileNames(None, '打开第二个文件', './')  # 注意这里返回值是元组
        if fnames[0]:
            for fname in fnames[0]:
                self.lineEdit_76.setText(fname)

    ##############################

    # 规范化1
    ##############################
    def xls_formatting_first_layer(self, path1):  # 用于一界面表格表头文字规范
        # path2 = ''.join([path1.replace('.' + path1.split('.')[-1], ''), '(已规范化).xls'])
        path2 = path1
        old_excel = xlrd.open_workbook(path1, formatting_info=True)
        row_num = old_excel.sheets()[0].nrows
        col_num = old_excel.sheets()[0].ncols
        # print(row_num, ' ', col_num)

        ############################################### 若为两大列则进行单列规范化
        if int(col_num) > 10:
            # 将操作文件对象拷贝，变成可写的workbook对象
            new_excel = copy(old_excel)
            # 获得第一个sheet的对象
            ws = new_excel.get_sheet(0)
            # 写入数据
            ws.write(2, 0, '解释序号')
            ws.write(2, 1, '井段(m)')
            ws.write(2, 2, '厚度(m)')
            ws.write(2, 3, '最大声幅(%)')
            ws.write(2, 4, '最小声幅(%)')
            ws.write(2, 5, '平均声幅(%)')
            ws.write(2, 6, '结论')

            # 写入第一大列
            for row in range(3, row_num):
                for col in range(0, col_num):
                    item = old_excel.sheets()[0].cell_value(row, col)
                    ws.write(row_num + row - 3, col, item)
            # 写入第二大列
            for row in range(3, row_num):
                for col in range(8, col_num):
                    item = old_excel.sheets()[0].cell_value(row, col)
                    ws.write(row_num + row - 3, col - 8, item)

            # 另存为excel文件，并将文件命名
            new_excel.save(path2)

            ######################## 删除无用列
            # 创建app，打开工作表
            app = xw.App(visible=False, add_book=False)
            app.screen_updating = False
            app.display_alerts = False
            load_wb = app.books.open(path2)
            load_ws = load_wb.sheets.active
            # 处理列，将指定列从大到小删除（避免先删除小列导致后续列号变动）
            # load_ws.api.columns('O').delete 这样会报错，不要用
            load_ws.range('O1').api.EntireColumn.Delete()
            load_ws.range('N1').api.EntireColumn.Delete()
            load_ws.range('M1').api.EntireColumn.Delete()
            load_ws.range('L1').api.EntireColumn.Delete()
            load_ws.range('K1').api.EntireColumn.Delete()
            load_ws.range('J1').api.EntireColumn.Delete()
            load_ws.range('I1').api.EntireColumn.Delete()
            load_ws.range('H1').api.EntireColumn.Delete()

            # 获取行数
            info = load_ws.used_range
            last_row = info.last_cell.row
            alpha_dict = {1: 'A', 2: 'B', 3: 'C', 4: 'D', 5: 'E', 6: 'F', 7: 'G', 8: 'H', 9: 'I', 10: 'J', 11: 'K',
                          12: 'L', 13: 'M', 14: 'N', 15: 'O', 16: 'P', 17: 'Q', 18: 'R', 19: 'S', 20: 'T', 21: 'U',
                          22: 'V', 23: 'W', 24: 'X', 25: 'Y', 26: 'Z'}
            last_column_num = info.last_cell.column
            last_column = alpha_dict[last_column_num]
            # print(last_row, ' ', last_column)

            # 删除最后的空行，单数行情况
            last_row_value = load_ws.range(last_row, last_column_num).value
            if last_row_value == None:
                # load_ws.api.rows(last_row).delete
                load_ws.range(f'{last_column}{last_row}').api.EntireRow.Delete()
                last_row = str(int(last_row) - 1)

            # 设置边框
            a_range = f'A1:{last_column}{last_row}'  # 生成表格的数据范围
            load_ws.range(a_range).api.Borders(8).LineStyle = 1  # 上边框
            load_ws.range(a_range).api.Borders(9).LineStyle = 1  # 下边框
            load_ws.range(a_range).api.Borders(7).LineStyle = 1  # 左边框
            load_ws.range(a_range).api.Borders(10).LineStyle = 1  # 右边框
            load_ws.range(a_range).api.Borders(12).LineStyle = 1  # 内横边框
            load_ws.range(a_range).api.Borders(11).LineStyle = 1  # 内纵边框
            load_ws.range(a_range).api.Font.Name = 'Times New Roman'
            load_ws.range(a_range).api.RowHeight = 20

            load_ws.range(a_range).columns.autofit()

            b_range = load_ws.range(a_range)
            # 设置单元格 字体格式
            b_range.color = 255, 255, 255  # 设置单元格的填充颜色
            b_range.api.Font.ColorIndex = 1  # 设置字体的颜色，具体颜色索引见下方。
            b_range.api.Font.Size = 11  # 设置字体的大小。
            b_range.api.Font.Bold = False  # 设置为粗体。
            b_range.api.HorizontalAlignment = -4108  # -4108 水平居中。 -4131 靠左，-4152 靠右。
            b_range.api.VerticalAlignment = -4108  # -4108 垂直居中（默认）。 -4160 靠上，-4107 靠下， -4130 自动换行对齐。
            # b_range.api.NumberFormat = "0.00"          # 设置单元格的数字格式。

            # 处理完毕，保存、关闭、退出Excel
            load_wb.save()
            load_wb.close()
            app.quit()

            # if os.path.exists(path1):
            #     # 删除文件，可使用以下两种方法
            #     os.remove(path1)
            #     # os.unlink(my_file)
            # else:
            #     print('no such file:%s' % path1)

        ############################################### 若为一大列则进行直接进行规范化
        elif int(col_num) < 10:
            # 将操作文件对象拷贝，变成可写的workbook对象
            new_excel = copy(old_excel)
            # 获得第一个sheet的对象
            ws = new_excel.get_sheet(0)
            # 写入数据
            ws.write(2, 0, '解释序号')
            ws.write(2, 1, '井段(m)')
            ws.write(2, 2, '厚度(m)')
            ws.write(2, 3, '最大声幅(%)')
            ws.write(2, 4, '最小声幅(%)')
            ws.write(2, 5, '平均声幅(%)')
            ws.write(2, 6, '结论')
            # 另存为excel文件，并将文件命名
            new_excel.save(path2)

            ######################## 格式进一步规范
            # 创建app，打开工作表
            app = xw.App(visible=False, add_book=False)
            app.screen_updating = False
            app.display_alerts = False
            load_wb = app.books.open(path2)
            load_ws = load_wb.sheets.active

            # 获取行数
            info = load_ws.used_range
            last_row = info.last_cell.row
            alpha_dict = {1: 'A', 2: 'B', 3: 'C', 4: 'D', 5: 'E', 6: 'F', 7: 'G', 8: 'H', 9: 'I', 10: 'J', 11: 'K',
                          12: 'L', 13: 'M', 14: 'N', 15: 'O', 16: 'P', 17: 'Q', 18: 'R', 19: 'S', 20: 'T', 21: 'U',
                          22: 'V', 23: 'W', 24: 'X', 25: 'Y', 26: 'Z'}
            last_column_num = info.last_cell.column
            last_column = alpha_dict[last_column_num]
            # print(last_row, ' ', last_column)

            # 删除最后的空行，单数行情况
            '''
            last_row_value = load_ws.range(last_row, last_column_num).value
            if last_row_value in [None, '']:
                while (last_row_value in [None, '']):
                    # load_ws.api.rows(last_row).delete
                    load_ws.range(f'{last_column}{last_row}').api.EntireRow.Delete()
                    last_row = str(int(last_row) - 1)
                    last_row_value = load_ws.range(last_row, last_column_num).value
            '''
            last_row_value = load_ws.range(last_row, last_column_num).value
            if last_row_value == None:
                # load_ws.api.rows(last_row).delete
                load_ws.range(f'{last_column}{last_row}').api.EntireRow.Delete()
                last_row = str(int(last_row) - 1)

            # 设置边框
            a_range = f'A1:{last_column}{last_row}'  # 生成表格的数据范围
            load_ws.range(a_range).api.Borders(8).LineStyle = 1  # 上边框
            load_ws.range(a_range).api.Borders(9).LineStyle = 1  # 下边框
            load_ws.range(a_range).api.Borders(7).LineStyle = 1  # 左边框
            load_ws.range(a_range).api.Borders(10).LineStyle = 1  # 右边框
            load_ws.range(a_range).api.Borders(12).LineStyle = 1  # 内横边框
            load_ws.range(a_range).api.Borders(11).LineStyle = 1  # 内纵边框
            load_ws.range(a_range).api.Font.Name = 'Times New Roman'
            load_ws.range(a_range).api.RowHeight = 20

            load_ws.range(a_range).columns.autofit()

            b_range = load_ws.range(a_range)
            # 设置单元格 字体格式
            b_range.color = 255, 255, 255  # 设置单元格的填充颜色
            b_range.api.Font.ColorIndex = 1  # 设置字体的颜色，具体颜色索引见下方。
            b_range.api.Font.Size = 11  # 设置字体的大小。
            b_range.api.Font.Bold = False  # 设置为粗体。
            b_range.api.HorizontalAlignment = -4108  # -4108 水平居中。 -4131 靠左，-4152 靠右。
            b_range.api.VerticalAlignment = -4108  # -4108 垂直居中（默认）。 -4160 靠上，-4107 靠下， -4130 自动换行对齐。
            # b_range.api.NumberFormat = "0.00"          # 设置单元格的数字格式。

            # 处理完毕，保存、关闭、退出Excel
            load_wb.save()
            load_wb.close()
            app.quit()

        else:
            pass

    def xls_formatting_second_layer(self, path1):  # 用于二界面表格表头文字规范
        # path2 = ''.join([path1.replace('.' + path1.split('.')[-1], ''), '(已规范化).xls'])
        path2 = path1
        old_excel = xlrd.open_workbook(path1, formatting_info=True)
        row_num = old_excel.sheets()[0].nrows
        col_num = old_excel.sheets()[0].ncols
        # print(row_num, ' ', col_num)

        ############################################### 若为两大列则进行单列规范化
        if int(col_num) > 10:
            # 将操作文件对象拷贝，变成可写的workbook对象
            new_excel = copy(old_excel)
            # 获得第一个sheet的对象
            ws = new_excel.get_sheet(0)
            # 写入数据
            ws.write(2, 0, '解释序号')
            ws.write(2, 1, '井段(m)')
            ws.write(2, 2, '厚度(m)')
            ws.write(2, 3, '最大指数')
            ws.write(2, 4, '最小指数')
            ws.write(2, 5, '平均指数')
            ws.write(2, 6, '结论')

            # 写入第一大列
            for row in range(3, row_num):
                for col in range(0, col_num):
                    item = old_excel.sheets()[0].cell_value(row, col)
                    ws.write(row_num + row - 3, col, item)
            # 写入第二大列
            for row in range(3, row_num):
                for col in range(8, col_num):
                    item = old_excel.sheets()[0].cell_value(row, col)
                    ws.write(row_num + row - 3, col - 8, item)

            # 另存为excel文件，并将文件命名
            new_excel.save(path2)

            ######################## 删除无用列
            # 创建app，打开工作表
            app = xw.App(visible=False, add_book=False)
            app.screen_updating = False
            app.display_alerts = False
            load_wb = app.books.open(path2)
            load_ws = load_wb.sheets.active
            # 处理列，将指定列从大到小删除（避免先删除小列导致后续列号变动）
            load_ws.range('O1').api.EntireColumn.Delete()
            load_ws.range('N1').api.EntireColumn.Delete()
            load_ws.range('M1').api.EntireColumn.Delete()
            load_ws.range('L1').api.EntireColumn.Delete()
            load_ws.range('K1').api.EntireColumn.Delete()
            load_ws.range('J1').api.EntireColumn.Delete()
            load_ws.range('I1').api.EntireColumn.Delete()
            load_ws.range('H1').api.EntireColumn.Delete()

            # 获取行数
            info = load_ws.used_range
            last_row = info.last_cell.row
            alpha_dict = {1: 'A', 2: 'B', 3: 'C', 4: 'D', 5: 'E', 6: 'F', 7: 'G', 8: 'H', 9: 'I', 10: 'J', 11: 'K',
                          12: 'L', 13: 'M', 14: 'N', 15: 'O', 16: 'P', 17: 'Q', 18: 'R', 19: 'S', 20: 'T', 21: 'U',
                          22: 'V', 23: 'W', 24: 'X', 25: 'Y', 26: 'Z'}
            last_column_num = info.last_cell.column
            last_column = alpha_dict[last_column_num]
            # print(last_row, ' ', last_column)

            # 删除最后的空行，单数行情况
            last_row_value = load_ws.range(last_row, last_column_num).value
            if last_row_value == None:
                # load_ws.api.rows(last_row).delete
                load_ws.range(f'{last_column}{last_row}').api.EntireRow.Delete()
                last_row = str(int(last_row) - 1)

            # 设置边框
            a_range = f'A1:{last_column}{last_row}'  # 生成表格的数据范围
            load_ws.range(a_range).api.Borders(8).LineStyle = 1  # 上边框
            load_ws.range(a_range).api.Borders(9).LineStyle = 1  # 下边框
            load_ws.range(a_range).api.Borders(7).LineStyle = 1  # 左边框
            load_ws.range(a_range).api.Borders(10).LineStyle = 1  # 右边框
            load_ws.range(a_range).api.Borders(12).LineStyle = 1  # 内横边框
            load_ws.range(a_range).api.Borders(11).LineStyle = 1  # 内纵边框
            load_ws.range(a_range).api.Font.Name = 'Times New Roman'
            load_ws.range(a_range).api.RowHeight = 20

            load_ws.range(a_range).columns.autofit()

            b_range = load_ws.range(a_range)
            # 设置单元格 字体格式
            b_range.color = 255, 255, 255  # 设置单元格的填充颜色
            b_range.api.Font.ColorIndex = 1  # 设置字体的颜色，具体颜色索引见下方。
            b_range.api.Font.Size = 11  # 设置字体的大小。
            b_range.api.Font.Bold = False  # 设置为粗体。
            b_range.api.HorizontalAlignment = -4108  # -4108 水平居中。 -4131 靠左，-4152 靠右。
            b_range.api.VerticalAlignment = -4108  # -4108 垂直居中（默认）。 -4160 靠上，-4107 靠下， -4130 自动换行对齐。
            # b_range.api.NumberFormat = "0.00"          # 设置单元格的数字格式。

            # 处理完毕，保存、关闭、退出Excel
            load_wb.save()
            load_wb.close()
            app.quit()

            # if os.path.exists(path1):
            #     # 删除文件，可使用以下两种方法
            #     os.remove(path1)
            #     # os.unlink(my_file)
            # else:
            #     print('no such file:%s' % path1)

        ############################################### 若为一大列则进行直接进行规范化
        elif int(col_num) < 10:
            # 将操作文件对象拷贝，变成可写的workbook对象
            new_excel = copy(old_excel)
            # 获得第一个sheet的对象
            ws = new_excel.get_sheet(0)
            # 写入数据
            ws.write(2, 0, '解释序号')
            ws.write(2, 1, '井段(m)')
            ws.write(2, 2, '厚度(m)')
            ws.write(2, 3, '最大指数')
            ws.write(2, 4, '最小指数')
            ws.write(2, 5, '平均指数')
            ws.write(2, 6, '结论')
            # 另存为excel文件，并将文件命名
            new_excel.save(path2)

            ######################## 格式进一步规范
            # 创建app，打开工作表
            app = xw.App(visible=False, add_book=False)
            app.screen_updating = False
            app.display_alerts = False
            load_wb = app.books.open(path2)
            load_ws = load_wb.sheets.active

            # 获取行数
            info = load_ws.used_range
            last_row = info.last_cell.row
            alpha_dict = {1: 'A', 2: 'B', 3: 'C', 4: 'D', 5: 'E', 6: 'F', 7: 'G', 8: 'H', 9: 'I', 10: 'J', 11: 'K',
                          12: 'L', 13: 'M', 14: 'N', 15: 'O', 16: 'P', 17: 'Q', 18: 'R', 19: 'S', 20: 'T', 21: 'U',
                          22: 'V', 23: 'W', 24: 'X', 25: 'Y', 26: 'Z'}
            last_column_num = info.last_cell.column
            last_column = alpha_dict[last_column_num]
            # print(last_row, ' ', last_column)

            # 删除最后的空行，单数行情况
            last_row_value = load_ws.range(last_row, last_column_num).value
            if last_row_value == None:
                # load_ws.api.rows(last_row).delete
                load_ws.range(f'{last_column}{last_row}').api.EntireRow.Delete()
                last_row = str(int(last_row) - 1)

            # 设置边框
            a_range = f'A1:{last_column}{last_row}'  # 生成表格的数据范围
            load_ws.range(a_range).api.Borders(8).LineStyle = 1  # 上边框
            load_ws.range(a_range).api.Borders(9).LineStyle = 1  # 下边框
            load_ws.range(a_range).api.Borders(7).LineStyle = 1  # 左边框
            load_ws.range(a_range).api.Borders(10).LineStyle = 1  # 右边框
            load_ws.range(a_range).api.Borders(12).LineStyle = 1  # 内横边框
            load_ws.range(a_range).api.Borders(11).LineStyle = 1  # 内纵边框
            load_ws.range(a_range).api.Font.Name = 'Times New Roman'
            load_ws.range(a_range).api.RowHeight = 20

            load_ws.range(a_range).columns.autofit()

            b_range = load_ws.range(a_range)
            # 设置单元格 字体格式
            b_range.color = 255, 255, 255  # 设置单元格的填充颜色
            b_range.api.Font.ColorIndex = 1  # 设置字体的颜色，具体颜色索引见下方。
            b_range.api.Font.Size = 11  # 设置字体的大小。
            b_range.api.Font.Bold = False  # 设置为粗体。
            b_range.api.HorizontalAlignment = -4108  # -4108 水平居中。 -4131 靠左，-4152 靠右。
            b_range.api.VerticalAlignment = -4108  # -4108 垂直居中（默认）。 -4160 靠上，-4107 靠下， -4130 自动换行对齐。
            # b_range.api.NumberFormat = "0.00"          # 设置单元格的数字格式。

            # 处理完毕，保存、关闭、退出Excel
            load_wb.save()
            load_wb.close()
            app.quit()

        else:
            pass

    def btnstate_table(self, btn):
        # 输出按钮1与按钮2的状态，选中还是没选中
        if btn.text() == '一界面':
            if btn.isChecked() == True:
                print(btn.text() + " 被选中")
                self.pushButton_14.clicked.connect(self.calculate_for_first_layer)
                self.pushButton_27.clicked.connect(self.table_process1)
            else:
                pass

        if btn.text() == "二界面":
            if btn.isChecked() == True:
                print(btn.text() + " 被选中")
                self.pushButton_14.clicked.connect(self.calculate_for_second_layer)
                self.pushButton_27.clicked.connect(self.table_process2)
            else:
                pass

    def reset_table_process(self):
        try:
            self.radioButton.toggled.disconnect(lambda: self.btnstate_table(self.radioButton))
        except:
            print('Error1')
        else:
            print('Disconnected1')

        try:
            self.radioButton_2.toggled.disconnect(lambda: self.btnstate_table(self.radioButton_2))
        except:
            print('Error2')
        else:
            print('Disconnected2')

        try:
            self.pushButton_14.clicked.disconnect(self.calculate_for_first_layer)
        except:
            print('Error3')
        else:
            print('Disconnected3')

        try:
            self.pushButton_27.clicked.disconnect(self.table_process1)
        except:
            print('Error4')
        else:
            print('Disconnected4')

        try:
            self.pushButton_14.clicked.disconnect(self.calculate_for_second_layer)
        except:
            print('Error5')
        else:
            print('Disconnected5')

        try:
            self.pushButton_27.clicked.disconnect(self.table_process2)
        except:
            print('Error6')
        else:
            print('Disconnected6')
        self.radioButton.toggled.connect(lambda: self.btnstate_table(self.radioButton))
        self.radioButton_2.toggled.connect(lambda: self.btnstate_table(self.radioButton_2))

    def table_process1(self):
        fileDir1 = self.lineEdit_48.text()
        fileDir2 = self.lineEdit_52.text()
        self.xls_formatting_first_layer(fileDir1)
        if fileDir1 != fileDir2:
            self.xls_formatting_first_layer(fileDir2)
        QMessageBox.information(self, "提示", "一界面表格数据规范化完毕")

    def table_process2(self):
        fileDir1 = self.lineEdit_48.text()
        fileDir2 = self.lineEdit_52.text()
        self.xls_formatting_second_layer(fileDir1)
        if fileDir1 != fileDir2:
            self.xls_formatting_second_layer(fileDir2)
        QMessageBox.information(self, "提示", "二界面表格数据规范化完毕")

    ##############################

    # 规范化2
    ##############################
    def table_process3(self):
        if self.lineEdit_75.text() != '' and self.lineEdit_73.text() == '':
            fileDir1 = self.lineEdit_75.text()
            self.xls_formatting_first_layer(fileDir1)
            QMessageBox.information(self, "提示", "一界面表格数据规范化完毕")
        elif self.lineEdit_73.text() != '' and self.lineEdit_75.text() == '':
            fileDir2 = self.lineEdit_73.text()
            self.xls_formatting_second_layer(fileDir2)
            QMessageBox.information(self, "提示", "二界面表格数据规范化完毕")
        elif self.lineEdit_73.text() != '' and self.lineEdit_75.text() != '':
            fileDir1 = self.lineEdit_75.text()
            self.xls_formatting_first_layer(fileDir1)
            fileDir2 = self.lineEdit_73.text()
            self.xls_formatting_second_layer(fileDir2)
            QMessageBox.information(self, "提示", "一二界面表格数据都规范化完毕")

    def calculate_for_first_layer(self):
        splicing_Depth = float(self.lineEdit_45.text())

        fileDir1 = self.lineEdit_48.text()
        fileDir2 = self.lineEdit_52.text()

        df1 = pd.read_excel(fileDir1, header=2)
        # df1.drop([0], inplace=True)
        df1 = df1.dropna(axis=0, how='any')  # 删除dataframe里NaN的所有行
        df1.loc[:, '井段(m)'] = df1['井段(m)'].str.replace(' ', '')  # 消除数据中空格
        # if len(df1) % 2 == 0: # 如果len(df1)为偶数需要删除最后一行NaN，一行的情况不用删
        #     df1.drop([len(df1)], inplace=True)
        df1['井段Start'] = df1['井段(m)'].map(lambda x: x.split("-")[0])
        df1['井段End'] = df1['井段(m)'].map(lambda x: x.split("-")[1])
        # 表格数据清洗
        df1.loc[:, "井段Start"] = df1["井段Start"].str.replace(" ", "").astype('float')
        df1.loc[:, "井段End"] = df1["井段End"].str.replace(" ", "").astype('float')

        # 截取拼接点以上的数据体
        df_temp1 = df1.loc[(df1['井段Start'] <= splicing_Depth), :].copy()  # 加上copy()可防止直接修改df报错
        # print(df_temp1)

        #####################################################
        df2 = pd.read_excel(fileDir2, header=2)
        # df2.drop([0], inplace=True)
        df2 = df2.dropna(axis=0, how='any')  # 删除dataframe里NaN的所有行
        df2.loc[:, '井段(m)'] = df2['井段(m)'].str.replace(' ', '')  # 消除数据中空格
        # if len(df2) % 2 == 0:#如果len(df2)为偶数需要删除最后一行NaN，一行的情况不用删
        #     df2.drop([len(df2)], inplace=True)
        df2['井段Start'] = df2['井段(m)'].map(lambda x: x.split("-")[0])
        df2['井段End'] = df2['井段(m)'].map(lambda x: x.split("-")[1])
        # 表格数据清洗
        df2.loc[:, "井段Start"] = df2["井段Start"].str.replace(" ", "").astype('float')
        df2.loc[:, "井段End"] = df2["井段End"].str.replace(" ", "").astype('float')

        # 截取拼接点以下的数据体
        df_temp2 = df2.loc[(df2['井段End'] >= splicing_Depth), :].copy()  # 加上copy()可防止直接修改df报错
        df_temp2.reset_index(drop=True, inplace=True)  # 重新设置列索引

        # print(df_temp2)

        df_all = df_temp1.append(df_temp2)
        df_all.reset_index(drop=True, inplace=True)  # 重新设置列索引
        # 对df_all进行操作
        df_all.loc[len(df_temp1) - 1, '井段(m)'] = ''.join([str(df_all.loc[len(df_temp1) - 1, '井段Start']), '-', \
                                                            str(df_all.loc[len(df_temp1), '井段End'])])
        df_all.loc[len(df_temp1) - 1, '厚度(m)'] = df_all.loc[len(df_temp1), '井段End'] - \
                                                   df_all.loc[len(df_temp1) - 1, '井段Start']
        df_all.loc[len(df_temp1) - 1, '最大声幅(%)'] = max(df_all.loc[len(df_temp1), '最大声幅(%)'], \
                                                           df_all.loc[len(df_temp1) - 1, '最大声幅(%)'])
        df_all.loc[len(df_temp1) - 1, '最小声幅(%)'] = min(df_all.loc[len(df_temp1), '最小声幅(%)'], \
                                                           df_all.loc[len(df_temp1) - 1, '最小声幅(%)'])
        df_all.loc[len(df_temp1) - 1, '平均声幅(%)'] = np.add(df_all.loc[len(df_temp1), '平均声幅(%)'], \
                                                              df_all.loc[len(df_temp1) - 1, '平均声幅(%)']) / 2
        df_all.loc[len(df_temp1) - 1, '井段End'] = df_all.loc[len(df_temp1), '井段End']  # 解决后续重计算厚度计算bug
        df_all.drop(len(df_temp1), inplace=True)
        df_all.set_index(["解释序号"], inplace=True)
        df_all.reset_index(drop=True, inplace=True)  # 重新设置列索引
        # print(df_all.columns)

        #################################################################
        # 在指定深度段统计

        # calculation_Start = float(input('请输入开始统计深度'))
        # calculation_End = float(input('请输入结束统计深度'))
        calculation_Start = float(self.lineEdit_46.text())
        calculation_End = float(self.lineEdit_47.text())

        start_Evaluation = df_all.loc[0, '井段(m)'].split('-')[0]
        end_Evaluation = df_all.loc[len(df_all) - 1, '井段(m)'].split('-')[1]
        if (calculation_End <= float(end_Evaluation)) & (calculation_Start >= float(start_Evaluation)):
            df_temp = df_all.loc[(df_all['井段Start'] >= calculation_Start) & (df_all['井段Start'] <= calculation_End),
                      :]
            # 获取起始深度到第一层井段底界的结论
            df_temp_start_to_first_layer = df_all.loc[(df_all['井段Start'] <= calculation_Start), :]
            start_to_upper_result = df_temp_start_to_first_layer.loc[len(df_temp_start_to_first_layer) - 1, '结论']
            # 获取calculation_Start所在段的声幅值
            df_temp_calculation_Start = df_all.loc[(df_all['井段Start'] <= calculation_Start) & (
                    df_all['井段End'] >= calculation_Start), :]
            df_temp_calculation_Start.reset_index(drop=True, inplace=True)  # 重新设置列索引#防止若截取中段，index不从0开始的bug
            # 补充储层界到井段的深度
            x, y = df_temp.shape
            df_temp = df_temp.reset_index()
            df_temp.drop(['index'], axis=1, inplace=True)
            if x != 0:  # 防止df_temp为空时，loc报错的bug
                first_layer_start = df_temp.loc[0, '井段Start']
            if x > 1 and first_layer_start != calculation_Start:
                upper = pd.DataFrame({'井段(m)': ''.join([str(calculation_Start), '-', str(first_layer_start)]),
                                      '厚度(m)': first_layer_start - calculation_Start,
                                      '最大声幅(%)': df_temp_calculation_Start.loc[0, '最大声幅(%)'],
                                      '最小声幅(%)': df_temp_calculation_Start.loc[0, '最小声幅(%)'],
                                      '平均声幅(%)': df_temp_calculation_Start.loc[0, '平均声幅(%)'],
                                      '结论': start_to_upper_result,
                                      '井段Start': calculation_Start,
                                      '井段End': first_layer_start},
                                     index=[1])  # 自定义索引为：1 ，这里也可以不设置index
                df_temp.loc[len(df_temp) - 1, '井段End'] = calculation_End
                df_temp = pd.concat([upper, df_temp], ignore_index=True)
                # df_temp = df_temp.append(new, ignore_index=True)  # ignore_index=True,表示不按原来的索引，从0开始自动递增
                df_temp.loc[:, "重计算厚度"] = df_temp.apply(self.get_thickness, axis=1)
                # 修改df_temp的最末一行
                df_temp.loc[len(df_temp) - 1, '井段(m)'] = ''.join([str(df_temp.loc[len(df_temp) - 1, '井段Start']), \
                                                                    '-', str(df_temp.loc[len(df_temp) - 1, '井段End'])])
                df_temp.loc[len(df_temp) - 1, '厚度(m)'] = df_temp.loc[len(df_temp) - 1, '重计算厚度']
            elif x > 1 and first_layer_start == calculation_Start:
                df_temp.loc[len(df_temp) - 1, '井段End'] = calculation_End
                df_temp.loc[:, "重计算厚度"] = df_temp.apply(self.get_thickness, axis=1)
                # 修改df_temp的最末一行
                df_temp.loc[len(df_temp) - 1, '井段(m)'] = ''.join([str(df_temp.loc[len(df_temp) - 1, '井段Start']), \
                                                                    '-', str(df_temp.loc[len(df_temp) - 1, '井段End'])])
                df_temp.loc[len(df_temp) - 1, '厚度(m)'] = df_temp.loc[len(df_temp) - 1, '重计算厚度']
            else:  # 储层包含在一个井段内的情况
                df_temp = pd.DataFrame({'井段(m)': ''.join([str(calculation_Start), '-', str(calculation_End)]),
                                        '厚度(m)': calculation_End - calculation_Start,
                                        '最大声幅(%)': df_temp_calculation_Start.loc[0, '最大声幅(%)'],
                                        '最小声幅(%)': df_temp_calculation_Start.loc[0, '最小声幅(%)'],
                                        '平均声幅(%)': df_temp_calculation_Start.loc[0, '平均声幅(%)'],
                                        '结论': start_to_upper_result,
                                        '井段Start': calculation_Start,
                                        '井段End': calculation_End},
                                       index=[1])  # 自定义索引为：1 ，这里也可以不设置index
                df_temp.loc[:, "重计算厚度"] = df_temp.apply(self.get_thickness, axis=1)
                # 修改df_temp的最末一行
                df_temp.loc[len(df_temp), '井段(m)'] = ''.join([str(df_temp.loc[len(df_temp), '井段Start']),
                                                                '-', str(df_temp.loc[len(df_temp), '井段End'])])
                df_temp.loc[len(df_temp), '厚度(m)'] = df_temp.loc[len(df_temp), '重计算厚度']
            print(df_temp)
            ratio_Series = df_temp.groupby(by=['结论'])['重计算厚度'].sum() / df_temp['重计算厚度'].sum() * 100
            if ratio_Series.__len__() == 2:
                if '胶结好' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                elif '胶结中等' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                elif '胶结差' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
            elif ratio_Series.__len__() == 1:
                if ('胶结好' not in ratio_Series) & ('胶结中等' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                elif ('胶结好' not in ratio_Series) & ('胶结差' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
                elif ('胶结中等' not in ratio_Series) & ('胶结差' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
        # print(ratio_Series)

        # 统计结论
        actual_Hao = str(round((calculation_End - calculation_Start) * (ratio_Series['胶结好'] / 100), 2))
        Hao_Ratio = str(round(ratio_Series['胶结好'], 2))

        actual_Zhong = str(round((calculation_End - calculation_Start) * (ratio_Series['胶结中等'] / 100), 2))
        Zhong_Ratio = str(round(ratio_Series['胶结中等'], 2))

        actual_Cha = str(round(calculation_End - calculation_Start - float(actual_Hao) - float(actual_Zhong), 2))
        Cha_Ratio = str(round(100.00 - float(Hao_Ratio) - float(Zhong_Ratio), 2))

        PATH = '.\\resources\\模板\\'
        wb = openpyxl.load_workbook(PATH + '1统模板.xlsx')
        sheet = wb[wb.sheetnames[0]]
        sheet['A1'] = ''.join(['第一界面水泥胶结统计表（', str(calculation_Start), '-', str(calculation_End), 'm）'])
        sheet['C4'] = actual_Hao
        sheet['D4'] = Hao_Ratio
        sheet['C5'] = actual_Zhong
        sheet['D5'] = Zhong_Ratio
        sheet['C6'] = actual_Cha
        sheet['D6'] = Cha_Ratio

        self.mkdir('.\\WorkSpace\\合并统计工区')
        wb.save(
            '.\\WorkSpace\\合并统计工区\\一界面水泥胶结统计表(' + str(calculation_Start) + '-' + str(
                calculation_End) + 'm).xlsx')

        # 保存指定起始截止深度的单层统计表
        df_temp.drop(['井段Start', '井段End', '重计算厚度'], axis=1, inplace=True)
        df_temp.reset_index(drop=True, inplace=True)  # 重新设置列索引
        df_temp.index = df_temp.index + 1
        writer = pd.ExcelWriter(
            '.\\WorkSpace\\合并统计工区\\一界面水泥胶结单层评价表(' + str(calculation_Start) + '-' + str(
                calculation_End) + 'm).xlsx')
        df_temp.to_excel(writer, 'Sheet1')
        writer.save()

        QMessageBox.information(self, "提示", "运行完毕，请查看WorkSpace")

    def calculate_for_second_layer(self):
        splicing_Depth = float(self.lineEdit_45.text())

        fileDir1 = self.lineEdit_48.text()
        fileDir2 = self.lineEdit_52.text()

        df1 = pd.read_excel(fileDir1, header=2)
        # df1.drop([0], inplace=True)
        # if df1.loc[0, '结论'] == '不确定':
        #     df1.drop([0], inplace=True)
        df1 = df1.dropna(axis=0, how='any')  # 删除dataframe里NaN的所有行
        df1.loc[:, '井段(m)'] = df1['井段(m)'].str.replace(' ', '')  # 消除数据中空格
        # if len(df1) % 2 == 0:#如果len(df1)为偶数需要删除最后一行NaN，一行的情况不用删
        #     df1.drop([len(df1)], inplace=True)
        df1['井段Start'] = df1['井段(m)'].map(lambda x: x.split("-")[0])
        df1['井段End'] = df1['井段(m)'].map(lambda x: x.split("-")[1])
        # 表格数据清洗
        df1.loc[:, "井段Start"] = df1["井段Start"].str.replace(" ", "").astype('float')
        df1.loc[:, "井段End"] = df1["井段End"].str.replace(" ", "").astype('float')

        # 截取拼接点以上的数据体
        df_temp1 = df1.loc[(df1['井段Start'] <= splicing_Depth), :].copy()  # 加上copy()可防止直接修改df报错
        # print(df_temp1)

        #####################################################
        df2 = pd.read_excel(fileDir2, header=2)
        # df2.drop([0], inplace=True)
        df2 = df2.dropna(axis=0, how='any')  # 删除dataframe里NaN的所有行
        df2.loc[:, '井段(m)'] = df2['井段(m)'].str.replace(' ', '')  # 消除数据中空格
        # if len(df2) % 2 == 0:#如果len(df2)为偶数需要删除最后一行NaN，一行的情况不用删
        #     df2.drop([len(df2)], inplace=True)
        df2['井段Start'] = df2['井段(m)'].map(lambda x: x.split("-")[0])
        df2['井段End'] = df2['井段(m)'].map(lambda x: x.split("-")[1])
        # 表格数据清洗
        df2.loc[:, "井段Start"] = df2["井段Start"].str.replace(" ", "").astype('float')
        df2.loc[:, "井段End"] = df2["井段End"].str.replace(" ", "").astype('float')

        # 截取拼接点以下的数据体
        df_temp2 = df2.loc[(df2['井段End'] >= splicing_Depth), :].copy()  # 加上copy()可防止直接修改df报错
        df_temp2.reset_index(drop=True, inplace=True)  # 重新设置列索引

        # print(df_temp2)

        df_all = df_temp1.append(df_temp2)
        df_all.reset_index(drop=True, inplace=True)  # 重新设置列索引
        # 对df_all进行操作
        df_all.loc[len(df_temp1) - 1, '井段(m)'] = ''.join([str(df_all.loc[len(df_temp1) - 1, '井段Start']), '-', \
                                                            str(df_all.loc[len(df_temp1), '井段End'])])
        df_all.loc[len(df_temp1) - 1, '厚度(m)'] = df_all.loc[len(df_temp1), '井段End'] - \
                                                   df_all.loc[len(df_temp1) - 1, '井段Start']
        df_all.loc[len(df_temp1) - 1, '最大指数'] = max(df_all.loc[len(df_temp1), '最大指数'], \
                                                        df_all.loc[len(df_temp1) - 1, '最大指数'])
        df_all.loc[len(df_temp1) - 1, '最小指数'] = min(df_all.loc[len(df_temp1), '最小指数'], \
                                                        df_all.loc[len(df_temp1) - 1, '最小指数'])
        df_all.loc[len(df_temp1) - 1, '平均指数'] = np.add(df_all.loc[len(df_temp1), '平均指数'], \
                                                           df_all.loc[len(df_temp1) - 1, '平均指数']) / 2

        df_all.drop(len(df_temp1), inplace=True)
        df_all.set_index(["解释序号"], inplace=True)
        df_all.reset_index(drop=True, inplace=True)  # 重新设置列索引
        # print(df_all.columns)

        #################################################################
        # 在指定深度段统计

        # calculation_Start = float(input('请输入开始统计深度'))
        # calculation_End = float(input('请输入结束统计深度'))
        calculation_Start = float(self.lineEdit_46.text())
        calculation_End = float(self.lineEdit_47.text())

        start_Evaluation = df_all.loc[0, '井段(m)'].split('-')[0]
        end_Evaluation = df_all.loc[len(df_all) - 1, '井段(m)'].split('-')[1]
        if (calculation_End <= float(end_Evaluation)) & (calculation_Start >= float(start_Evaluation)):
            df_temp = df_all.loc[(df_all['井段Start'] >= calculation_Start) & (df_all['井段Start'] <= calculation_End),
                      :]
            # 获取起始深度到第一层井段底界的结论
            df_temp_start_to_first_layer = df_all.loc[(df_all['井段Start'] <= calculation_Start), :]
            start_to_upper_result = df_temp_start_to_first_layer.loc[len(df_temp_start_to_first_layer) - 1, '结论']
            # 获取calculation_Start所在段的声幅值
            df_temp_calculation_Start = df_all.loc[(df_all['井段Start'] <= calculation_Start) & (
                    df_all['井段End'] >= calculation_Start), :]
            df_temp_calculation_Start.reset_index(drop=True, inplace=True)  # 重新设置列索引#防止若截取中段，index不从0开始的bug
            # 补充储层界到井段的深度
            x, y = df_temp.shape
            df_temp = df_temp.reset_index()
            df_temp.drop(['index'], axis=1, inplace=True)
            if x != 0:  # 防止df_temp为空时，loc报错的bug
                first_layer_start = df_temp.loc[0, '井段Start']
            if x > 1 and first_layer_start != calculation_Start:
                upper = pd.DataFrame({'井段(m)': ''.join([str(calculation_Start), '-', str(first_layer_start)]),
                                      '厚度(m)': first_layer_start - calculation_Start,
                                      '最大指数': df_temp_calculation_Start.loc[0, '最大指数'],
                                      '最小指数': df_temp_calculation_Start.loc[0, '最小指数'],
                                      '平均指数': df_temp_calculation_Start.loc[0, '平均指数'],
                                      '结论': start_to_upper_result,
                                      '井段Start': calculation_Start,
                                      '井段End': first_layer_start},
                                     index=[1])  # 自定义索引为：1 ，这里也可以不设置index
                df_temp.loc[len(df_temp) - 1, '井段End'] = calculation_End
                df_temp = pd.concat([upper, df_temp], ignore_index=True)
                # df_temp = df_temp.append(new, ignore_index=True)  # ignore_index=True,表示不按原来的索引，从0开始自动递增
                df_temp.loc[:, "重计算厚度"] = df_temp.apply(self.get_thickness, axis=1)
                # 修改df_temp的最末一行
                df_temp.loc[len(df_temp) - 1, '井段(m)'] = ''.join([str(df_temp.loc[len(df_temp) - 1, '井段Start']), \
                                                                    '-',
                                                                    str(df_temp.loc[len(df_temp) - 1, '井段End'])])
                df_temp.loc[len(df_temp) - 1, '厚度(m)'] = df_temp.loc[len(df_temp) - 1, '重计算厚度']
            elif x > 1 and first_layer_start == calculation_Start:
                df_temp.loc[len(df_temp) - 1, '井段End'] = calculation_End
                df_temp.loc[:, "重计算厚度"] = df_temp.apply(self.get_thickness, axis=1)
                # 修改df_temp的最末一行
                df_temp.loc[len(df_temp) - 1, '井段(m)'] = ''.join([str(df_temp.loc[len(df_temp) - 1, '井段Start']), \
                                                                    '-',
                                                                    str(df_temp.loc[len(df_temp) - 1, '井段End'])])
                df_temp.loc[len(df_temp) - 1, '厚度(m)'] = df_temp.loc[len(df_temp) - 1, '重计算厚度']
            else:  # 储层包含在一个井段内的情况
                df_temp = pd.DataFrame({'井段(m)': ''.join([str(calculation_Start), '-', str(calculation_End)]),
                                        '厚度(m)': calculation_End - calculation_Start,
                                        '最大指数': df_temp_calculation_Start.loc[0, '最大指数'],
                                        '最小指数': df_temp_calculation_Start.loc[0, '最小指数'],
                                        '平均指数': df_temp_calculation_Start.loc[0, '平均指数'],
                                        '结论': start_to_upper_result,
                                        '井段Start': calculation_Start,
                                        '井段End': calculation_End},
                                       index=[1])  # 自定义索引为：1 ，这里也可以不设置index
                df_temp.loc[:, "重计算厚度"] = df_temp.apply(self.get_thickness, axis=1)
                # 修改df_temp的最末一行
                df_temp.loc[len(df_temp), '井段(m)'] = ''.join([str(df_temp.loc[len(df_temp), '井段Start']),
                                                                '-', str(df_temp.loc[len(df_temp), '井段End'])])
                df_temp.loc[len(df_temp), '厚度(m)'] = df_temp.loc[len(df_temp), '重计算厚度']
            print(df_temp)
            ratio_Series = df_temp.groupby(by=['结论'])['重计算厚度'].sum() / df_temp['重计算厚度'].sum() * 100
            if ratio_Series.__len__() == 2:
                if '胶结好' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                elif '胶结中等' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                elif '胶结差' not in ratio_Series:
                    ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
            elif ratio_Series.__len__() == 1:
                if ('胶结好' not in ratio_Series) & ('胶结中等' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                elif ('胶结好' not in ratio_Series) & ('胶结差' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
                elif ('胶结中等' not in ratio_Series) & ('胶结差' not in ratio_Series):
                    ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                    ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))

        # 统计结论
        actual_Hao = str(round((calculation_End - calculation_Start) * (ratio_Series['胶结好'] / 100), 2))
        Hao_Ratio = str(round(ratio_Series['胶结好'], 2))

        actual_Zhong = str(round((calculation_End - calculation_Start) * (ratio_Series['胶结中等'] / 100), 2))
        Zhong_Ratio = str(round(ratio_Series['胶结中等'], 2))

        actual_Cha = str(round(calculation_End - calculation_Start - float(actual_Hao) - float(actual_Zhong), 2))
        Cha_Ratio = str(round(100.00 - float(Hao_Ratio) - float(Zhong_Ratio), 2))

        PATH = '.\\resources\\模板\\'
        wb = openpyxl.load_workbook(PATH + '2统模板.xlsx')
        sheet = wb[wb.sheetnames[0]]
        sheet['A1'] = ''.join(['二界面水泥胶结统计表（', str(calculation_Start), '-', str(calculation_End), 'm）'])
        sheet['C4'] = actual_Hao
        sheet['D4'] = Hao_Ratio
        sheet['C5'] = actual_Zhong
        sheet['D5'] = Zhong_Ratio
        sheet['C6'] = actual_Cha
        sheet['D6'] = Cha_Ratio

        self.mkdir('.\\WorkSpace\\合并统计工区')
        wb.save(
            '.\\WorkSpace\\合并统计工区\\二界面水泥胶结统计表(' + str(calculation_Start) + '-' + str(
                calculation_End) + 'm).xlsx')

        # 保存指定起始截止深度的单层统计表
        df_temp.drop(['井段Start', '井段End', '重计算厚度'], axis=1, inplace=True)
        df_temp.reset_index(drop=True, inplace=True)  # 重新设置列索引
        df_temp.index = df_temp.index + 1
        writer = pd.ExcelWriter(
            '.\\WorkSpace\\合并统计工区\\二界面水泥胶结单层评价表(' + str(calculation_Start) + '-' + str(
                calculation_End) + 'm).xlsx')
        df_temp.to_excel(writer, 'Sheet1')
        writer.save()

        # 单层统计表保存为Excel
        # df_all.drop(['井段Start', '井段End'], axis=1, inplace=True)
        # df_all.index = df_all.index + 1
        # writer = pd.ExcelWriter(
        #     '.\\WorkSpace\\合并统计工区\\单层评价表(合并)(' + str(start_Evaluation) + '-' + str(end_Evaluation) + 'm).xlsx')
        # df_all.to_excel(writer, 'Sheet1')
        # writer.save()

        QMessageBox.information(self, "提示", "运行完毕，请查看WorkSpace")

    def open_table_process_directory(self):
        path = '.\\WorkSpace\\合并统计工区'
        if not os.path.exists(path):
            os.makedirs(path)
            print(path, ' has been created.')
        os.startfile(path)

    def open_table_fusion_directory(self):
        path = '.\\WorkSpace\\综合评价工区'
        if not os.path.exists(path):
            os.makedirs(path)
            print(path, ' has been created.')
        os.startfile(path)

    def table_fusion_reaction(self):
        fileDir1 = self.lineEdit_75.text()
        fileDir2 = self.lineEdit_73.text()

        # 获取一界面单层评价表的深度界限
        df1 = pd.read_excel(fileDir1, header=2)
        # df1.drop([0], inplace=True)
        df1 = df1.dropna(axis=0, how='any')  # 删除dataframe里NaN的所有行
        df1.loc[:, '井段(m)'] = df1['井段(m)'].str.replace(' ', '')  # 消除数据中空格
        # if len(df1) % 2 == 0:#如果len(df1)为偶数需要删除最后一行NaN，一行的情况不用删
        #     df1.drop([len(df1)], inplace=True)
        df1['井段Start'] = df1['井段(m)'].map(lambda x: x.split("-")[0])
        df1['井段End'] = df1['井段(m)'].map(lambda x: x.split("-")[1])
        # 表格数据清洗
        df1.loc[:, "井段Start"] = df1["井段Start"].str.replace(" ", "").astype('float')
        df1.loc[:, "井段End"] = df1["井段End"].str.replace(" ", "").astype('float')

        # 获取二界面单层评价表的深度界限
        df2 = pd.read_excel(fileDir2, header=2)
        # df2.drop([0], inplace=True)
        # if df2.loc[1, '结论'] == '不确定':
        #     df2.drop([1], inplace=True)
        df2 = df2.dropna(axis=0, how='any')  # 删除dataframe里NaN的所有行
        df2.loc[:, '井段(m)'] = df2['井段(m)'].str.replace(' ', '')  # 消除数据中空格
        # if len(df2) % 2 == 0:#如果len(df2)为偶数需要删除最后一行NaN，一行的情况不用删
        #     df2.drop([len(df2)], inplace=True)
        df2['井段Start'] = df2['井段(m)'].map(lambda x: x.split("-")[0])
        df2['井段End'] = df2['井段(m)'].map(lambda x: x.split("-")[1])
        # 表格数据清洗
        df2.loc[:, "井段Start"] = df2["井段Start"].str.replace(" ", "").astype('float')
        df2.loc[:, "井段End"] = df2["井段End"].str.replace(" ", "").astype('float')

        list1 = df1['井段Start'].values.tolist()
        list2 = df1['井段End'].values.tolist()
        list3 = df2['井段Start'].values.tolist()
        list4 = df2['井段End'].values.tolist()
        # list合并并去重
        for item in list2:
            if item not in list1:
                list1.append(item)
        for item in list3:
            if item not in list1:
                list1.append(item)
        for item in list4:
            if item not in list1:
                list1.append(item)
        list1.sort(key=lambda x: float(x))
        print(list1)
        data = pd.DataFrame()
        for i in range(0, len(list1) - 1):
            j = i + 1
            evaluation_of_formation1 = self.layer_evaluation1(df1, list1[i], list1[j])[1]  # 调取一界面评价函数
            evaluation_of_formation2 = self.layer_evaluation2(df2, list1[i], list1[j])[1]  # 调取二界面评价函数

            # 综合评价
            if evaluation_of_formation1 == '胶结好' and evaluation_of_formation2 in ['胶结好', '胶结中等', '不确定']:
                evaluation_of_formation = '胶结好'
                evaluation_of_formation_normal = '胶结好'
            elif evaluation_of_formation1 == '胶结中等' and evaluation_of_formation2 == '胶结好':
                evaluation_of_formation = '胶结好'
                evaluation_of_formation_normal = '胶结好'
            elif evaluation_of_formation1 == '胶结中等' and evaluation_of_formation2 in ['胶结中等', '不确定']:
                evaluation_of_formation = '胶结中等'
                evaluation_of_formation_normal = '胶结中等'
            elif evaluation_of_formation1 == '胶结差' or evaluation_of_formation2 == ['胶结差', '胶结中等', '不确定']:
                evaluation_of_formation = '胶结差'
                evaluation_of_formation_normal = '胶结差'

            thickness = round(list1[j] - list1[i], 2)
            interval = '-'.join([('%.2f' % list1[i]), ('%.2f' % list1[j])])
            print(interval, thickness, evaluation_of_formation1, evaluation_of_formation2, evaluation_of_formation,
                  evaluation_of_formation_normal, '\n')
            series = pd.Series({"井段(m)": interval, "厚度(m)": thickness, "一界面评价": evaluation_of_formation1, \
                                "二界面评价": evaluation_of_formation2, "综合评价": evaluation_of_formation,
                                "综合评价(好中差)": evaluation_of_formation_normal}, name=i + 1)
            data = data.append(series)
        # dataframe排序
        data = data[['井段(m)', '厚度(m)', '一界面评价', '二界面评价', '综合评价', '综合评价(好中差)']]
        print(data)

        # 获取开始结束深度
        workbook = xlrd.open_workbook(fileDir1)
        sheet = workbook.sheets()[0]

        # 获得表单的行数及列数
        nrow = sheet.nrows
        ncol = sheet.ncols
        # 处理评价井段
        calculation_Start = str(sheet.cell_value(3, 1)).strip()
        calculation_Start = calculation_Start.split('-')[0]
        calculation_End = str(sheet.cell_value(nrow - 1, 1)).strip('')
        calculation_End = ''.join(calculation_End.split())  # 去除所有空格
        calculation_End = calculation_End.split('-')[1]

        # 保存为excel
        writer = pd.ExcelWriter(
            '.\\WorkSpace\\综合评价工区\\综合评价表(' + str(calculation_Start) + '-' + str(calculation_End) + 'm).xlsx')
        data.to_excel(writer, 'Sheet1')
        writer.save()

        # 统计好中差比例
        df = pd.read_excel(
            '.\\WorkSpace\\综合评价工区\\综合评价表(' + str(calculation_Start) + '-' + str(calculation_End) + 'm).xlsx',
            header=0)
        df = df.reset_index()
        ratio_Series = df.groupby(by=['综合评价(好中差)'])['厚度(m)'].sum() / df['厚度(m)'].sum() * 100
        if ratio_Series.__len__() == 2:
            if '胶结好' not in ratio_Series:
                ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
            elif '胶结中等' not in ratio_Series:
                ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
            elif '胶结差' not in ratio_Series:
                ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
        elif ratio_Series.__len__() == 1:
            if ('胶结好' not in ratio_Series) & ('胶结中等' not in ratio_Series):
                ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
            elif ('胶结好' not in ratio_Series) & ('胶结差' not in ratio_Series):
                ratio_Series = ratio_Series.append(pd.Series({'胶结好': 0}))
                ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
            elif ('胶结中等' not in ratio_Series) & ('胶结差' not in ratio_Series):
                ratio_Series = ratio_Series.append(pd.Series({'胶结中等': 0}))
                ratio_Series = ratio_Series.append(pd.Series({'胶结差': 0}))
        # print(ratio_Series)

        wb = openpyxl.load_workbook(
            '.\\WorkSpace\\综合评价工区\\综合评价表(' + str(calculation_Start) + '-' + str(calculation_End) + 'm).xlsx')
        sheet = wb[wb.sheetnames[0]]
        max_row = sheet.max_row
        calculation_End2 = float(sheet['B' + str(max_row)].value.split('-')[1])
        calculation_Start2 = float(sheet['B2'].value.split('-')[0])
        # 统计结论
        actual_Hao = str(round((calculation_End2 - calculation_Start2) * (ratio_Series['胶结好'] / 100), 2))
        Hao_Ratio = str(round(ratio_Series['胶结好'], 2))

        actual_Zhong = str(round((calculation_End2 - calculation_Start2) * (ratio_Series['胶结中等'] / 100), 2))
        Zhong_Ratio = str(round(ratio_Series['胶结中等'], 2))

        actual_Cha = str(
            round(calculation_End2 - calculation_Start2 - float(actual_Hao) - float(actual_Zhong), 2))
        Cha_Ratio = str(round(100.00 - float(Hao_Ratio) - float(Zhong_Ratio), 2))

        PATH = '.\\resources\\模板\\'
        wb = openpyxl.load_workbook(PATH + '综合评价统计模板.xlsx')
        sheet = wb[wb.sheetnames[0]]
        sheet['A1'] = ''.join(['综合评价统计表（', str(calculation_Start2), '-', str(calculation_End2), 'm）'])
        sheet['B4'] = actual_Hao
        sheet['C4'] = Hao_Ratio
        sheet['B5'] = actual_Zhong
        sheet['C5'] = Zhong_Ratio
        sheet['B6'] = actual_Cha
        sheet['C6'] = Cha_Ratio

        self.mkdir('.\\WorkSpace\\综合评价工区')
        wb.save('.\\WorkSpace\\综合评价工区\\综合评价统计表(' + str(calculation_Start2) + '-' + str(
            calculation_End2) + 'm).xlsx')
        QMessageBox.information(self, "提示", "运行完毕，请查看工区")

    def table_process4(self):
        fileDir1 = self.lineEdit_77.text()
        fileDir2 = self.lineEdit_76.text()
        if fileDir1 != '':
            self.xls_formatting_first_layer(fileDir1)
            QMessageBox.information(self, "提示", "一界面表格数据规范化完毕")
        else:
            pass
        if fileDir2 != '':
            self.xls_formatting_second_layer(fileDir1)
            QMessageBox.information(self, "提示", "二界面表格数据规范化完毕")
        else:
            pass

    def search_for_statistic_result(self):
        start_depth = float(self.lineEdit_78.text())
        end_depth = float(self.lineEdit_79.text())

        if self.lineEdit_77.text() != '' and self.lineEdit_76.text() != '':
            QMessageBox.information(self, "提示", "暂不支持两个评价表同时统计，请删除一个后重试")
        elif self.lineEdit_77.text() != '':
            fileDir1 = self.lineEdit_77.text()

            # 获取一界面单层评价表的深度界限
            df1 = pd.read_excel(fileDir1, header=2)
            # df1.drop([0], inplace=True)
            df1 = df1.dropna(axis=0, how='any')  # 删除dataframe里NaN的所有行
            df1.loc[:, '井段(m)'] = df1['井段(m)'].str.replace(' ', '')  # 消除数据中空格
            # if len(df1) % 2 == 0:#如果len(df1)为偶数需要删除最后一行NaN，一行的情况不用删
            #     df1.drop([len(df1)], inplace=True)
            df1['井段Start'] = df1['井段(m)'].map(lambda x: x.split("-")[0])
            df1['井段End'] = df1['井段(m)'].map(lambda x: x.split("-")[1])
            # 表格数据清洗
            df1.loc[:, "井段Start"] = df1["井段Start"].str.replace(" ", "").astype('float')
            df1.loc[:, "井段End"] = df1["井段End"].str.replace(" ", "").astype('float')
            evaluation_of_formation1 = self.layer_evaluation1(df1, start_depth, end_depth)[0]  # 调取一界面评价函数
            self.lineEdit_83.setText(('%.2f' % evaluation_of_formation1['胶结好']))
            self.lineEdit_81.setText(('%.2f' % evaluation_of_formation1['胶结中等']))
            self.lineEdit_80.setText(('%.2f' % evaluation_of_formation1['胶结差']))
            not_sure = 100 - evaluation_of_formation1['胶结好'] - evaluation_of_formation1['胶结中等'] - \
                       evaluation_of_formation1[
                           '胶结差']
            self.lineEdit_82.setText(('%.2f' % not_sure))

        elif self.lineEdit_76.text() != '':
            fileDir2 = self.lineEdit_76.text()

            # 获取二界面单层评价表的深度界限
            df1 = pd.read_excel(fileDir2, header=2)
            # df1.drop([0], inplace=True)
            df1 = df1.dropna(axis=0, how='any')  # 删除dataframe里NaN的所有行
            df1.loc[:, '井段(m)'] = df1['井段(m)'].str.replace(' ', '')  # 消除数据中空格
            # if len(df1) % 2 == 0:#如果len(df1)为偶数需要删除最后一行NaN，一行的情况不用删
            #     df1.drop([len(df1)], inplace=True)
            df1['井段Start'] = df1['井段(m)'].map(lambda x: x.split("-")[0])
            df1['井段End'] = df1['井段(m)'].map(lambda x: x.split("-")[1])
            # 表格数据清洗
            df1.loc[:, "井段Start"] = df1["井段Start"].str.replace(" ", "").astype('float')
            df1.loc[:, "井段End"] = df1["井段End"].str.replace(" ", "").astype('float')
            evaluation_of_formation2 = self.layer_evaluation2(df1, start_depth, end_depth)[0]  # 调取二界面评价函数
            self.lineEdit_83.setText(('%.2f' % evaluation_of_formation2['胶结好']))
            self.lineEdit_81.setText(('%.2f' % evaluation_of_formation2['胶结中等']))
            self.lineEdit_80.setText(('%.2f' % evaluation_of_formation2['胶结差']))
            not_sure = 100 - evaluation_of_formation2['胶结好'] - evaluation_of_formation2['胶结中等'] - \
                       evaluation_of_formation2[
                           '胶结差']
            self.lineEdit_82.setText(('%.2f' % not_sure))

    def error_animation(self):  # 窗口错误摇摆提示
        animation = QPropertyAnimation(self)
        animation.setTargetObject(self.tabWidget_2)
        animation.setPropertyName(b"pos")
        animation.setKeyValueAt(0, self.tabWidget_2.pos())
        animation.setKeyValueAt(0.3, self.tabWidget_2.pos() + QPoint(-10, 0))
        animation.setKeyValueAt(0.5, self.tabWidget_2.pos())
        animation.setKeyValueAt(0.7, self.tabWidget_2.pos() + QPoint(10, 0))
        animation.setKeyValueAt(1, self.tabWidget_2.pos())
        animation.setDuration(120)
        animation.setLoopCount(3)
        animation.start(QAbstractAnimation.DeleteWhenStopped)


class Chain_Pane(QWidget, Ui_Form):
    def __init__(self):
        super(Chain_Pane, self).__init__()
        self.setupUi(self)
        # self.setWindowOpacity(0.9)
        self.setObjectName("proChain")

        self.pushButton.clicked.connect(Main_window_show)
        self.pushButton_3.clicked.connect(self.Check_updates)
        self.pushButton_4.clicked.connect(self.Save_to_FTP)
        self.pushButton_5.clicked.connect(self.Button_Rock_and_Roll)

    def Button_Rock_and_Roll(self):
        animation = QPropertyAnimation(self)
        animation.setTargetObject(self.pushButton_5)
        animation.setPropertyName(b"pos")
        animation.setKeyValueAt(0, self.pushButton_5.pos() + QPoint(-50, 0))
        animation.setKeyValueAt(0.05, self.pushButton_5.pos() + QPoint(-50, 0))
        animation.setKeyValueAt(0.1, self.pushButton_5.pos() + QPoint(-30, -30))
        animation.setKeyValueAt(0.15, self.pushButton_5.pos() + QPoint(0, -50))
        animation.setKeyValueAt(0.2, self.pushButton_5.pos() + QPoint(30, -30))
        animation.setKeyValueAt(0.25, self.pushButton_5.pos() + QPoint(50, 0))
        animation.setKeyValueAt(0.3, self.pushButton_5.pos() + QPoint(30, 30))
        animation.setKeyValueAt(0.35, self.pushButton_5.pos() + QPoint(0, 50))
        animation.setKeyValueAt(0.4, self.pushButton_5.pos() + QPoint(-30, 30))
        animation.setKeyValueAt(0.45, self.pushButton_5.pos() + QPoint(-50, 0))
        animation.setKeyValueAt(0.5, self.pushButton_5.pos() + QPoint(-30, -30))
        animation.setKeyValueAt(0.55, self.pushButton_5.pos() + QPoint(0, -50))
        animation.setKeyValueAt(0.6, self.pushButton_5.pos() + QPoint(30, -30))
        animation.setKeyValueAt(0.65, self.pushButton_5.pos() + QPoint(50, 0))
        animation.setKeyValueAt(0.7, self.pushButton_5.pos() + QPoint(30, 30))
        animation.setKeyValueAt(0.75, self.pushButton_5.pos() + QPoint(0, 50))
        animation.setKeyValueAt(0.8, self.pushButton_5.pos() + QPoint(-30, 30))
        animation.setKeyValueAt(0.85, self.pushButton_5.pos() + QPoint(-50, 0))
        animation.setKeyValueAt(0.9, self.pushButton_5.pos() + QPoint(-30, -30))
        animation.setKeyValueAt(0.95, self.pushButton_5.pos() + QPoint(0, -50))
        animation.setKeyValueAt(0.96, self.pushButton_5.pos() + QPoint(30, -30))
        animation.setKeyValueAt(0.97, self.pushButton_5.pos() + QPoint(50, 0))
        animation.setKeyValueAt(0.98, self.pushButton_5.pos() + QPoint(30, 30))
        animation.setKeyValueAt(0.99, self.pushButton_5.pos() + QPoint(0, 50))
        animation.setKeyValueAt(0.995, self.pushButton_5.pos() + QPoint(-30, 30))
        animation.setKeyValueAt(0.999, self.pushButton_5.pos() + QPoint(-50, 0))

        animation.setKeyValueAt(1, self.pushButton_5.pos() + QPoint(0, 0))

        animation.setDuration(300)
        animation.setLoopCount(1)
        animation.start(QAbstractAnimation.DeleteWhenStopped)

    def Save_to_FTP(self):
        try:  # 将工区备份至FTP
            self.auto_upload_to_FTP()
        except:
            print('Faild to connect to oracle_data')

    def auto_upload_to_FTP(self):
        ftp = MyFTP('10.132.203.206')
        ftp.Login('zonghs', 'zonghs123')
        local_path = './WorkSpace'
        # local_path = r'C:\Users\YANGYI\source\repos\GC_Logging_Helper_Release'
        remote_path = '/oracle_data9/arc_data/SGI1/2016年油套管检测归档/工区备份'

        # 备份文件夹改名
        myname = socket.getfqdn(socket.gethostname())  # 获取本机电脑名
        myaddr = socket.gethostbyname(myname)  # 获取本机ip
        myaddr = myaddr.replace('.', '-')
        timeStr = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
        timeStr = timeStr.replace(':', '-').replace(' ', '-')
        # print(remote_path + '/' + timeStr + '_' + myaddr + '_' + myname)
        ftp.Mkd(remote_path + '/' + timeStr + '_' + myaddr + '_' + myname)
        ftp.UpLoadFileTree(local_path, remote_path + '/' + timeStr + '_' + myaddr + '_' + myname)

        QMessageBox.information(self, "提示", "工区文件已备份至FTP.")

    def Check_updates(self):
        try:
            # 先检查更新
            PATH = ".\\"
            listdir = []

            ftp = MyFTP('10.132.203.206')
            ftp.Login('zonghs', 'zonghs123')
            local_path = './'
            # local_path = r'C:\Users\YANGYI\source\repos\GC_Logging_Helper_Release'
            remote_path = '/oracle_data9/arc_data/SGI1/2016年油套管检测归档/工程测井助手最新版本(全部更新)'

            # 打开本地版本号
            with open(local_path + 'resources/版本号.txt', "r") as f:
                license_str = f.read()
            local_license_date = int(license_str)

            # 打开服务器版本号
            ftp.Cwd(remote_path)
            filenames = ftp.Nlst()
            filename = 'resources/版本号.txt'
            LocalFile = local_path + 'temp/版本号.txt'
            RemoteFile = filename

            # 接收服务器上文件并写入本地文件
            if not os.path.exists(local_path + 'temp'):
                os.makedirs(local_path + 'temp')
            ftp.DownLoadFile(LocalFile, RemoteFile)

            with open(local_path + 'temp/版本号.txt', "r") as f:
                license_str = f.read()
            remote_license_date = int(license_str)

            # 比较版本号信息
            if local_license_date < remote_license_date:
                Chain.label.setText("Download to update.")

            elif local_license_date >= remote_license_date:
                Chain.label.setText("All have been updated.")

        except:
            Chain.label.setText("Connect failed, continue?")
            pass


def Main_window_show():
    main.show()


def show_r():
    animation = QPropertyAnimation(app)
    animation.setTargetObject(main)
    animation.setPropertyName(b"pos")
    # randon_num1 = random.randint(0, 2000)
    # randon_num2 = random.randint(0, 2000)
    animation.setStartValue(QPoint(200, 0))
    animation.setEndValue(QPoint(200, 200))
    animation.setDuration(3500)
    animation.setEasingCurve(QEasingCurve.OutBounce)
    animation.start()


import resource_rc

if __name__ == "__main__":
    # 运行主程序
    QApplication.setAttribute(
        QtCore.Qt.AA_EnableHighDpiScaling)  # Enables high-DPI scaling in Qt on supported platforms
    app = QApplication(sys.argv)
    QApplication.setStyle(QStyleFactory.create("Fusion"))
    Chain = Chain_Pane()
    # main.show()
    Chain.show()
    main = Main_window()
    # Chain.label.setText("未检查更新。直接启动?")
    Chain.label.setText("本地软件版本已经是最新")
    show_r()  # Main_window进场动画
    sys.exit(app.exec_())
