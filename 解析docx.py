# coding=utf-8
import os
from docx import Document

# 开始解析docx
PATH = "Others/Examples/测试数据-报告生成/顺探1原始资料收集登记表_20240120.docx"
document = Document(PATH)

# 打印所有段落
print('一共有', str(len(document.paragraphs) - 1), '个文本段落。')
count = 0
for paragraph in document.paragraphs:
    print('第', str(count), '个段落')
    # print(paragraph.text)  # 打印各段落内容文本
    count += 1

# 打印所有表格
print('一共有', str(len(document.tables) - 1), '个表。')
count = 0
for table in document.tables:
    print('第', str(count), '个表')
    count += 1
    for row in range(len(table.rows)):
        for col in range(len(table.columns)):
            table.cell(row, col).text += '({0},{1})'.format(row, col)#给文本中的单元格添加表格坐标
            print('(', str(row), ',', str(col), '):', table.cell(row, col).text)

table = document.tables[2]
for row in range(len(table.rows)):
        for col in range(len(table.columns)):
            table.cell(row, col).text += '({0},{1})'.format(row, col)#给文本中的单元格添加表格坐标
            print('(', str(row), ',', str(col), '):', table.cell(row, col).text)